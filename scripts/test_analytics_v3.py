from __future__ import annotations

"""Behavioral tests for Analytics exact-latest methodology and narrative contract.

Run from repository root:
    python -m unittest scripts.test_analytics_v3 -v
"""

import unittest
from unittest.mock import patch
import ast
from io import BytesIO
from pathlib import Path

import pandas as pd

from core import analytics_calculations as calc
from core.analytics_text import (
    build_context, generate_analytics_note, detect_signals, derive_findings,
    SUPPORTED_SIGNAL_CODES, SUPPORTED_FINDING_CODES, SCENARIOS, QUESTIONS,
)
from core.analytics_text.models import AnalyticalFinding
from core.analytics_text.validation import validate_text, validate_finding_numeric_provenance
from core.analytics_text.analytical_metrics import _Builder, metric_code_of


def _snapshot(*rows):
    defaults = {
        "code": "1.1.1", "parent_goal_code": "1.", "parent_goal_name": "Ціль 1",
        "parent_task_code": "1.1.", "parent_task_name": "Завдання 1.1",
        "product_type": "НПА", "main_ssp": "1", "department": "ССП 1",
        "deputy_minister_by_ssp": "Заступник 1", "submitted": True,
        "missing_required_submission": False, "preliminary_attention": False,
        "included_in_risk_assessment": False, "risk_level": "",
        "forecast_kind": "", "result_achieved": False, "execution_score": 50.0,
        "coverage_eligible": True,
    }
    data = []
    for row in rows or ({},):
        item = dict(defaults); item.update(row); data.append(item)
    return pd.DataFrame(data)


def _goal_scores(value, coverage=100.0, code="1.", name="Ціль 1"):
    return pd.DataFrame([{"goal_code": code, "goal_name": name, "by_tasks": value, "coverage": coverage}])


def _task_scores(value, coverage=100.0, code="1.1.", name="Завдання 1.1"):
    return pd.DataFrame([{"goal_code": "1.", "task_code": code, "task_name": name, "execution": value, "coverage": coverage}])


def _period(execution, coverage, *, goal=None, snapshot=None):
    snap = _snapshot() if snapshot is None else snapshot
    goal_value = execution if goal is None else goal
    return {
        "snapshot": snap,
        "execution_by_measures": execution,
        "execution_by_goals": goal_value,
        "coverage": coverage,
        "goal_scores": _goal_scores(goal_value, coverage),
        "task_scores": _task_scores(execution, coverage),
        "risk_summary": {},
    }


class ExactLatestExecutionTests(unittest.TestCase):
    def test_single_period(self):
        plan = calc.build_analytics_plan_summary({(2026, "I"): _period(41.0, 80.0)})
        self.assertEqual(plan["execution_by_measures"], 41.0)
        self.assertIsNone(plan["execution_by_measures_change"])

    def test_i_ii_uses_ii_not_average(self):
        results = {(2026, "I"): _period(20.0, 70.0), (2026, "II"): _period(60.0, 90.0)}
        plan = calc.build_analytics_plan_summary(results)
        self.assertEqual(plan["execution_by_measures"], 60.0)
        self.assertEqual(plan["execution_by_measures_change"], 40.0)
        self.assertEqual(plan["coverage_average"], 80.0)
        self.assertEqual(plan["coverage_latest"], 90.0)
        self.assertNotIn("execution_by_measures_average", plan)

    def test_non_contiguous_uses_chronological_latest(self):
        results = {(2026, "I"): _period(10.0, 40.0), (2026, "III"): _period(73.0, 95.0)}
        self.assertEqual(calc.build_analytics_plan_summary(results)["execution_by_measures"], 73.0)

    def test_multiple_years_uses_latest_chronological_period(self):
        results = {(2026, "IV"): _period(55.0, 80.0), (2027, "II"): _period(68.0, 90.0)}
        plan = calc.build_analytics_plan_summary(results)
        self.assertEqual(plan["latest_period"], (2027, "II"))
        self.assertEqual(plan["execution_by_measures"], 68.0)

    def test_latest_missing_does_not_carry_forward(self):
        results = {(2026, "I"): _period(45.0, 80.0), (2026, "II"): _period(None, 90.0)}
        plan = calc.build_analytics_plan_summary(results)
        self.assertIsNone(plan["execution_by_measures"])
        self.assertIsNone(plan["execution_by_measures_change"])

    def test_legitimate_zero_latest_is_preserved(self):
        results = {(2026, "I"): _period(45.0, 80.0), (2026, "II"): _period(0.0, 0.0)}
        plan = calc.build_analytics_plan_summary(results)
        self.assertEqual(plan["execution_by_measures"], 0.0)
        self.assertEqual(plan["coverage_latest"], 0.0)

    def test_overachievement_is_not_silently_rewritten_by_adapter(self):
        results = {(2026, "II"): _period(120.0, 100.0)}
        self.assertEqual(calc.build_analytics_plan_summary(results)["execution_by_measures"], 120.0)

    def test_goal_and_task_exact_latest_have_no_carry(self):
        results = {
            (2026, "I"): _period(50.0, 80.0, goal=60.0),
            (2026, "II"): {
                **_period(None, 90.0, goal=None),
                "goal_scores": _goal_scores(None, 90.0),
                "task_scores": _task_scores(None, 90.0),
            },
        }
        active = calc.snapshot_rows_from_period_results(results)
        goals = calc.build_analytics_goal_summary(results, active)
        tasks = calc.build_analytics_task_summary(results, active)
        self.assertTrue(pd.isna(goals.loc[0, "Виконання"]))
        self.assertTrue(pd.isna(tasks.loc[0, "Виконання"]))


class CoverageTests(unittest.TestCase):
    def test_average_and_latest_are_separate(self):
        results = {
            (2026, "I"): _period(20.0, 50.0),
            (2026, "II"): _period(40.0, None),
            (2026, "III"): _period(60.0, 100.0),
        }
        plan = calc.build_analytics_plan_summary(results)
        self.assertEqual(plan["coverage_average"], 75.0)
        self.assertEqual(plan["coverage_latest"], 100.0)

    def test_latest_unavailable_is_not_previous_coverage(self):
        results = {(2026, "I"): _period(20.0, 50.0), (2026, "II"): _period(40.0, None)}
        plan = calc.build_analytics_plan_summary(results)
        self.assertEqual(plan["coverage_average"], 50.0)
        self.assertIsNone(plan["coverage_latest"])


class ManagementAttentionTests(unittest.TestCase):
    def test_q1_preliminary_only(self):
        snap = _snapshot(
            {"code":"A","preliminary_attention":True,"risk_level":"Критичний ризик"},
            {"code":"B","preliminary_attention":False,"risk_level":"Критичний ризик"},
        )
        info = calc.management_attention_info({(2026,"I"):_period(30,80,snapshot=snap)})
        self.assertEqual(info["type"], "preliminary_attention")
        self.assertEqual(info["count"], 1)

    def test_q2_high_and_critical_only(self):
        snap = _snapshot(
            {"code":"A","included_in_risk_assessment":True,"risk_level":"Високий ризик"},
            {"code":"B","included_in_risk_assessment":True,"risk_level":"Критичний ризик"},
            {"code":"C","included_in_risk_assessment":True,"risk_level":"Середній ризик"},
        )
        info = calc.management_attention_info({(2026,"II"):_period(40,90,snapshot=snap)})
        self.assertEqual(info["count"], 2)
        self.assertEqual(info["assessed_count"], 3)

    def test_q3_latest_deduplicates_code_and_ignores_q2_history(self):
        q2 = _snapshot({"code":"A","included_in_risk_assessment":True,"risk_level":"Критичний ризик"})
        q3 = _snapshot(
            {"code":"A","included_in_risk_assessment":True,"risk_level":"Високий ризик"},
            {"code":"A","included_in_risk_assessment":True,"risk_level":"Високий ризик"},
        )
        results = {(2026,"II"):_period(40,80,snapshot=q2), (2026,"III"):_period(50,90,snapshot=q3)}
        self.assertEqual(calc.management_attention_info(results)["count"], 1)

    def test_q4_only_assessed_final_nonachievement(self):
        snap = _snapshot(
            {"code":"A","forecast_kind":"final","execution_score":50.0,"result_achieved":False},
            {"code":"B","forecast_kind":"final","execution_score":100.0,"result_achieved":True},
            {"code":"C","forecast_kind":"final","execution_score":None,"result_achieved":False},
            {"code":"D","forecast_kind":"forecast","execution_score":50.0,"result_achieved":False},
        )
        info = calc.management_attention_info({(2026,"IV"):_period(50,90,snapshot=snap)})
        self.assertEqual(info["type"], "final_nonachievement")
        self.assertEqual(info["count"], 1)
        self.assertEqual(info["assessed_count"], 2)


class DataContractTests(unittest.TestCase):
    def test_metrics_have_no_accumulated_problem(self):
        results = {(2026,"I"):_period(20,50), (2026,"II"):_period(30,60)}
        active = calc.snapshot_rows_from_period_results(results)
        metrics = calc.build_metrics(active, results)
        self.assertNotIn("problem", metrics)
        self.assertNotIn("is_problem_status", active.columns)

    def test_ssp_exact_latest_does_not_use_last_present_group_period(self):
        results = {
            (2026, "I"): _period(50.0, 80.0),
            (2026, "II"): {**_period(None, None), "snapshot": pd.DataFrame()},
        }
        shared = pd.DataFrame([{
            "ssp": "1", "average": 50.0, "latest": 50.0, "change": None,
            "average_coverage": 80.0, "latest_coverage": 80.0, "latest_period": (2026, "I"),
        }])
        period_frame = pd.DataFrame([{
            "ssp": "1", "year": 2026, "quarter": "I", "execution": 50.0, "coverage": 80.0,
        }])
        out = calc._override_group_summary_exact_latest(
            shared, period_frame, results, group_col="ssp"
        )
        self.assertTrue(pd.isna(out.loc[0, "latest"]))
        self.assertTrue(pd.isna(out.loc[0, "latest_coverage"]))
        self.assertEqual(out.loc[0, "latest_period"], (2026, "II"))

    def test_deputy_api_has_exact_latest_contract_and_no_average(self):
        results = {
            (2026, "I"): _period(50.0, 80.0),
            (2026, "II"): {**_period(None, None), "snapshot": pd.DataFrame()},
        }
        shared = pd.DataFrame([{
            "deputy": "Заступник 1", "average": 50.0, "latest": 50.0, "change": None,
            "average_coverage": 80.0, "latest_coverage": 80.0, "latest_period": (2026, "I"),
        }])
        period_frame = pd.DataFrame([{
            "deputy": "Заступник 1", "year": 2026, "quarter": "I", "execution": 50.0, "coverage": 80.0,
        }])
        with patch.object(calc, "_dashboard_deputy_summary", return_value=shared), \
             patch.object(calc, "deputy_period_frame", return_value=period_frame):
            out = calc.build_analytics_deputy_summary(results)
        self.assertNotIn("average", out.columns)
        self.assertTrue(pd.isna(out.loc[0, "Виконання"]))
        self.assertEqual(out.loc[0, "latest_period"], (2026, "II"))

    def test_latest_breakdown_uses_current_attention(self):
        q1 = _snapshot({"code":"A","parent_goal_code":"1.","parent_goal_name":"Ціль 1","preliminary_attention":True})
        q2 = _snapshot(
            {"code":"A","parent_goal_code":"1.","parent_goal_name":"Ціль 1","included_in_risk_assessment":True,"risk_level":"Середній ризик"},
            {"code":"B","parent_goal_code":"1.","parent_goal_name":"Ціль 1","included_in_risk_assessment":True,"risk_level":"Критичний ризик"},
        )
        results={(2026,"I"):_period(20,50,snapshot=q1),(2026,"II"):_period(40,80,snapshot=q2)}
        active=calc.snapshot_rows_from_period_results(results)
        goals=calc.build_analytics_goal_summary(results,active)
        self.assertEqual(int(goals.loc[0,"Актуальна_увага"]),1)
        self.assertNotIn("Проблемних", goals.columns)


class NarrativeTests(unittest.TestCase):
    def _context(self, *, completion=72.0, coverage=90.0, coverage_latest=95.0, attention=2, attention_type="forecast_risk", dynamics=(60.0,72.0)):
        metrics={
            "completion":completion,"goal_completion":70.0,
            "coverage":coverage,"coverage_latest":coverage_latest,
            "completion_change":12.0 if len(dynamics)>1 else None,"coverage_change":5.0 if len(dynamics)>1 else None,
            "total_rows":10,"unique_measures":10,"latest_measure_count":10,"goals":2,"tasks":3,"no_data":1,
            "attention_count":attention,"attention_type":attention_type,"attention_label":"Високий або критичний ризик",
            "attention_assessed_count":8,"latest_risk_summary":{},
        }
        goals=pd.DataFrame([
            {"goal_code":"1.","strategic_goal":"Ціль 1","Виконання":45.0,"Зміна":-5.0,"Покриття_середнє_%":80.0,"Покриття_останній_%":85.0,"Унікальних_заходів":5,"Актуальна_увага":2,"Без_даних":1},
            {"goal_code":"2.","strategic_goal":"Ціль 2","Виконання":95.0,"Зміна":8.0,"Покриття_середнє_%":100.0,"Покриття_останній_%":100.0,"Унікальних_заходів":5,"Актуальна_увага":0,"Без_даних":0},
        ])
        tasks=pd.DataFrame([
            {"goal_code":"1.","task_code":"1.1.","task_name":"Завдання А","Виконання":30.0,"Зміна":-10.0,"Актуальна_увага":2,"Без_даних":1},
            {"goal_code":"1.","task_code":"1.2.","task_name":"Завдання Б","Виконання":70.0,"Зміна":0.0,"Актуальна_увага":0,"Без_даних":0},
        ])
        departments=pd.DataFrame([
            {"ssp_index":"1","department":"ССП 1","Виконання":50.0,"Зміна":-3.0,"Покриття_середнє_%":85.0,"Покриття_останній_%":90.0,"Унікальних_заходів":6,"Актуальна_увага":2,"Без_даних":1,"portfolio_weight_pct":60.0,"underperformance_contribution_pct":75.0},
            {"ssp_index":"2","department":"ССП 2","Виконання":90.0,"Зміна":5.0,"Покриття_середнє_%":100.0,"Покриття_останній_%":100.0,"Унікальних_заходів":4,"Актуальна_увага":0,"Без_даних":0,"portfolio_weight_pct":40.0,"underperformance_contribution_pct":25.0},
        ])
        products=pd.DataFrame([{"product_type":"НПА","Унікальних_заходів":10,"Виконання":72.0,"Актуальна_увага":2,"Без_даних":1}])
        statuses=pd.DataFrame([{"status":"Виконано","Кількість":8},{"status":"Частково виконано","Кількість":2}])
        period_rows=[]
        for i,value in enumerate(dynamics):
            q=("I","II","III","IV")[i]
            period_rows.append({"report_year":2026,"report_quarter":q,"report_quarter_num":i+1,"Період":f"2026 {q}","Виконання":value,"Покриття_%":coverage if i==0 else coverage_latest})
        active=pd.DataFrame([{"code":str(i),"report_year":2026,"report_quarter":"I","missing_required_submission":False} for i in range(10)])
        return build_context(
            filters={"years":[2026],"quarters":[("I","II")[0],("I","II")[1]],"ssp":[],"ssp_indices":[],"deputies":[],"goal_labels":[],"task_labels":[],"product_types":[]},
            metrics=metrics,goal_progress=goals,task_progress=tasks,department_progress=departments,
            product_progress=products,status_counts=statuses,period_dynamics=pd.DataFrame(period_rows),active=active,
            mio_goal_evaluation=pd.DataFrame(),mio_goal_task_evaluation=pd.DataFrame(),mio_measure_evaluation=pd.DataFrame(),mio_financing=pd.DataFrame(),
        )

    def test_deterministic_and_no_obsolete_concepts(self):
        ctx=self._context()
        a=generate_analytics_note(context=ctx); b=generate_analytics_note(context=ctx)
        self.assertEqual(a,b)
        low=a.lower()
        self.assertNotIn("рік до року",low)
        self.assertNotIn("середнє виконання за квартал",low)
        self.assertNotIn("проблемних позицій",low)

    def test_drilldown_and_management_implication_present_for_salient_local_problem(self):
        text=generate_analytics_note(context=self._context()).lower()
        self.assertTrue("завдання а" in text or "1.1." in text)
        self.assertTrue("управлін" in text or "уваг" in text)

    def test_small_context_does_not_force_long_paragraph_quota(self):
        ctx=self._context(attention=0,dynamics=(90.0,))
        result=generate_analytics_note(context=ctx,debug=True)
        paragraphs=[p for p in result.text.split("\n\n") if p.strip()]
        self.assertTrue(paragraphs)
        self.assertIsNone(result.debug.target_paragraph_count)
        self.assertEqual(result.debug.important_findings_skipped,[])


class AcceptanceMatrixTests(unittest.TestCase):
    def test_execution_selection_matrix(self):
        cases = [
            ({(2026, "I"):_period(11,70)}, 11),
            ({(2026, "I"):_period(11,70),(2026,"II"):_period(22,80)},22),
            ({(2026, "I"):_period(11,70),(2026,"II"):_period(22,80),(2026,"III"):_period(33,90)},33),
            ({(2026, "I"):_period(11,70),(2026,"II"):_period(22,80),(2026,"III"):_period(33,90),(2026,"IV"):_period(44,100)},44),
            ({(2026, "II"):_period(22,80),(2026,"III"):_period(33,90)},33),
            ({(2026, "I"):_period(11,70),(2026,"III"):_period(33,90)},33),
            ({(2026, "IV"):_period(44,100),(2027,"I"):_period(55,90)},55),
        ]
        for results, expected in cases:
            with self.subTest(expected=expected, periods=list(results)):
                self.assertEqual(calc.build_analytics_plan_summary(results)["execution_by_measures"], expected)

    def test_qualitative_and_yes_no_scores_are_passed_through_exactly(self):
        # Analytics does not reinterpret the canonical Dashboard score; it uses the exact latest prepared KPI.
        for canonical in (0.0, 75.0, 100.0):
            with self.subTest(canonical=canonical):
                result=calc.build_analytics_plan_summary({(2026,"IV"):_period(canonical,100.0)})
                self.assertEqual(result["execution_by_measures"],canonical)

    def test_attention_semantics_matrix(self):
        q1_none=_snapshot({"code":"A","preliminary_attention":False})
        self.assertEqual(calc.management_attention_info({(2026,"I"):_period(10,80,snapshot=q1_none)})["count"],0)
        q2_medium=_snapshot({"code":"A","included_in_risk_assessment":True,"risk_level":"Середній ризик"})
        self.assertEqual(calc.management_attention_info({(2026,"II"):_period(20,80,snapshot=q2_medium)})["count"],0)
        q3_critical=_snapshot({"code":"A","included_in_risk_assessment":True,"risk_level":"Критичний ризик"})
        info=calc.management_attention_info({(2026,"III"):_period(30,80,snapshot=q3_critical)})
        self.assertEqual((info["type"],info["count"]),("forecast_risk",1))
        q4_achieved=_snapshot({"code":"A","forecast_kind":"final","execution_score":100.0,"result_achieved":True})
        self.assertEqual(calc.management_attention_info({(2026,"IV"):_period(100,100,snapshot=q4_achieved)})["count"],0)
        q4_unassessed=_snapshot({"code":"A","forecast_kind":"final","execution_score":None,"result_achieved":False})
        self.assertEqual(calc.management_attention_info({(2026,"IV"):_period(None,100,snapshot=q4_unassessed)})["count"],0)

    def test_coverage_boundary_matrix(self):
        single=calc.build_analytics_plan_summary({(2026,"I"):_period(10,0.0)})
        self.assertEqual(single["coverage_average"],0.0); self.assertEqual(single["coverage_latest"],0.0)
        full=calc.build_analytics_plan_summary({(2026,"I"):_period(10,100.0),(2026,"II"):_period(20,100.0)})
        self.assertEqual(full["coverage_average"],100.0); self.assertEqual(full["coverage_latest"],100.0)
        mixed=calc.build_analytics_plan_summary({(2026,"I"):_period(10,0.0),(2026,"II"):_period(20,None),(2026,"III"):_period(30,100.0)})
        self.assertEqual(mixed["coverage_average"],50.0); self.assertEqual(mixed["coverage_latest"],100.0)

    def test_latest_missing_breakdowns_do_not_carry_old_object_values(self):
        results={
            (2026,"I"):_period(80,80,goal=90),
            (2026,"II"):{**_period(None,90,goal=None),"goal_scores":_goal_scores(None,90),"task_scores":_task_scores(None,90)},
        }
        active=calc.snapshot_rows_from_period_results(results)
        self.assertTrue(pd.isna(calc.build_analytics_goal_summary(results,active).iloc[0]["Виконання"]))
        self.assertTrue(pd.isna(calc.build_analytics_task_summary(results,active).iloc[0]["Виконання"]))
        products=calc.aggregate_product_progress(results,active)
        self.assertTrue(pd.isna(products.iloc[0]["Виконання"]))


class NarrativeAcceptanceTests(NarrativeTests):
    def test_debug_provenance_and_no_fatal_warnings(self):
        result=generate_analytics_note(context=self._context(),debug=True)
        self.assertTrue(result.text.strip())
        self.assertFalse([w for w in result.debug.validation_warnings if w.startswith(("polarity/sign contradiction","finding contradiction","unsupported numeric tokens","obsolete Analytics narrative concept"))])
        self.assertTrue(result.debug.numeric_provenance)
        self.assertTrue(all(row.get("source") for row in result.debug.numeric_provenance))

    def test_synthetic_narrative_scenarios_do_not_crash_or_revive_removed_concepts(self):
        scenarios=[
            dict(completion=95,coverage=100,coverage_latest=100,attention=0,dynamics=(94,95)),
            dict(completion=82,coverage=95,coverage_latest=98,attention=1,dynamics=(50,70,82)),
            dict(completion=45,coverage=90,coverage_latest=92,attention=4,dynamics=(75,60,45)),
            dict(completion=68,coverage=45,coverage_latest=30,attention=2,dynamics=(65,68)),
            dict(completion=50,coverage=85,coverage_latest=88,attention=5,attention_type="final_nonachievement",dynamics=(55,50)),
            dict(completion=90,coverage=100,coverage_latest=100,attention=0,dynamics=(90,)),
        ]
        for kwargs in scenarios:
            with self.subTest(kwargs=kwargs):
                result=generate_analytics_note(context=self._context(**kwargs),debug=True)
                low=result.text.lower()
                self.assertNotIn("рік до року",low)
                self.assertNotIn("середнє виконання за квартал",low)
                self.assertNotIn("проблемних позицій",low)
                self.assertFalse([w for w in result.debug.validation_warnings if w.startswith(("polarity/sign contradiction","finding contradiction","unsupported numeric tokens","obsolete Analytics narrative concept"))])

    def test_no_forced_paragraph_quota_when_information_is_redundant(self):
        result=generate_analytics_note(context=self._context(completion=90,coverage=100,coverage_latest=100,attention=0,dynamics=(90,90,90)),debug=True)
        paragraphs=[p for p in result.text.split("\n\n") if p.strip()]
        self.assertTrue(paragraphs)
        self.assertIsNone(result.debug.target_paragraph_count)
        self.assertEqual(len(paragraphs),len(set(paragraphs)))
        self.assertEqual(result.debug.important_findings_skipped,[])

    def test_exact_latest_execution_is_never_described_as_average(self):
        result=generate_analytics_note(context=self._context(completion=75.0,coverage=90.0,coverage_latest=90.0,attention=0,dynamics=(75.0,)),debug=True)
        low=result.text.lower()
        self.assertNotIn("середній рівень виконання",low)
        self.assertNotIn("середнє виконання",low)
        self.assertIn("останньому обраному періоді",low)

    def test_final_visible_text_has_no_marker_punctuation_or_locative_errors(self):
        result=generate_analytics_note(context=self._context(),debug=True)
        low=result.text.lower()
        self.assertNotIn("в.п..",result.text)
        self.assertNotIn("у поточному зріз відсутн",low)
        self.assertNotIn("вище за середній.",low)
        self.assertNotIn("подання за 1 захід",low)
        for term in ("execution","coverage","snapshot","drill-down","quarter-aware"):
            self.assertNotIn(term,low)
        self.assertNotIn("найбільший внесок у відхилення формують",low)

    def test_overachievement_context_keeps_prepared_value_and_does_not_crash(self):
        result=generate_analytics_note(context=self._context(completion=120.0,coverage=100.0,coverage_latest=100.0,attention=0,dynamics=(100.0,120.0)),debug=True)
        self.assertIn("120",result.text)
        self.assertFalse([w for w in result.debug.validation_warnings if w.startswith(("polarity/sign contradiction","finding contradiction","unsupported numeric tokens","obsolete Analytics narrative concept"))])



class ContradictionValidationTests(unittest.TestCase):
    def test_positive_number_cannot_be_worded_as_decline(self):
        ctx=NarrativeTests()._context()
        warnings=validate_text("Найбільше зниження становить +19,4 в.п.",ctx,[],[])
        self.assertTrue(any(w.startswith("polarity/sign contradiction") for w in warnings))

    def test_broad_negative_requires_declines_without_improvements(self):
        ctx=NarrativeTests()._context()
        finding=AnalyticalFinding(code="goal_change_broad_negative",topic="goal",importance=90,facts={"declined":0,"improved":2})
        warnings=validate_text("Нейтральний текст без чисел.",ctx,[],[finding])
        self.assertTrue(any(w.startswith("finding contradiction") for w in warnings))

    def test_ratio_rejects_mixed_observation_units(self):
        builder=_Builder({})
        with self.assertRaises(ValueError):
            builder.ratio_pct("bad",1,2,source="test",aggregation="invalid",numerator_unit="measure-period",denominator_unit="unique-measure")



class ExportContractTests(unittest.TestCase):
    @staticmethod
    def _load_export_functions():
        page=Path(__file__).resolve().parents[1] / "pages" / "7_Аналітика.py"
        tree=ast.parse(page.read_text(encoding="utf-8"))
        wanted={"format_pct","_excel_safe_value","_excel_safe_frame","create_excel_report","create_docx_report"}
        module=ast.Module(body=[node for node in tree.body if isinstance(node,(ast.FunctionDef,ast.AsyncFunctionDef)) and node.name in wanted],type_ignores=[])
        from datetime import datetime
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        class Recorder:
            sheets=None
            kwargs=None
            @classmethod
            def write_styled_excel(cls,sheets,**kwargs):
                cls.sheets=sheets; cls.kwargs=kwargs
                return b"xlsx"
        ns={
            "pd":pd,"BytesIO":BytesIO,"Document":Document,"Pt":Pt,"Inches":Inches,
            "WD_ALIGN_PARAGRAPH":WD_ALIGN_PARAGRAPH,"core_exports":Recorder,
            "kyiv_now":lambda:datetime(2026,8,26,9,0),
        }
        exec(compile(module,str(page),"exec"),ns)
        ns["build_report_charts"]=lambda *args,**kwargs: []
        return ns,Recorder

    def test_excel_export_uses_new_analytics_contract_and_preserves_metadata(self):
        ns,rec=self._load_export_functions()
        metrics={"completion":77.0,"coverage":88.0,"coverage_latest":92.0,"attention_label":"Високий або критичний ризик","attention_count":3,"unique_measures":10,"latest_period":(2026,"II")}
        filters={"years":[2026],"quarters":["I","II"],"ssp":[],"deputies":[],"goal_labels":[],"task_labels":[],"product_types":[]}
        empty=pd.DataFrame()
        out=ns["create_excel_report"](empty,empty,empty,empty,empty,empty,empty,empty,metrics,filters)
        self.assertIsInstance(out,BytesIO)
        self.assertIn("Пояснення",rec.sheets)
        self.assertIn("Підсумок",rec.sheets)
        self.assertIn("Аналітичний масив",rec.sheets)
        self.assertIn("Реєстр заявок",rec.sheets)
        self.assertTrue(rec.kwargs.get("freeze_first_col") == 1)
        self.assertFalse(any("yoy" in name.lower() or "рік до року" in name.lower() for name in rec.sheets))
        summary=rec.sheets["Підсумок"]
        labels=set(summary["Показник"].astype(str))
        self.assertIn("Рівень виконання — останній обраний період",labels)
        self.assertIn("Покриття — середнє за вибраний діапазон",labels)
        self.assertIn("Покриття — останній обраний період",labels)
        self.assertFalse(any("проблем" in x.lower() for x in labels))

    def test_docx_export_uses_latest_execution_dual_coverage_and_keeps_graph_section(self):
        ns,_=self._load_export_functions()
        metrics={"completion":77.0,"coverage":88.0,"coverage_latest":92.0,"attention_label":"Фінальний результат не досягнуто","attention_count":2,"unique_measures":10,"total_rows":20,"no_data":1}
        out=ns["create_docx_report"]("Короткий управлінський висновок.",metrics,{"years":[2026],"quarters":["IV"]})
        self.assertIsInstance(out,BytesIO)
        from docx import Document
        doc=Document(out)
        chunks=[p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                chunks.extend(cell.text for cell in row.cells)
        text="\n".join(chunks).lower()
        self.assertIn("останній обраний період",text)
        self.assertIn("покриття за діапазон",text)
        self.assertIn("останній період",text)
        self.assertIn("графічні матеріали не додано",text)
        self.assertNotIn("рік до року",text)
        self.assertNotIn("проблемних позицій",text)

    def test_page_source_keeps_graphical_export_builder(self):
        source=(Path(__file__).resolve().parents[1]/"pages"/"7_Аналітика.py").read_text(encoding="utf-8")
        self.assertIn("def build_report_charts",source)
        self.assertIn("Графічні матеріали",source)
        self.assertIn("fig_png_bytes",source)


class RebaseRegressionTests(unittest.TestCase):
    def test_mio_compatibility_logic_is_not_oversimplified(self):
        source=(Path(__file__).resolve().parents[1]/"pages"/"7_Аналітика.py").read_text(encoding="utf-8")
        self.assertIn("_integral_compatible",source)
        self.assertIn("_mio_indicator_compatible",source)
        self.assertIn("selected_goals",source)
        self.assertIn("selected_tasks",source)
        self.assertIn("_active_codes_for_mio",source)
        # A goal filter alone is not part of the incompatibility tuple for integral MIO.
        self.assertIn("(selected_ssp, selected_deputies, selected_tasks, selected_products)",source)

    def test_calculations_page_retains_light_service_structure(self):
        source=(Path(__file__).resolve().parents[1]/"pages"/"9_Розрахунки.py").read_text(encoding="utf-8")
        for token in (
            "def _safe_display_frame", "def _quarter_result_table", "def _reason_table",
            "Dashboard — розрахунки", "Аналітика — розрахунки", "Графіки Аналітики",
            "Показники аналітичної довідки", "Технічна звірка",
        ):
            self.assertIn(token,source)
        self.assertGreaterEqual(len(source.splitlines()),650)

    def test_runtime_overlay_preserves_base_composer_instead_of_replacing_it(self):
        overlay=(Path(__file__).resolve().parents[1]/"core"/"analytics_text"/"composer_overlay.py").read_text(encoding="utf-8")
        init=(Path(__file__).resolve().parents[1]/"core"/"analytics_text"/"__init__.py").read_text(encoding="utf-8")
        self.assertIn("from . import composer as base",overlay)
        self.assertIn("base._render_block",overlay)
        self.assertIn("from .composer_overlay import compose_note",init)
        self.assertNotIn("найбільший внесок у відхилення формують",overlay)

    def test_managerial_language_cleanup_removes_internal_english(self):
        from core.analytics_text.composer_overlay import _language_cleanup
        raw="execution змінився; coverage не оцінюється; у точному останньому snapshot; Drill-down; execution alone; quarter-aware показник."
        cleaned=_language_cleanup(raw).lower()
        for word in ("execution","coverage","snapshot","drill-down","quarter-aware"):
            self.assertNotIn(word,cleaned)
        self.assertIn("рівень виконання",cleaned)
        self.assertIn("покриття моніторингом",cleaned)


    def test_management_priorities_pass_production_finding_provenance_gate(self):
        ctx=NarrativeTests()._context()
        signals=detect_signals(ctx)
        _, findings=derive_findings(ctx,signals)
        priority=next((f for f in findings if f.code=="management_priorities"),None)
        self.assertIsNotNone(priority)
        warnings=validate_finding_numeric_provenance(ctx,[priority])
        self.assertEqual(warnings,[])
        # Regression guard: all public priority numerics that production validator
        # treats as analytical values must remain bound factual metric values.
        numeric_keys={"execution","change","attention","missing","portfolio_weight","underperformance_contribution"}
        for item in priority.facts.get("priorities",[]):
            for key in numeric_keys:
                value=item.get(key)
                if value is not None:
                    self.assertIsNotNone(metric_code_of(value),f"{key} lost factual provenance")
                    self.assertIsNotNone(ctx.analytical_facts.metric(metric_code_of(value)))

    def test_analytics_coverage_average_does_not_invoke_shared_execution_aggregate(self):
        source=(Path(__file__).resolve().parents[1]/"core"/"analytics_calculations.py").read_text(encoding="utf-8")
        tree=ast.parse(source)
        imported=[]
        calls=[]
        for node in ast.walk(tree):
            if isinstance(node,ast.ImportFrom):
                imported.extend(alias.name for alias in node.names)
            elif isinstance(node,ast.Call) and isinstance(node.func,ast.Name):
                calls.append(node.func.id)
        self.assertNotIn("aggregate_plan",imported)
        self.assertNotIn("aggregate_plan",calls)
        self.assertIn('_selection_average(period_results, "coverage")',source)

    def test_management_rendering_has_no_retired_problem_position_wording(self):
        result=generate_analytics_note(context=NarrativeTests()._context(),debug=True)
        low=result.text.lower()
        for fragment in ("проблемних позицій","проблемними позиціями","проблемні позиції","проблемна позиція"):
            self.assertNotIn(fragment,low)
        self.assertIn("актуаль",low)
        self.assertTrue("управлін" in low or "уваг" in low)

    def test_current_context_never_constructs_retired_facts_and_registers_current_attention(self):
        ctx=NarrativeTests()._context()
        metrics=getattr(ctx.analytical_facts,"metrics",{})
        structures=getattr(ctx.analytical_facts,"structures",{})
        self.assertFalse(any("yoy" in str(code).lower() or "problem" in str(code).lower() for code in metrics))
        self.assertFalse(any("yoy" in str(code).lower() or "problem" in str(code).lower() for code in structures))
        self.assertIn("page.attention_count",metrics)
        self.assertEqual(metrics["page.attention_count"].observation_unit,"unique-measure")
        self.assertIn("goal.attention",structures)

    def test_page_shell_and_exports_are_preserved_from_production_shape(self):
        source=(Path(__file__).resolve().parents[1]/"pages"/"7_Аналітика.py").read_text(encoding="utf-8")
        for token in ("backdrop-filter: blur(8px)",".badge-wrap",".report-meta","def build_report_charts","Графічні матеріали","Таблиці для перевірки","Аналіз повернень на доопрацювання","Швидкість погодження"):
            self.assertIn(token,source)
        self.assertNotIn("render_year_over_year_block",source)
        self.assertNotIn("build_year_over_year_comparison",source)

    def test_retired_context_api_is_physically_absent(self):
        import inspect
        from core.analytics_text.models import AnalyticsContext, NoteQualityMetrics
        from core.analytics_text.context import build_context
        self.assertNotIn("yoy_comparison", AnalyticsContext.__dataclass_fields__)
        self.assertNotIn("yoy_comparison", inspect.signature(build_context).parameters)
        self.assertEqual(
            set(NoteQualityMetrics.__dataclass_fields__),
            {"word_count","paragraph_count","sentence_count","important_finding_coverage","repeated_phrase_count","unique_fact_count","median_sentences_per_paragraph"},
        )

    def test_planner_has_no_hard_paragraph_quota(self):
        from core.analytics_text.planner import build_text_plan
        ctx=NarrativeTests()._context()
        sig=detect_signals(ctx); _,findings=derive_findings(ctx,sig)
        from core.analytics_text.scenarios_current import activate_scenarios
        plan=build_text_plan(ctx,sig,activate_scenarios(sig,findings),findings)
        self.assertIsNone(plan.target_paragraphs)
        source=(Path(__file__).resolve().parents[1]/"core"/"analytics_text"/"planner.py").read_text(encoding="utf-8")
        self.assertNotIn("target_max",source)
        self.assertNotIn("def _targets",source)
    def test_public_text_engine_apis_do_not_expose_retired_families(self):
        ctx=NarrativeTests()._context()
        signals=detect_signals(ctx)
        _, findings=derive_findings(ctx,signals)
        retired=("pro"+"blem","yo"+"y")
        self.assertFalse(any(any(token in s.code.lower() for token in retired) for s in signals))
        self.assertFalse(any(any(token in f.code.lower() or token in f.topic.lower() for token in retired) for f in findings))
        self.assertFalse(any(any(token in str(code).lower() for token in retired) for code in SUPPORTED_SIGNAL_CODES))
        self.assertFalse(any(any(token in str(code).lower() for token in retired) for code in SUPPORTED_FINDING_CODES))
        self.assertFalse(any(any(token in sc.code.lower() for token in retired) for sc in SCENARIOS))
        self.assertFalse(any(any(token in q.code.lower() for token in retired) for q in QUESTIONS))
        self.assertNotIn("what_changed_yoy", {q.code for q in QUESTIONS})
        self.assertNotIn("where_are_problems", {q.code for q in QUESTIONS})

    def test_production_validation_diagnostics_are_preserved(self):
        source=(Path(__file__).resolve().parents[1]/"pages"/"7_Аналітика.py").read_text(encoding="utf-8")
        self.assertIn("validation_warnings",source)
        self.assertIn("st.code",source)
        self.assertIn('diagnostics={"validation_warnings":_validation_warnings,"filters":filters}',source)

    def test_managerial_ui_contains_no_latest_snapshot_english(self):
        source=(Path(__file__).resolve().parents[1]/"pages"/"7_Аналітика.py").read_text(encoding="utf-8")
        self.assertNotIn("Унікальні заходи у latest snapshot",source)
        self.assertNotIn("Середнє за діапазон · latest",source)
        self.assertIn("Унікальні заходи в актуальному зрізі",source)
        self.assertIn("Середнє за діапазон · останній період",source)

    def test_attention_factual_observation_unit_is_unique_measure(self):
        builder=_Builder({})
        value=builder.add("x",2,unit="count",source="test",aggregation="current",observation_unit="unique-measure")
        self.assertEqual(builder.metrics["x"].observation_unit,"unique-measure")
        self.assertEqual(int(value),2)


class MioRegressionTests(unittest.TestCase):
    """Behavioral gates for preserving the target MіO analytical contract."""

    @staticmethod
    def _context_with_mio():
        base = NarrativeTests()._context()
        mio_goals = pd.DataFrame([
            {
                "Код": "1.", "Ціль": "Ціль 1",
                "Інтеграл 2026": 58.0, "Заходи 2026": 72.0,
                "Завдання 2026": 55.0, "Прогрес 2026": 48.0,
            },
            {
                "Код": "2.", "Ціль": "Ціль 2",
                "Інтеграл 2026": 88.0, "Заходи 2026": 92.0,
                "Завдання 2026": 85.0, "Прогрес 2026": 87.0,
            },
        ])
        mio_tasks = pd.DataFrame([
            {"Рівень": "task", "Код": "1.1.", "Оцінка 2026": 80.0},
            {"Рівень": "task", "Код": "1.2.", "Оцінка 2026": 40.0},
        ])
        mio_measures = pd.DataFrame([
            {"Захід": "1.1.1", "Факт/План, %": 40.0},
            {"Захід": "1.1.2", "Факт/План, %": 100.0},
            {"Захід": "1.2.1", "Факт/План, %": 150.0},
        ])
        mio_financing = pd.DataFrame([
            {
                "Захід": "1.1.1", "Назва заходу": "Захід А",
                "План, млрд грн": 1.0, "Факт, млрд грн": 0.8,
                "% виконання": 80.0, "Стан виконання заходу, %": 50.0,
            },
            {
                "Захід": "1.2.1", "Назва заходу": "Захід Б",
                "План, млрд грн": 2.0, "Факт, млрд грн": 0.8,
                "% виконання": 40.0, "Стан виконання заходу, %": 90.0,
            },
        ])
        return build_context(
            filters=base.filters,
            metrics=base.metrics,
            goal_progress=base.goal_progress,
            task_progress=base.task_progress,
            department_progress=base.department_progress,
            product_progress=base.product_progress,
            status_counts=base.status_counts,
            period_dynamics=base.period_dynamics,
            active=base.active,
            mio_goal_evaluation=mio_goals,
            mio_goal_task_evaluation=mio_tasks,
            mio_measure_evaluation=mio_measures,
            mio_financing=mio_financing,
        )

    def test_financing_contract_restores_paired_averages_and_largest_gaps(self):
        ctx = self._context_with_mio()
        self.assertAlmostEqual(float(ctx.factual_value("mio.fin.avg_financial_execution")), 60.0)
        self.assertAlmostEqual(float(ctx.factual_value("mio.fin.avg_physical_execution")), 70.0)
        self.assertEqual(int(ctx.factual_value("mio.fin.paired_count")), 2)
        financing = ctx.factual_structure("mio.financing")
        self.assertEqual(len(financing.get("largest_gaps", [])), 2)
        self.assertEqual(financing["largest_gaps"][0].get("Захід"), "1.2.1")
        self.assertAlmostEqual(float(financing["largest_gaps"][0]["_gap"]), -50.0)

    def test_financial_card_source_receives_real_factual_value(self):
        import ast
        ctx = self._context_with_mio()
        self.assertAlmostEqual(float(ctx.factual_value("mio.fin.avg_financial_execution")), 60.0)
        page_path = Path(__file__).resolve().parents[1] / "pages" / "7_Аналітика.py"
        source = page_path.read_text(encoding="utf-8")
        tree = ast.parse(source)
        found = False
        for node in ast.walk(tree):
            if not isinstance(node, ast.Call):
                continue
            if not (isinstance(node.func, ast.Name) and node.func.id == "format_pct" and node.args):
                continue
            rendered_source = ast.get_source_segment(source, node.args[0]) or ""
            if "factual_value" in rendered_source and "mio.fin.avg_financial_execution" in rendered_source:
                found = True
                break
        self.assertTrue(found, "Financial MіO card must render the factual avg_financial_execution metric")

    def test_task_profile_best_worst_gap_and_divergence_codes_are_preserved(self):
        ctx = self._context_with_mio()
        tasks = ctx.factual_structure("mio.tasks")
        self.assertEqual(tasks["best_task"], "1.1.")
        self.assertEqual(tasks["worst_task"], "1.2.")
        self.assertAlmostEqual(float(tasks["best_task_progress"]), 80.0)
        self.assertAlmostEqual(float(tasks["worst_task_progress"]), 40.0)
        self.assertAlmostEqual(float(tasks["gap"]), 40.0)
        self.assertTrue(tasks.get("divergences"))
        signals = detect_signals(ctx)
        _, findings = derive_findings(ctx, signals)
        codes = {item.code for item in findings}
        self.assertIn("mio_task_indicator_profile", codes)
        self.assertIn("mio_task_execution_result_divergence", codes)
        self.assertNotIn("mio_task_divergence", codes)

    def test_measure_profile_preserves_average_and_median(self):
        ctx = self._context_with_mio()
        measures = ctx.factual_structure("mio.measures")
        self.assertEqual(int(measures["measures_count"]), 3)
        self.assertEqual(int(measures["evaluated_measures"]), 3)
        self.assertAlmostEqual(float(measures["average_fact_plan"]), (40.0 + 100.0 + 150.0) / 3.0)
        self.assertAlmostEqual(float(measures["median_fact_plan"]), 100.0)

    def test_mio_findings_render_through_production_mio_block(self):
        ctx = self._context_with_mio()
        result = generate_analytics_note(context=ctx, debug=True)
        codes = set(result.debug.analytical_findings)
        for code in (
            "mio_task_indicator_profile",
            "mio_task_execution_result_divergence",
            "mio_measure_profile",
            "mio_financing_profile",
        ):
            self.assertIn(code, codes)
        low = result.text.lower()
        self.assertIn("на рівні завдань", low)
        self.assertIn("розриви між виконанням завдань", low)
        self.assertIn("медіан", low)
        self.assertIn("фінансове виконання становить", low)
        self.assertIn("фінансово-фізичних розривів", low)
        _, all_findings = derive_findings(ctx, detect_signals(ctx))
        mio_findings = [item for item in all_findings if item.code.startswith("mio_")]
        self.assertEqual(validate_finding_numeric_provenance(ctx, mio_findings), [])
        self.assertFalse(result.debug.validation_warnings)

    def test_mio_task_execution_uses_exact_latest_not_temporal_average(self):
        source = (Path(__file__).resolve().parents[1] / "core" / "analytics_text" / "analytical_metrics.py").read_text(encoding="utf-8")
        self.assertIn("analytics.task_progress.exact_latest", source)
        self.assertNotIn('groupby("_code")["_exec"].mean()', source)
        ctx = self._context_with_mio()
        divergences = ctx.factual_structure("mio.tasks").get("divergences", [])
        by_code = {item["code"]: item for item in divergences}
        self.assertAlmostEqual(float(by_code["1.1."]["execution"]), 30.0)
        self.assertAlmostEqual(float(by_code["1.1."]["indicator_progress"]), 80.0)
        self.assertAlmostEqual(float(by_code["1.1."]["gap"]), -50.0)

    def test_mio_restoration_does_not_revive_retired_analytics_families(self):
        ctx = self._context_with_mio()
        metrics = ctx.analytical_facts.metrics
        structures = ctx.analytical_facts.structures
        retired = ("problem", "yoy")
        self.assertFalse(any(any(token in str(code).lower() for token in retired) for code in metrics))
        self.assertFalse(any(any(token in str(code).lower() for token in retired) for code in structures))
        source = (Path(__file__).resolve().parents[1] / "core" / "analytics_text" / "analytical_metrics.py").read_text(encoding="utf-8")
        self.assertNotIn("execution_by_measures_average", source)

    def test_public_signal_registry_includes_overlay_management_codes(self):
        self.assertIn("management_attention_present", SUPPORTED_SIGNAL_CODES)
        self.assertIn("management_attention_none", SUPPORTED_SIGNAL_CODES)

    def test_current_trajectory_codes_match_base_renderer_contract(self):
        source = (Path(__file__).resolve().parents[1] / "core" / "analytics_text" / "findings_current.py").read_text(encoding="utf-8")
        self.assertIn("trajectory_net_growth", source)
        self.assertIn("trajectory_net_decline", source)
        self.assertIn("trajectory_reversal_negative", source)
        self.assertNotIn("trajectory_net_growth_mixed", source)
        self.assertNotIn("trajectory_net_decline_mixed", source)
        self.assertNotIn("trajectory_negative_reversal", source)

    def test_rcv4_goal_drilldown_uses_current_children_contract(self):
        ctx = self._context_with_mio()
        drill = ctx.factual_structure("drilldown.goal", {}) or {}
        self.assertIn("children", drill)
        self.assertTrue(drill["children"])
        child = drill["children"][0]
        self.assertIn("execution", child)
        self.assertIn("attention_count", child)
        self.assertIn("missing_count", child)
        _, findings = derive_findings(ctx, detect_signals(ctx))
        goal_drill = next(item for item in findings if item.code == "goal_drilldown")
        self.assertEqual(goal_drill.topic, "task")

    def test_rcv4_ssp_without_child_evidence_has_no_distributed_inference(self):
        result = generate_analytics_note(context=NarrativeTests()._context(), debug=True)
        low = result.text.lower()
        self.assertIn("ssp_drilldown", result.debug.supporting_findings)
        self.assertNotIn("ssp_drilldown", {code for codes in result.debug.block_findings.values() for code in codes})
        self.assertNotIn("відхилення розподілені між кількома завданнями", low)
        self.assertEqual(result.debug.important_findings_skipped, [])

    def test_rcv4_acceleration_code_matches_production_renderer(self):
        ctx = NarrativeTests()._context(dynamics=(50.0, 55.0, 70.0))
        result = generate_analytics_note(context=ctx, debug=True)
        self.assertIn("trajectory_late_acceleration", result.debug.analytical_findings)
        self.assertIn("trajectory_late_acceleration", result.debug.block_findings.get("dynamics", []))
        self.assertEqual(result.debug.important_findings_skipped, [])

    def test_rcv4_finding_planner_taxonomy_is_aligned(self):
        from core.analytics_text.planner import build_text_plan
        from core.analytics_text.scenarios_current import activate_scenarios
        ctx = NarrativeTests()._context()
        signals = detect_signals(ctx); _, findings = derive_findings(ctx, signals)
        plan = build_text_plan(ctx, signals, activate_scenarios(signals, findings), findings)
        task_plan = plan.block_plan("tasks")
        self.assertIsNotNone(task_plan)
        self.assertIn("goal_drilldown", task_plan.finding_codes)
        result = generate_analytics_note(context=ctx, debug=True)
        self.assertIn("goal_drilldown", result.debug.block_findings.get("tasks", []))

    def test_rcv4_compatibility_registry_covers_active_finding_families(self):
        from core.analytics_text.compatibility import FINDING_COMPATIBILITY, RENDERED, SUPPORTING_ONLY, INTERNAL_ONLY
        self.assertEqual(set(FINDING_COMPATIBILITY), set(SUPPORTED_FINDING_CODES))
        self.assertEqual(len(FINDING_COMPATIBILITY), 75)
        self.assertTrue(all(item.disposition in {RENDERED, SUPPORTING_ONLY, INTERNAL_ONLY} for item in FINDING_COMPATIBILITY.values()))
        result = generate_analytics_note(context=NarrativeTests()._context(), debug=True)
        self.assertEqual(result.debug.important_findings_skipped, [])
        self.assertFalse([w for w in result.debug.validation_warnings if w.startswith("compatibility:")])


if __name__ == "__main__":
    unittest.main(verbosity=2)
