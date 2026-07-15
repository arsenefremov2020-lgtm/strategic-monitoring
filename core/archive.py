"""Архівні знімки ДЕМО 2.0: побудова, стиснення, читання та експорти."""

from __future__ import annotations

import base64
import gzip
import io
import json
from datetime import date, datetime, timezone
from decimal import Decimal
from typing import Any
from zoneinfo import ZoneInfo

import pandas as pd

from core.db import fetch_all
from core.strategic_data import load_strat_matrix
from core.statuses import normalize_to_model_status

KYIV_TZ = ZoneInfo("Europe/Kyiv")
ARCHIVE_SCHEMA_VERSION = "DEMO2-ARCHIVE-1"

STATUS_SCORE = {
    "Виконано": 100.0,
    "Частково виконано": 75.0,
    "Не виконано": 0.0,
    "Не настав час": None,
    "Втратило актуальність": None,
}

QUARTER_MAP = {
    1: "I",
    2: "II",
    3: "III",
    4: "IV",
    "1": "I",
    "2": "II",
    "3": "III",
    "4": "IV",
    "I": "I",
    "II": "II",
    "III": "III",
    "IV": "IV",
}


def _clean(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    return str(value).strip()


def _json_value(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, Decimal):
        return float(value)
    if hasattr(value, "item"):
        try:
            return value.item()
        except (ValueError, TypeError):
            pass
    if isinstance(value, float) and pd.isna(value):
        return None
    if not isinstance(value, (list, dict, tuple, set)):
        try:
            if pd.isna(value):
                return None
        except (TypeError, ValueError):
            pass
    return value


def _records(frame: pd.DataFrame | None) -> list[dict[str, Any]]:
    if frame is None or frame.empty:
        return []
    rows: list[dict[str, Any]] = []
    for raw in frame.to_dict(orient="records"):
        rows.append({str(key): _json_value(value) for key, value in raw.items()})
    return rows


def _safe_fetch(client: Any, table: str, select: str = "*") -> list[dict[str, Any]]:
    try:
        return fetch_all(table, select, order=("id", False), client=client)
    except Exception:
        # Не всі історичні інсталяції мають допоміжні таблиці.
        return []


def _quarter_key(value: Any) -> str:
    raw = _clean(value).upper()
    return QUARTER_MAP.get(value, QUARTER_MAP.get(raw, raw))


def _numeric(value: Any) -> float | None:
    if value in (None, ""):
        return None
    try:
        return float(str(value).replace(" ", "").replace(",", "."))
    except (TypeError, ValueError):
        return None


def _approved_request_index(requests: list[dict[str, Any]]) -> dict[tuple[str, int, str], dict[str, Any]]:
    approved = [row for row in requests if _clean(row.get("approval_status")) == "Погоджено"]
    approved.sort(key=lambda row: _clean(row.get("submitted_at")))
    index: dict[tuple[str, int, str], dict[str, Any]] = {}
    for row in approved:
        if _clean(row.get("object_kind")).lower() == "indicator":
            continue
        code = _clean(row.get("strat_code"))
        try:
            year = int(row.get("year"))
        except (TypeError, ValueError):
            continue
        quarter = _quarter_key(row.get("quarter"))
        if code and quarter:
            index[(code, year, quarter)] = row
    return index


def _is_yes_no(unit: Any, plan: Any) -> bool:
    unit_text = _clean(unit).lower()
    plan_text = _clean(plan).lower()
    return "так/ні" in unit_text or "так / ні" in unit_text or plan_text in {"так", "ні"}


def _annual_ratio(unit: Any, fact: Any, plan: Any) -> float | None:
    if _is_yes_no(unit, plan):
        if not _clean(fact) or not _clean(plan):
            return None
        return 100.0 if _clean(fact).lower() == _clean(plan).lower() == "так" else 0.0
    fact_n = _numeric(fact)
    plan_n = _numeric(plan)
    if fact_n is None or plan_n in (None, 0):
        return None
    return round(fact_n / plan_n * 100.0, 2)




def build_main_snapshot(
    strat_df: pd.DataFrame,
    requests: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Повна таблиця Головної: структура плюс всі накопичені звітні періоди."""
    if strat_df is None or strat_df.empty:
        return []

    frame = strat_df.copy()
    if "code" not in frame.columns:
        return _records(frame)

    periods: set[tuple[int, str]] = set()
    latest: dict[tuple[str, int, str], dict[str, Any]] = {}
    ordered_requests = sorted(
        requests,
        key=lambda row: (_clean(row.get("submitted_at")), int(row.get("id") or 0)),
    )
    for row in ordered_requests:
        code = _clean(row.get("strat_code"))
        try:
            year = int(row.get("year"))
        except (TypeError, ValueError):
            continue
        quarter = _quarter_key(row.get("quarter"))
        if not code or quarter not in {"I", "II", "III", "IV"}:
            continue
        periods.add((year, quarter))
        latest[(code, year, quarter)] = row

    quarter_order = {"I": 1, "II": 2, "III": 3, "IV": 4}
    for year, quarter in sorted(periods, key=lambda item: (item[0], quarter_order[item[1]])):
        status_col = f"{year} {quarter} кв. · Статус виконання"
        fact_col = f"{year} {quarter} кв. · Фактичне значення"
        approval_col = f"{year} {quarter} кв. · Статус погодження"
        progress_col = f"{year} {quarter} кв. · Опис прогресу"
        risk_col = f"{year} {quarter} кв. · Ризики"

        def request_for_code(code_value: Any) -> dict[str, Any]:
            return latest.get((_clean(code_value), year, quarter), {})

        rows = frame["code"].map(request_for_code)
        frame[status_col] = rows.map(lambda row: _clean(row.get("status")))
        frame[fact_col] = rows.map(
            lambda row: _json_value(
                row.get("numeric_value")
                if row.get("numeric_value") not in (None, "")
                else row.get("value_text")
            )
        )
        frame[approval_col] = rows.map(lambda row: _clean(row.get("approval_status")))
        frame[progress_col] = rows.map(lambda row: _clean(row.get("progress_text")))
        frame[risk_col] = rows.map(lambda row: _clean(row.get("risks")))

    return _records(frame)


def _mio_plan_is_x(plan: Any) -> bool:
    return _clean(plan).lower() == "х"


def _mio_ratio(unit: Any, s1: str, s2: str, s3: str, year_fact: Any, plan: Any) -> float | str:
    """Формула «Факт/План, %» з чинного режиму М_заходи."""
    if "Втратило актуальність" in (s1, s2, s3):
        return "в/а"
    if _mio_plan_is_x(plan) or not _clean(year_fact):
        return "х"
    if _is_yes_no(unit, plan):
        fact_t = _clean(year_fact).lower()
        plan_t = _clean(plan).lower()
        return 100.0 if fact_t == plan_t == "так" else 0.0
    fact_n = _numeric(year_fact)
    plan_n = _numeric(plan)
    if fact_n is None or plan_n in (None, 0):
        return "х"
    return round(fact_n / plan_n * 100.0, 6)


def _mio_year_status(
    unit: Any,
    s1: str,
    s2: str,
    s3: str,
    year_fact: Any,
    plan: Any,
    ratio: float | str,
) -> str:
    """Формула річного статусу з чинного режиму М_заходи."""
    if "Втратило актуальність" in (s1, s2, s3):
        return "Втратило актуальність"
    if _mio_plan_is_x(plan) or not _clean(year_fact):
        return "Не настав час"
    if _is_yes_no(unit, plan):
        fact_t = _clean(year_fact).lower()
        plan_t = _clean(plan).lower()
        if fact_t == plan_t:
            return "Виконано"
        if fact_t == "ні":
            return "Не виконано"
        return "Не настав час"
    if not isinstance(ratio, (int, float)):
        return "Не настав час"
    if ratio > 99.99:
        return "Виконано"
    if 74.99 < ratio < 100:
        return "Частково виконано"
    if ratio == 0:
        return "Не настав час"
    return "Не виконано"


def build_mio_snapshot(
    strat_df: pd.DataFrame,
    requests: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Фіксує чинні складові режиму «М_заходи» для всіх ССП і 2026–2028 років."""
    if strat_df is None or strat_df.empty or "object_type" not in strat_df.columns:
        return [], []

    # Погоджені подання заходів; останнє подання періоду перемагає.
    approved = [
        row for row in requests
        if _clean(row.get("approval_status")) == "Погоджено"
        and _clean(row.get("object_kind")).lower() != "indicator"
    ]
    approved.sort(key=lambda row: (_clean(row.get("submitted_at")), int(row.get("id") or 0)))
    index: dict[tuple[str, int, str], dict[str, Any]] = {}
    current_names = {
        _clean(row.get("code")): _clean(row.get("name"))
        for _, row in strat_df[strat_df["object_type"] == "measure"].iterrows()
    }
    for row in approved:
        code = _clean(row.get("strat_code"))
        stored_name = _clean(row.get("object_name"))
        current_name = current_names.get(code, "")
        if stored_name and current_name:
            norm_stored = " ".join(stored_name.casefold().split())
            norm_current = " ".join(current_name.casefold().split())
            if norm_stored != norm_current:
                continue
        try:
            year = int(row.get("year"))
        except (TypeError, ValueError):
            continue
        quarter = _quarter_key(row.get("quarter"))
        if code and quarter in {"I", "II", "III", "IV"}:
            index[(code, year, quarter)] = row

    measures = strat_df[strat_df["object_type"] == "measure"].copy()
    detail: list[dict[str, Any]] = []
    score_map = {"Виконано": 100.0, "Частково виконано": 75.0, "Не виконано": 0.0}

    for _, measure in measures.iterrows():
        code = _clean(measure.get("code"))
        department = _clean(measure.get("resp_main"))
        unit = _clean(measure.get("unit")).replace("\n", " ")
        for year in (2026, 2027, 2028):
            plan = measure.get(f"target_{year}", "")
            facts: dict[str, Any] = {}
            statuses: dict[str, str] = {}
            for quarter in ("I", "II", "III", "IV"):
                row = index.get((code, year, quarter), {})
                fact = row.get("numeric_value")
                if fact in (None, ""):
                    fact = row.get("value_text")
                facts[quarter] = _json_value(fact)
                statuses[quarter] = normalize_to_model_status(row.get("status"))

            ratio = _mio_ratio(
                unit,
                statuses["I"],
                statuses["II"],
                statuses["III"],
                facts["IV"],
                plan,
            )
            year_status = _mio_year_status(
                unit,
                statuses["I"],
                statuses["II"],
                statuses["III"],
                facts["IV"],
                plan,
                ratio,
            )
            detail.append({
                "Рік": year,
                "ССП": department,
                "Стратегічна ціль": (
                    f"{_clean(measure.get('parent_goal_code'))} {_clean(measure.get('parent_goal_name'))}"
                ).strip(),
                "Завдання": (
                    f"{_clean(measure.get('parent_task_code'))} {_clean(measure.get('parent_task_name'))}"
                ).strip(),
                "Код заходу": code,
                "Назва заходу": _clean(measure.get("name")),
                "Індикатор": _clean(measure.get("indicator")),
                "Одиниця виміру": unit,
                "Факт · I кв": facts["I"],
                "Стан · I кв": statuses["I"],
                "Факт · I пів": facts["II"],
                "Стан · I пів": statuses["II"],
                "Факт · 9 міс": facts["III"],
                "Стан · 9 міс": statuses["III"],
                "Факт · РІК": facts["IV"],
                "Стан · РІК": year_status,
                "План (ціль. орієнтир)": _json_value(plan),
                "Факт/План, %": ratio,
                "Річний бал виконання": score_map.get(year_status),
            })

    detail_df = pd.DataFrame(detail)
    if detail_df.empty:
        return detail, []

    summary_rows: list[dict[str, Any]] = []
    for (year, department), group in detail_df.groupby(["Рік", "ССП"], dropna=False):
        scores = pd.to_numeric(group["Річний бал виконання"], errors="coerce").dropna()
        numeric_ratios = pd.to_numeric(group["Факт/План, %"], errors="coerce").dropna()
        status_counts = group["Стан · РІК"].value_counts().to_dict()
        summary_rows.append({
            "Рік": int(year),
            "ССП": _clean(department),
            "Кількість заходів": int(len(group)),
            "Заходів з оцінюваним річним статусом": int(scores.count()),
            "Середній річний бал виконання": round(float(scores.mean()), 2) if not scores.empty else None,
            "Середнє Факт/План, %": round(float(numeric_ratios.mean()), 2) if not numeric_ratios.empty else None,
            "Виконано": int(status_counts.get("Виконано", 0)),
            "Частково виконано": int(status_counts.get("Частково виконано", 0)),
            "Не виконано": int(status_counts.get("Не виконано", 0)),
            "Не настав час": int(status_counts.get("Не настав час", 0)),
            "Втратило актуальність": int(status_counts.get("Втратило актуальність", 0)),
        })
    return detail, summary_rows


def _coverage(requests: list[dict[str, Any]]) -> tuple[str, int | None, int | None]:
    periods: set[tuple[int, int]] = set()
    for row in requests:
        try:
            year = int(row.get("year"))
            quarter_raw = row.get("quarter")
            if isinstance(quarter_raw, int):
                quarter = quarter_raw
            else:
                q = _quarter_key(quarter_raw)
                quarter = {"I": 1, "II": 2, "III": 3, "IV": 4}.get(q)
            if quarter:
                periods.add((year, quarter))
        except (TypeError, ValueError):
            continue
    if not periods:
        return "Накопичені дані за всі доступні періоди", None, None
    ordered = sorted(periods)
    labels = [f"{QUARTER_MAP[q]} кв. {y}" for y, q in ordered]
    return ", ".join(labels), ordered[0][0], ordered[-1][0]


def build_archive_payload(client: Any, actor: dict[str, Any], reason: str) -> tuple[dict[str, Any], str]:
    """Будує повну фотографію системи та повертає metadata + gzip/base64 payload."""
    strat_df = load_strat_matrix().copy()
    requests = _safe_fetch(client, "monitoring_requests")
    versions = _safe_fetch(client, "monitoring_request_versions")
    logs = _safe_fetch(client, "monitoring_logs")
    closeouts = _safe_fetch(client, "closeout_requests")
    closeout_versions = _safe_fetch(client, "closeout_request_versions")
    main_snapshot = build_main_snapshot(strat_df, requests)
    mio_detail, mio_summary = build_mio_snapshot(strat_df, requests)

    coverage_label, year_from, year_to = _coverage(requests)
    generated_at = datetime.now(timezone.utc)
    payload = {
        "schema_version": ARCHIVE_SCHEMA_VERSION,
        "generated_at_utc": generated_at.isoformat(),
        "generated_by": {
            "email": _clean(actor.get("email")),
            "name": _clean(actor.get("full_name") or actor.get("name")),
            "role": _clean(actor.get("role")),
        },
        "reason": reason,
        "coverage": coverage_label,
        "main_table": main_snapshot,
        "monitoring_requests": requests,
        "monitoring_request_versions": versions,
        "mio_components": mio_detail,
        "mio_ssp_summary": mio_summary,
        "monitoring_logs": logs,
        "closeout_requests": closeouts,
        "closeout_request_versions": closeout_versions,
    }
    raw = json.dumps(payload, ensure_ascii=False, separators=(",", ":"), default=_json_value).encode("utf-8")
    compressed = gzip.compress(raw, compresslevel=9)
    encoded = base64.b64encode(compressed).decode("ascii")

    measure_count = 0
    if not strat_df.empty and "object_type" in strat_df.columns:
        measure_count = int((strat_df["object_type"] == "measure").sum())

    metadata = {
        "schema_version": ARCHIVE_SCHEMA_VERSION,
        "generated_at_utc": generated_at.isoformat(),
        "coverage_label": coverage_label,
        "coverage_year_from": year_from,
        "coverage_year_to": year_to,
        "structure_row_count": len(main_snapshot),
        "request_count": len(requests),
        "version_count": len(versions),
        "measure_count": measure_count,
        "mio_record_count": len(mio_detail),
        "log_count": len(logs),
        "closeout_count": len(closeouts),
        "closeout_version_count": len(closeout_versions),
        "payload_size_bytes": len(compressed),
        "anchor_year": generated_at.astimezone(KYIV_TZ).year,
        "anchor_quarter": ((generated_at.astimezone(KYIV_TZ).month - 1) // 3) + 1,
    }
    return metadata, encoded


def create_archive_snapshot(
    client: Any,
    *,
    actor: dict[str, Any],
    reason: str,
    snapshot_type: str,
    replaces_snapshot_id: int | None = None,
    replacement_reason: str = "",
) -> dict[str, Any]:
    reason = _clean(reason)
    if not reason:
        raise ValueError("Причина створення архівного знімка є обов’язковою.")
    if replaces_snapshot_id is not None and not _clean(replacement_reason):
        raise ValueError("Для знімка-заміни обов’язково зазначте причину заміни.")

    metadata, encoded = build_archive_payload(client, actor, reason)
    meta = {
        **metadata,
        "reason": reason,
        "snapshot_type": snapshot_type,
        "replaces_snapshot_id": replaces_snapshot_id,
        "replacement_reason": _clean(replacement_reason),
    }
    response = client.rpc(
        "transition_create_archive_snapshot",
        {
            "p_snapshot_meta": meta,
            "p_snapshot_gzip_b64": encoded,
            "p_actor": {
                "email": _clean(actor.get("email")),
                "name": _clean(actor.get("full_name") or actor.get("name")),
                "role": _clean(actor.get("role")),
            },
        },
    ).execute()
    result = response.data
    if isinstance(result, list) and len(result) == 1 and isinstance(result[0], dict):
        result = result[0]
    return result if isinstance(result, dict) else {"success": False, "message": "Некоректна відповідь сервера."}


def decode_snapshot_payload(value: str | None) -> dict[str, Any]:
    if not value:
        return {}
    raw = gzip.decompress(base64.b64decode(value.encode("ascii")))
    return json.loads(raw.decode("utf-8"))


def format_kyiv(value: Any) -> str:
    if not value:
        return "—"
    ts = pd.to_datetime(value, errors="coerce", utc=True)
    if pd.isna(ts):
        return _clean(value) or "—"
    return ts.to_pydatetime().astimezone(KYIV_TZ).strftime("%d.%m.%Y %H:%M")


def snapshot_marker(snapshot: dict[str, Any]) -> str:
    return f"Сформовано з архівного знімка від {format_kyiv(snapshot.get('archived_at'))}"


def _frame(payload: dict[str, Any], key: str) -> pd.DataFrame:
    value = payload.get(key, [])
    return pd.DataFrame(value if isinstance(value, list) else [])


def export_snapshot_excel(snapshot: dict[str, Any], payload: dict[str, Any]) -> bytes:
    output = io.BytesIO()
    marker = snapshot_marker(snapshot)
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        title = pd.DataFrame([
            ["Архівний знімок", format_kyiv(snapshot.get("archived_at"))],
            ["Позначка", marker],
            ["Хто створив", _clean(snapshot.get("archived_by"))],
            ["Тип", "Автоматичний" if snapshot.get("snapshot_type") == "automatic" else "Ручний"],
            ["Охоплені періоди", _clean(snapshot.get("coverage_label"))],
            ["Причина", _clean(snapshot.get("reason"))],
            ["Причина заміни", _clean(snapshot.get("replacement_reason"))],
        ], columns=["Поле", "Значення"])
        title.to_excel(writer, sheet_name="Титулка", index=False)
        sheets = [
            ("Головна", "main_table"),
            ("Заявки", "monitoring_requests"),
            ("Версії", "monitoring_request_versions"),
            ("МіО_складові", "mio_components"),
            ("МіО_ССП", "mio_ssp_summary"),
            ("Журнал", "monitoring_logs"),
            ("Ручні_закриття", "closeout_requests"),
        ]
        for sheet_name, key in sheets:
            frame = _frame(payload, key)
            pd.DataFrame([[marker]]).to_excel(
                writer, sheet_name=sheet_name, index=False, header=False, startrow=0
            )
            frame.to_excel(writer, sheet_name=sheet_name, index=False, startrow=2)
            worksheet = writer.sheets[sheet_name]
            worksheet.freeze_panes(2, 0)
            worksheet.autofilter(2, 0, max(2, len(frame) + 2), max(0, len(frame.columns) - 1))
    return output.getvalue()


def export_snapshot_docx(snapshot: dict[str, Any], payload: dict[str, Any]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Архівний знімок системи моніторингу", 0)
    doc.add_paragraph(snapshot_marker(snapshot))
    doc.add_paragraph(f"Створено: {format_kyiv(snapshot.get('archived_at'))}")
    doc.add_paragraph(f"Автор: {_clean(snapshot.get('archived_by'))}")
    doc.add_paragraph(f"Охоплені періоди: {_clean(snapshot.get('coverage_label'))}")
    doc.add_paragraph(f"Причина: {_clean(snapshot.get('reason'))}")
    if _clean(snapshot.get("replacement_reason")):
        doc.add_paragraph(f"Причина заміни: {_clean(snapshot.get('replacement_reason'))}")

    sections = [
        ("Заявки", "monitoring_requests", ["id", "strat_code", "year", "quarter", "department", "status", "approval_status"]),
        ("Оцінки МіО", "mio_ssp_summary", None),
        ("Журнал дій", "monitoring_logs", ["changed_at", "request_id", "action", "old_status", "new_status", "changed_by"]),
    ]
    for title, key, preferred in sections:
        doc.add_heading(title, level=1)
        frame = _frame(payload, key)
        if frame.empty:
            doc.add_paragraph("Немає даних.")
            continue
        columns = [col for col in (preferred or list(frame.columns)) if col in frame.columns]
        if not columns:
            columns = list(frame.columns)
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        for idx, col in enumerate(columns):
            table.rows[0].cells[idx].text = str(col)
        for _, row in frame.iterrows():
            cells = table.add_row().cells
            for idx, col in enumerate(columns):
                cells[idx].text = _clean(row.get(col))
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def export_snapshot_pdf(snapshot: dict[str, Any], payload: dict[str, Any]) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    output = io.BytesIO()
    try:
        pdfmetrics.registerFont(TTFont("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
        font = "DejaVuSans"
    except Exception:
        font = "Helvetica"

    styles = getSampleStyleSheet()
    for style in styles.byName.values():
        style.fontName = font
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=14 * mm,
        leftMargin=14 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )
    story = [
        Paragraph("Архівний знімок системи моніторингу", styles["Title"]),
        Paragraph(snapshot_marker(snapshot), styles["Heading2"]),
        Spacer(1, 6),
        Paragraph(f"Автор: {_clean(snapshot.get('archived_by'))}", styles["BodyText"]),
        Paragraph(f"Охоплені періоди: {_clean(snapshot.get('coverage_label'))}", styles["BodyText"]),
        Paragraph(f"Причина: {_clean(snapshot.get('reason'))}", styles["BodyText"]),
    ]
    if _clean(snapshot.get("replacement_reason")):
        story.append(Paragraph(f"Причина заміни: {_clean(snapshot.get('replacement_reason'))}", styles["BodyText"]))

    sections = [
        ("Заявки", "monitoring_requests", ["id", "strat_code", "year", "quarter", "department", "status", "approval_status"]),
        ("Оцінки МіО по ССП", "mio_ssp_summary", None),
        ("Журнал дій", "monitoring_logs", ["changed_at", "request_id", "action", "old_status", "new_status", "changed_by"]),
    ]
    for title, key, preferred in sections:
        story.extend([PageBreak(), Paragraph(title, styles["Heading1"]), Paragraph(snapshot_marker(snapshot), styles["BodyText"])])
        frame = _frame(payload, key)
        if frame.empty:
            story.append(Paragraph("Немає даних.", styles["BodyText"]))
            continue
        columns = [col for col in (preferred or list(frame.columns)) if col in frame.columns]
        if not columns:
            columns = list(frame.columns)
        for _, row in frame.iterrows():
            text = " · ".join(f"{col}: {_clean(row.get(col))}" for col in columns)
            story.append(Paragraph(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["BodyText"]))
            story.append(Spacer(1, 3))
    doc.build(story)
    return output.getvalue()
