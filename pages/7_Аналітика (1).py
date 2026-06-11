import re
from datetime import datetime
from io import BytesIO
from html import escape

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from supabase import create_client

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# Page config
# ============================================================

st.set_page_config(page_title="Аналітика", layout="wide")

try:
    st.logo("assets/Мінекономіки.png", size="large")
except Exception:
    pass

FILE_PATH = "Під моніторинг СП.xlsx"
SHEET_NAME = "Страт_матриця"

supabase = create_client(
    st.secrets["SUPABASE_URL"],
    st.secrets["SUPABASE_KEY"]
)

QUARTERS = ["I", "II", "III", "IV"]
YEAR_OPTIONS = [2026, 2027, 2028]

DEPUTY_MINISTER_BY_SSP = {
    "20": "ПИВОВАРОВ Андрій Андрійович",
    "21": "ПИВОВАРОВ Андрій Андрійович",
    "22": "ПИВОВАРОВ Андрій Андрійович",
    "23": "ЦИБОРТ Олександр Сергійович",
    "24": "ПИВОВАРОВ Андрій Андрійович",
    "25": "СОБОЛЕВ Олексій Дмитрович",
    "26": "БЕЗКАРАВАЙНИЙ Ігор Володимирович",
    "27": "КІНДРАТІВ Віталій Зіновійович",
    "28": "АРТЕМЕНКО Анна Ігорівна",
    "29": "КІНДРАТІВ Віталій Зіновійович",
    "30": "МАРЧАК Дарія Миколаївна",
    "31": "ПЕРЕЛИГІН Єгор Євгенович",
    "32": "МАРЧАК Дарія Миколаївна",
    "33": "АРТЕМЕНКО Анна Ігорівна",
    "34": "ПЕТРУК Віталій Вікторович",
    "35": "АРТЕМЕНКО Анна Ігорівна",
    "36": "ЦИБОРТ Олександр Сергійович",
    "37": "МАРЧАК Дарія Миколаївна",
    "38": "КІНДРАТІВ Віталій Зіновійович",
    "39": "ПЕРЕЛИГІН Єгор Євгенович",
    "40": "ЦИБОРТ Олександр Сергійович",
    "41": "ПЕТРУК Віталій Вікторович",
    "42": "АРТЕМЕНКО Анна Ігорівна",
    "43": "МАРЧАК Дарія Миколаївна",
    "44": "КІНДРАТІВ Віталій Зіновійович",
    "45": "ПИВОВАРОВ Андрій Андрійович",
    "46": "КІНДРАТІВ Віталій Зіновійович",
    "47": "МАРЧАК Дарія Миколаївна",
    "48": "МАРЧАК Дарія Миколаївна",
    "49": "АРТЕМЕНКО Анна Ігорівна",
    "50": "СОБОЛЕВ Олексій Дмитрович",
    "51": "ПИВОВАРОВ Андрій Андрійович",
    "52": "БЕЗКАРАВАЙНИЙ Ігор Володимирович",
    "54": "ПИВОВАРОВ Андрій Андрійович",
    "55": "ПИВОВАРОВ Андрій Андрійович",
    "56": "КРАСНОЛУЦЬКИЙ Олександр Васильович",
    "57": "ПИВОВАРОВ Андрій Андрійович",
    "58": "ПИВОВАРОВ Андрій Андрійович",
    "59": "СОБОЛЕВ Олексій Дмитрович",
    "60": "КРАСНОЛУЦЬКИЙ Олександр Васильович",
    "61": "КІНДРАТІВ Віталій Зіновійович",
    "62": "ПЕРЕЛИГІН Єгор Євгенович",
    "63": "КРАСНОЛУЦЬКИЙ Олександр Васильович",
    "64": "КРАСНОЛУЦЬКИЙ Олександр Васильович",
    "65": "КРАСНОЛУЦЬКИЙ Олександр Васильович",
    "67": "КРАСНОЛУЦЬКИЙ Олександр Васильович",
    "68": "КРАСНОЛУЦЬКИЙ Олександр Васильович",
    "69": "ВИСОЦЬКИЙ Тарас Миколайович",
    "70": "ВИСОЦЬКИЙ Тарас Миколайович",
    "71": "БАШЛИК Денис Олександрович",
    "72": "ВИСОЦЬКИЙ Тарас Миколайович",
    "73": "БАШЛИК Денис Олександрович",
    "74": "ОВЧАРЕНКО Ірина Іванівна",
    "75": "ПИВОВАРОВ Андрій Андрійович",
    "76": "ПИВОВАРОВ Андрій Андрійович",
    "77": "МАРЧАК Дарія Миколаївна",
    "78": "",
    "79": "",
    "80": "ВИСОЦЬКИЙ Тарас Миколайович",
}


# ============================================================
# CSS — responsive, fluid, single source
# ============================================================

st.markdown("""
<style>
header[data-testid="stHeader"] { background: transparent !important; }

.stApp {
    background:
        radial-gradient(circle at top right, rgba(37,99,235,0.08), transparent 28%),
        radial-gradient(circle at bottom left, rgba(22,163,74,0.07), transparent 30%),
        linear-gradient(180deg, #f6f8fb 0%, #eef2f7 100%);
}
.stApp::before {
    content:""; position:fixed; top:-160px; right:-120px;
    width:460px; height:460px; border-radius:50%;
    background:rgba(37,99,235,0.04); z-index:0;
}
.stApp::after {
    content:""; position:fixed; bottom:-180px; left:-120px;
    width:390px; height:390px; border-radius:50%;
    background:rgba(22,163,74,0.04); z-index:0;
}
.main .block-container {
    max-width: 1560px;
    padding: clamp(0.6rem,2vw,1.2rem) clamp(0.6rem,2vw,2rem);
    position: relative; z-index:1;
}

/* ── Typography scale ── */
.ua-line {
    height:7px; border-radius:999px;
    background:linear-gradient(90deg,#005BBB 50%,#FFD500 50%);
    margin-bottom:14px;
}
.ministry-label {
    text-align:right; color:#475569;
    font-size: clamp(11px,1vw,14px); font-weight:700; margin-bottom:8px;
}

/* ── Shared card base ── */
.header-box, .filter-box, .summary-banner,
.report-box, .section-card, .export-box {
    background: rgba(255,255,255,0.96);
    border: 1px solid #d8dee9;
    border-radius: 16px;
    box-shadow: 0 6px 20px rgba(15,23,42,0.055);
}

/* ── Header ── */
.header-box {
    padding: clamp(16px,2.2vw,26px) clamp(16px,2.5vw,28px);
    margin-bottom: 16px;
    backdrop-filter: blur(8px);
}
.header-title {
    font-size: clamp(20px,2.6vw,32px);
    font-weight:950; color:#0f172a; margin-bottom:6px;
}
.header-subtitle {
    font-size: clamp(12px,1vw,15px);
    color:#475569; line-height:1.55;
}
.badge-wrap { display:flex; flex-wrap:wrap; gap:8px; margin-top:12px; }
.badge {
    background:#eef6ff; border:1px solid #bfdbfe; color:#1d4ed8;
    border-radius:999px; padding:5px 11px;
    font-size: clamp(11px,0.85vw,13px); font-weight:700;
}
.badge-green { background:#dcfce7; border-color:#bbf7d0; color:#166534; }
.badge-yellow { background:#fef9c3; border-color:#fde68a; color:#854d0e; }

/* ── Summary banner (live metrics at top, before filters) ── */
.summary-banner {
    padding: clamp(14px,2vw,22px) clamp(14px,2.2vw,26px);
    margin-bottom:16px;
    border-left: 6px solid #2563eb;
}
.summary-banner-title {
    font-size: clamp(13px,1.1vw,15px);
    font-weight:900; color:#1e293b; margin-bottom:12px;
}
.summary-grid {
    display:grid;
    grid-template-columns: repeat(auto-fit, minmax(140px,1fr));
    gap: clamp(8px,1.2vw,14px);
}
.summary-cell {
    background:#f8fafc; border:1px solid #e2e8f0;
    border-radius:12px; padding:12px 14px;
    display:flex; flex-direction:column; gap:2px;
}
.summary-cell-label {
    font-size: clamp(10px,0.8vw,12px);
    font-weight:700; color:#64748b; line-height:1.3;
}
.summary-cell-value {
    font-size: clamp(20px,2.2vw,28px);
    font-weight:950; color:#0f172a; line-height:1;
}
.summary-cell-note {
    font-size: clamp(10px,0.75vw,11px);
    color:#94a3b8; margin-top:2px;
}
.summary-cell.blue  { border-top:3px solid #2563eb; }
.summary-cell.green { border-top:3px solid #16a34a; }
.summary-cell.yellow{ border-top:3px solid #d97706; }
.summary-cell.red   { border-top:3px solid #dc2626; }
.summary-cell.gray  { border-top:3px solid #64748b; }

/* ── Filter box ── */
.filter-box {
    padding: clamp(16px,2vw,24px) clamp(16px,2.2vw,26px);
    margin-bottom:16px;
    background: linear-gradient(180deg, rgba(255,255,255,0.99), rgba(241,246,253,0.99));
    border-color:#cbd8ea;
}
.filter-box-title {
    font-size: clamp(16px,1.4vw,21px);
    font-weight:950; color:#0f172a; margin-bottom:6px;
}
.filter-box-hint {
    font-size: clamp(11px,0.85vw,13px);
    color:#64748b; margin-bottom:14px; line-height:1.5;
}
.filter-section-title {
    color:#1e293b; font-size: clamp(12px,0.95vw,14px);
    font-weight:900; margin:12px 0 8px 0;
    padding-bottom:4px; border-bottom:1px solid rgba(148,163,184,0.35);
}

/* widget overrides */
div[data-testid="stSelectbox"] div[data-baseweb="select"] > div,
div[data-testid="stMultiSelect"] div[data-baseweb="select"] > div {
    background-color:#d7eaff !important;
    border:1px solid #8fb3df !important;
    border-radius:10px !important;
    min-height:40px !important;
}
div[data-testid="stSelectbox"] label,
div[data-testid="stMultiSelect"] label {
    font-weight:800 !important; color:#1e293b !important;
    font-size: clamp(11px,0.85vw,13px) !important;
}

/* ── Section card ── */
.section-card {
    padding: clamp(16px,2vw,24px) clamp(16px,2.2vw,26px);
    margin-bottom:16px;
}
.section-card-title {
    font-size: clamp(16px,1.4vw,21px);
    font-weight:950; color:#0f172a; margin-bottom:6px;
}
.section-card-sub {
    font-size: clamp(11px,0.85vw,13px);
    color:#64748b; margin-bottom:14px;
}
/* primary section = blue left border */
.section-card.primary { border-left:6px solid #2563eb; }
/* warning section = yellow left border */
.section-card.warning { border-left:6px solid #d97706; }

/* ── Analytical report text ── */
.report-meta {
    background:#f8fafc; border:1px solid #e2e8f0;
    border-radius:10px; padding:10px 14px;
    color:#475569; font-size: clamp(11px,0.85vw,13px);
    font-weight:700; margin-bottom:14px;
    display:flex; flex-wrap:wrap; gap:8px;
}
.report-meta-chip {
    background:rgba(255,255,255,0.8); border:1px solid #cbd5e1;
    border-radius:20px; padding:3px 10px;
    font-size: clamp(10px,0.8vw,12px); color:#334155;
}
.report-text {
    font-size: clamp(13px,0.95vw,15px);
    line-height:1.8; color:#334155;
    text-align:justify; margin-bottom:12px;
}
.report-text:first-of-type { font-weight:700; color:#0f172a; }

/* ── Chart wrapper ── */
.chart-wrap {
    background:#ffffff; border:1px solid #e9eef5;
    border-radius:12px; padding: clamp(10px,1.2vw,16px);
    margin-bottom:12px;
}
.chart-title {
    font-size: clamp(12px,0.95vw,14px);
    font-weight:800; color:#0f172a; margin-bottom:6px;
}
.chart-sub {
    font-size: clamp(10px,0.8vw,12px);
    color:#64748b; margin-bottom:8px;
}

/* ── Comparison row (two-period) ── */
.comparison-grid {
    display:grid;
    grid-template-columns: repeat(auto-fit, minmax(220px,1fr));
    gap: clamp(8px,1.2vw,14px);
    margin: 12px 0 16px;
}
.comparison-cell {
    background:#f8fafc; border:1px solid #e2e8f0;
    border-radius:12px; padding:14px 16px;
}
.comparison-period {
    font-size: clamp(10px,0.8vw,12px);
    font-weight:700; color:#64748b; margin-bottom:6px;
}
.comparison-value {
    font-size: clamp(22px,2vw,30px);
    font-weight:950; color:#0f172a; line-height:1;
}
.comparison-delta {
    font-size: clamp(11px,0.85vw,13px);
    font-weight:700; margin-top:4px;
}
.delta-up   { color:#16a34a; }
.delta-down { color:#dc2626; }
.delta-flat { color:#64748b; }

/* ── Export box ── */
.export-box {
    padding: clamp(16px,2vw,24px) clamp(16px,2.2vw,26px);
    margin-bottom:16px;
}
.export-title {
    font-size: clamp(16px,1.4vw,21px);
    font-weight:950; color:#0f172a; margin-bottom:4px;
}
.export-sub {
    font-size: clamp(11px,0.85vw,13px);
    color:#64748b; margin-bottom:16px;
}

div.stDownloadButton > button,
div.stButton > button {
    border-radius:10px; padding:10px 16px; font-weight:800;
    font-size: clamp(12px,0.9vw,14px) !important;
}

/* ── Tables section ── */
.tables-section {
    background:rgba(255,255,255,0.96); border:1px solid #d8dee9;
    border-radius:16px; box-shadow:0 6px 20px rgba(15,23,42,0.055);
    padding: clamp(14px,1.8vw,22px) clamp(14px,2vw,24px);
    margin-bottom:16px;
}
.tables-title {
    font-size: clamp(15px,1.2vw,18px);
    font-weight:950; color:#0f172a; margin-bottom:4px;
}
.tables-sub {
    font-size: clamp(11px,0.8vw,13px);
    color:#64748b; margin-bottom:12px;
}

/* ── Footer ── */
.footer {
    text-align:center; color:#64748b;
    font-size: clamp(10px,0.8vw,13px);
    margin-top:40px; padding:18px 0 10px;
    border-top:1px solid #d8dee9;
}

/* ── Responsive breakpoints ── */
@media (max-width:900px) {
    .summary-grid { grid-template-columns: repeat(2, 1fr); }
    .comparison-grid { grid-template-columns: 1fr 1fr; }
}
@media (max-width:600px) {
    .summary-grid { grid-template-columns: 1fr 1fr; }
    .comparison-grid { grid-template-columns: 1fr; }
    .badge-wrap { gap:5px; }
}
</style>
""", unsafe_allow_html=True)


# ============================================================
# Helpers
# ============================================================

def raw_value(value):
    if value is None or pd.isna(value) or str(value) == "None":
        return ""
    return str(value).strip()


def clean(value):
    return escape(raw_value(value))


def normalize_text(value):
    return raw_value(value).lower().replace("і", "i")


def is_empty_or_nd(value):
    return normalize_text(value).replace(" ", "") in {"", "н.д.", "нд", "nan", "none", "-", "—"}


def extract_ssp_index(value):
    text = raw_value(value)
    match = re.search(r"\d+", text)
    return match.group(0) if match else ""


def split_ssp_values(value):
    return re.findall(r"\d+", raw_value(value))


def get_deputy_minister_by_main_ssp(value):
    index = extract_ssp_index(value)
    return DEPUTY_MINISTER_BY_SSP.get(index, "")


def ssp_sort_key(value):
    index = extract_ssp_index(value)
    return (int(index) if index.isdigit() else 10_000, raw_value(value))


def get_goal_code(code):
    parts = raw_value(code).split(".")
    return parts[0] + "." if parts and parts[0] else ""


def get_task_code(code):
    parts = raw_value(code).split(".")
    if len(parts) >= 2:
        return f"{parts[0]}.{parts[1]}."
    return ""


def quarter_to_number(q):
    text = raw_value(q).replace(" квартал", "")
    return {"I": 1, "II": 2, "III": 3, "IV": 4, "1": 1, "2": 2, "3": 3, "4": 4}.get(text, 1)


def parse_period(value):
    text = normalize_text(value).strip()
    if text in ["", "nan", "none", "н.д.", "нд"]:
        return None
    q = None
    year = None
    if "1 квартал" in text or "i квартал" in text:
        q = 1
    elif "2 квартал" in text or "ii квартал" in text:
        q = 2
    elif "3 квартал" in text or "iii квартал" in text:
        q = 3
    elif "4 квартал" in text or "iv квартал" in text:
        q = 4
    year_match = re.search(r"20\d{2}", text)
    if year_match:
        year = int(year_match.group())
    if year and q:
        return year * 10 + q
    return None


def to_number(value):
    text = raw_value(value).replace(" ", "").replace(",", ".")
    if normalize_text(text) in ["", "nan", "none", "н.д.", "нд", "x", "х", "так", "ні", "yes", "no"]:
        return None
    match = re.search(r"-?\d+(\.\d+)?", text)
    if not match:
        return None
    try:
        return float(match.group())
    except Exception:
        return None


def status_score(status):
    status = raw_value(status)
    if status == "Виконано":
        return 100
    if status == "Виконано частково":
        return 50
    if status == "Виконується":
        return 40
    if status == "Потребує уваги":
        return 25
    if status in ["Прострочено", "Не розпочато", "Не подано"]:
        return 0
    return 0


def plan_fact_percent(actual, target):
    actual_num = to_number(actual)
    target_num = to_number(target)
    actual_text = normalize_text(actual)
    target_text = normalize_text(target)
    if actual_num is not None and target_num is not None and target_num != 0:
        return round(min((actual_num / target_num) * 100, 150), 1)
    if target_text in ["так", "yes"] or actual_text in ["так", "ні", "yes", "no"]:
        if actual_text in ["так", "yes"]:
            return 100
        if actual_text in ["ні", "no"]:
            return 0
    return None


def deviation_label(value):
    sign = "+" if value > 0 else ""
    return f"{sign}{round(value, 1)} в.п."


def deviation_color_class(value):
    if value >= 0:
        return "delta-up"
    if value >= -15:
        return "delta-flat"
    return "delta-down"


def traffic_light(score):
    if score >= 70:
        return "У графіку"
    if score >= 35:
        return "Потребує уваги"
    return "Відстає"


def get_indicator_type(row):
    unit = normalize_text(row.get("unit", ""))
    if "так/нi" in unit or ("так" in unit and "нi" in unit):
        return "Так/ні"
    if "%" in unit or "відсот" in unit:
        return "Відсотковий"
    if any(x in unit for x in ["грн", "дол", "євро", "eur", "usd"]):
        return "Фінансовий"
    return "Кількісний"


def concise_list(values, limit=5):
    values = [raw_value(v) for v in values if raw_value(v)]
    if not values:
        return "н/д"
    if len(values) <= limit:
        return "; ".join(values)
    return "; ".join(values[:limit]) + f" та ще {len(values) - limit}"


def filter_label(selected, default_label="усі"):
    if not selected:
        return default_label
    return concise_list(selected, limit=4)


def plotly_theme():
    """Common Plotly layout kwargs for consistent look."""
    return dict(
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="#f8fafc",
        font=dict(family="Inter, Helvetica Neue, Arial, sans-serif", size=12, color="#334155"),
        margin=dict(l=10, r=10, t=40, b=10),
        title_font=dict(size=14, color="#0f172a", family="Inter, sans-serif"),
    )


# ============================================================
# Data loading
# ============================================================

@st.cache_data
def load_strat_matrix():
    source_df = pd.read_excel(FILE_PATH, sheet_name=SHEET_NAME, header=None, engine="openpyxl")
    data = source_df.iloc[7:].copy()

    def safe_col(index):
        if index < source_df.shape[1]:
            return data.iloc[:, index]
        return pd.Series([""] * len(data), index=data.index)

    def find_col_by_keywords(keywords):
        keywords = [k.lower() for k in keywords]
        header_area = source_df.iloc[:7, :].copy()
        for col_idx in range(source_df.shape[1]):
            joined = " ".join(
                raw_value(header_area.iloc[row_idx, col_idx]).lower()
                for row_idx in range(len(header_area))
            )
            if all(keyword in joined for keyword in keywords):
                return col_idx
        return None

    def safe_keyword_col(keywords):
        col_idx = find_col_by_keywords(keywords)
        if col_idx is None:
            return pd.Series([""] * len(data), index=data.index)
        return safe_col(col_idx)

    result = pd.DataFrame({
        "type_marker":          safe_col(1),
        "code":                 safe_col(2),
        "name":                 safe_col(3),
        "product_type":         safe_col(4),
        "indicator":            safe_col(5),
        "unit":                 safe_col(6),
        "base_2021":            safe_col(7),
        "fact_2024":            safe_col(8),
        "fact_2025":            safe_col(9),
        "target_2026":          safe_col(10),
        "target_2027":          safe_col(11),
        "target_2028":          safe_col(12),
        "strategic_target_2028": safe_col(13),
        "strategic_target_2034": safe_col(14),
        "source_global":        safe_col(15),
        "source_national":      safe_col(16),
        "resp_main":            safe_col(17),
        "resp_co_1":            safe_col(18),
        "resp_co_2":            safe_col(19),
        "measure_period_years": safe_keyword_col(["період", "років"]),
        "measure_start_date":   safe_keyword_col(["початкова", "дата"]),
        "measure_end_date":     safe_keyword_col(["кінцева", "дата"]),
        "start_period":         safe_keyword_col(["початок", "квартал"]),
        "end_period":           safe_keyword_col(["кінець", "квартал"]),
        "budget_kpkvk":         safe_col(24),
        "budget_2026_approved": safe_col(25),
        "budget_2027_forecast": safe_col(26),
        "budget_2028_forecast": safe_col(27),
        "other_source":         safe_col(28),
        "other_2026_plan":      safe_col(29),
        "other_2027_forecast":  safe_col(30),
        "other_2028_forecast":  safe_col(31),
        "department":           safe_col(17),
    })

    result = result.dropna(subset=["code"])
    result["code"] = result["code"].astype(str).str.strip()
    result["type_marker"] = result["type_marker"].astype(str).str.strip()

    current_goal_code = ""
    current_goal_name = ""
    current_task_code = ""
    current_task_name = ""
    object_types, parent_goal_codes, parent_goal_names, parent_task_codes, parent_task_names = [], [], [], [], []

    for _, row in result.iterrows():
        marker = normalize_text(row["type_marker"])
        code = raw_value(row["code"])
        dots = code.count(".")
        if "стратегічна ціль" in marker:
            object_type = "goal"
            current_goal_code = code
            current_goal_name = raw_value(row["name"])
            current_task_code = ""
            current_task_name = ""
        elif "завдання" in marker:
            object_type = "task"
            current_task_code = code
            current_task_name = raw_value(row["name"])
        elif "заход" in marker or dots >= 3:
            object_type = "measure"
        else:
            object_type = "task_indicator" if current_task_code else "goal_indicator" if current_goal_code else "other"

        object_types.append(object_type)
        parent_goal_codes.append(current_goal_code)
        parent_goal_names.append(current_goal_name)
        parent_task_codes.append(current_task_code)
        parent_task_names.append(current_task_name)

    result["object_type"] = object_types
    result["parent_goal_code"] = parent_goal_codes
    result["parent_goal_name"] = parent_goal_names
    result["parent_task_code"] = parent_task_codes
    result["parent_task_name"] = parent_task_names

    if result["start_period"].apply(raw_value).eq("").all() and source_df.shape[1] > 22:
        result["start_period"] = safe_col(22)
    if result["end_period"].apply(raw_value).eq("").all() and source_df.shape[1] > 23:
        result["end_period"] = safe_col(23)

    result["ssp_index"] = result["resp_main"].apply(extract_ssp_index)
    result["deputy_minister"] = result["resp_main"].apply(get_deputy_minister_by_main_ssp)
    result["indicator_type"] = result.apply(get_indicator_type, axis=1)
    return result


@st.cache_data(ttl=60)
def load_requests():
    try:
        response = supabase.table("monitoring_requests").select("*").execute()
        return pd.DataFrame(response.data or [])
    except Exception:
        return pd.DataFrame()


def ensure_request_columns(requests_df):
    required = [
        "id", "year", "quarter", "department", "strat_code", "status", "numeric_value",
        "risks", "progress_text", "approval_status", "submitted_at", "responsible_person",
        "phone", "email", "file_names", "file_urls", "admin_comment", "start_date", "end_date"
    ]
    for col in required:
        if col not in requests_df.columns:
            requests_df[col] = ""
    return requests_df


# ============================================================
# Period and filtering logic
# ============================================================

def base_measures(strat_df):
    measures = strat_df[strat_df["object_type"] == "measure"].copy()
    goals = strat_df[strat_df["object_type"] == "goal"].copy()
    tasks = strat_df[strat_df["object_type"] == "task"].copy()

    goal_names = goals.set_index("code")["name"].to_dict()
    task_names = tasks.set_index("code")["name"].to_dict()

    measures["goal_code"] = measures["parent_goal_code"].where(
        measures["parent_goal_code"].astype(str).str.strip() != "",
        measures["code"].apply(get_goal_code)
    )
    measures["task_code"] = measures["parent_task_code"].where(
        measures["parent_task_code"].astype(str).str.strip() != "",
        measures["code"].apply(get_task_code)
    )
    measures["strategic_goal"] = measures["goal_code"].map(goal_names).fillna(measures["parent_goal_name"])
    measures["task_name"] = measures["task_code"].map(task_names).fillna(measures["parent_task_name"])
    measures["start_num"] = measures["start_period"].apply(parse_period)
    measures["end_num"] = measures["end_period"].apply(parse_period)
    measures["ssp_index"] = measures["resp_main"].apply(extract_ssp_index)
    measures["deputy_minister"] = measures["resp_main"].apply(get_deputy_minister_by_main_ssp)
    measures["indicator_type"] = measures.apply(get_indicator_type, axis=1)
    return measures


def is_active_for_period(row, year, quarter):
    selected_period_num = int(year) * 10 + quarter_to_number(quarter)
    target_col = f"target_{year}"
    has_target_for_year = True
    if target_col in row:
        has_target_for_year = not is_empty_or_nd(row.get(target_col, ""))
    start_num = row.get("start_num", None)
    end_num = row.get("end_num", None)
    starts_ok = pd.isna(start_num) or start_num is None or start_num <= selected_period_num
    ends_ok = pd.isna(end_num) or end_num is None or end_num >= selected_period_num
    return bool(has_target_for_year and starts_ok and ends_ok)


def latest_approved_records(requests_df, year, quarter):
    if requests_df.empty:
        return pd.DataFrame(columns=["strat_code"])
    data = requests_df[
        (requests_df["year"].astype(str).str.strip() == str(year)) &
        (requests_df["quarter"].astype(str).str.strip() == str(quarter)) &
        (requests_df["approval_status"].astype(str).str.strip() == "Погоджено")
    ].copy()
    if data.empty:
        return data
    data["submitted_at_sort"] = pd.to_datetime(data["submitted_at"], errors="coerce")
    data = data.sort_values(["strat_code", "submitted_at_sort"], na_position="first")
    return data.groupby("strat_code", as_index=False).tail(1)


def prepare_period_slice(measures, requests_df, year, quarter):
    active = measures[measures.apply(
        lambda row: is_active_for_period(row, year, quarter), axis=1
    )].copy()

    target_col = f"target_{year}"
    active["selected_target"] = active[target_col] if target_col in active.columns else ""
    active["report_year"] = int(year)
    active["report_quarter"] = quarter
    active["report_quarter_num"] = quarter_to_number(quarter)
    active["report_period"] = f"{year} {quarter} квартал"
    active["expected_progress"] = active["report_quarter_num"] * 25

    period_requests = latest_approved_records(requests_df, year, quarter)
    merge_cols = [
        "strat_code", "status", "numeric_value", "risks", "progress_text",
        "submitted_at", "responsible_person", "phone", "email",
        "file_names", "file_urls", "admin_comment"
    ]
    for col in merge_cols:
        if col not in period_requests.columns:
            period_requests[col] = ""

    active = active.merge(period_requests[merge_cols], left_on="code", right_on="strat_code", how="left")
    active["status"] = active["status"].fillna("Не подано").replace("", "Не подано")
    active["numeric_value"] = active["numeric_value"].fillna("")
    active["risks"] = active["risks"].fillna("")
    active["progress_text"] = active["progress_text"].fillna("")
    active["submitted_at"] = active["submitted_at"].fillna("")

    active["status_score"] = active["status"].apply(status_score)
    active["plan_fact_percent"] = active.apply(
        lambda r: plan_fact_percent(r["numeric_value"], r["selected_target"]), axis=1
    )
    active["performance_score"] = active.apply(
        lambda r: r["plan_fact_percent"] if pd.notna(r["plan_fact_percent"]) else r["status_score"], axis=1
    )
    active["period_deviation"] = active["performance_score"] - active["expected_progress"]
    active["traffic_light"] = active["performance_score"].apply(traffic_light)
    active["has_submission"] = active["status"] != "Не подано"
    active["has_text_risk"] = active["risks"].astype(str).str.strip() != ""
    active["is_problem_status"] = active["status"].isin(["Потребує уваги", "Прострочено", "Не розпочато", "Не подано"])
    return active


def prepare_analysis_data(strat_df, requests_df, years, quarters):
    measures = base_measures(strat_df)
    parts = [
        prepare_period_slice(measures, requests_df, year, quarter)
        for year in years
        for quarter in quarters
    ]
    if not parts:
        return pd.DataFrame()
    return pd.concat(parts, ignore_index=True)


def apply_dimension_filters(data, selected_ssp, selected_deputies, selected_goals, selected_tasks, selected_product_types):
    filtered = data.copy()
    if selected_ssp:
        filtered = filtered[filtered["ssp_index"].astype(str).isin({str(x) for x in selected_ssp})]
    if selected_deputies:
        filtered = filtered[filtered["deputy_minister"].astype(str).isin(selected_deputies)]
    if selected_goals:
        filtered = filtered[filtered["goal_code"].astype(str).isin(selected_goals)]
    if selected_tasks:
        filtered = filtered[filtered["task_code"].astype(str).isin(selected_tasks)]
    if selected_product_types:
        filtered = filtered[filtered["product_type"].astype(str).isin(selected_product_types)]
    return filtered.copy()


# ============================================================
# Aggregations
# ============================================================

def build_metrics(active):
    total = len(active)
    submitted = int(active["has_submission"].sum()) if total else 0
    coverage = round(submitted / total * 100, 1) if total else 0
    completion = round(active["performance_score"].mean(), 1) if total else 0
    expected = round(active["expected_progress"].mean(), 1) if total else 0
    deviation = round(completion - expected, 1)
    unique_measures = active["code"].nunique() if total else 0
    goals = active["goal_code"].nunique() if total else 0
    tasks = active["task_code"].nunique() if total else 0
    no_data = int((active["status"] == "Не подано").sum()) if total else 0
    completed = int((active["status"] == "Виконано").sum()) if total else 0
    problem = int((active["is_problem_status"] | (active["period_deviation"] < -25)).sum()) if total else 0
    return {
        "total_rows": total, "unique_measures": unique_measures,
        "submitted": submitted, "coverage": coverage,
        "completion": completion, "expected": expected,
        "deviation": deviation, "goals": goals, "tasks": tasks,
        "no_data": no_data, "completed": completed, "problem": problem,
    }


def aggregate_goal_progress(active):
    if active.empty:
        return pd.DataFrame()
    result = (
        active.groupby(["goal_code", "strategic_goal"], dropna=False)
        .agg(
            Заходів_періодів=("code", "count"),
            Унікальних_заходів=("code", "nunique"),
            Виконання=("performance_score", "mean"),
            Очікуваний_темп=("expected_progress", "mean"),
            Подано=("has_submission", "sum"),
            Без_даних=("status", lambda x: (x == "Не подано").sum()),
            Проблемних=("is_problem_status", "sum"),
        ).reset_index()
    )
    result["Виконання"] = result["Виконання"].round(1)
    result["Очікуваний_темп"] = result["Очікуваний_темп"].round(1)
    result["Відхилення"] = (result["Виконання"] - result["Очікуваний_темп"]).round(1)
    result["Покриття_%"] = (result["Подано"] / result["Заходів_періодів"] * 100).round(1)
    return result


def aggregate_dep_progress(active):
    if active.empty:
        return pd.DataFrame()
    result = (
        active.groupby(["ssp_index", "department", "deputy_minister"], dropna=False)
        .agg(
            Заходів_періодів=("code", "count"),
            Унікальних_заходів=("code", "nunique"),
            Виконання=("performance_score", "mean"),
            Очікуваний_темп=("expected_progress", "mean"),
            Подано=("has_submission", "sum"),
            Без_даних=("status", lambda x: (x == "Не подано").sum()),
            Проблемних=("is_problem_status", "sum"),
        ).reset_index()
        .sort_values("ssp_index", key=lambda s: s.apply(lambda x: int(x) if str(x).isdigit() else 10_000))
    )
    result["Виконання"] = result["Виконання"].round(1)
    result["Очікуваний_темп"] = result["Очікуваний_темп"].round(1)
    result["Відхилення"] = (result["Виконання"] - result["Очікуваний_темп"]).round(1)
    result["Покриття_%"] = (result["Подано"] / result["Заходів_періодів"] * 100).round(1)
    return result


def aggregate_task_progress(active):
    if active.empty:
        return pd.DataFrame()
    result = (
        active.groupby(["goal_code", "task_code", "task_name"], dropna=False)
        .agg(
            Заходів_періодів=("code", "count"),
            Унікальних_заходів=("code", "nunique"),
            Виконання=("performance_score", "mean"),
            Очікуваний_темп=("expected_progress", "mean"),
            Подано=("has_submission", "sum"),
            Без_даних=("status", lambda x: (x == "Не подано").sum()),
        ).reset_index()
    )
    result["Виконання"] = result["Виконання"].round(1)
    result["Очікуваний_темп"] = result["Очікуваний_темп"].round(1)
    result["Відхилення"] = (result["Виконання"] - result["Очікуваний_темп"]).round(1)
    result["Покриття_%"] = (result["Подано"] / result["Заходів_періодів"] * 100).round(1)
    return result.sort_values(["goal_code", "task_code"])


def aggregate_product_progress(active):
    if active.empty:
        return pd.DataFrame()
    result = (
        active.groupby("product_type", dropna=False)
        .agg(
            Заходів_періодів=("code", "count"),
            Унікальних_заходів=("code", "nunique"),
            Виконання=("performance_score", "mean"),
            Очікуваний_темп=("expected_progress", "mean"),
            Подано=("has_submission", "sum"),
            Без_даних=("status", lambda x: (x == "Не подано").sum()),
        ).reset_index()
    )
    result["product_type"] = result["product_type"].replace("", "н/д")
    result["Виконання"] = result["Виконання"].round(1)
    result["Очікуваний_темп"] = result["Очікуваний_темп"].round(1)
    result["Відхилення"] = (result["Виконання"] - result["Очікуваний_темп"]).round(1)
    result["Покриття_%"] = (result["Подано"] / result["Заходів_періодів"] * 100).round(1)
    return result.sort_values("Заходів_періодів", ascending=False)


def aggregate_status(active):
    if active.empty:
        return pd.DataFrame(columns=["status", "Кількість"])
    return active.groupby("status").size().reset_index(name="Кількість").sort_values("Кількість", ascending=False)


def aggregate_period_dynamics(active):
    if active.empty:
        return pd.DataFrame()
    result = (
        active.groupby(["report_year", "report_quarter", "report_quarter_num"], dropna=False)
        .agg(
            Заходів_періодів=("code", "count"),
            Виконання=("performance_score", "mean"),
            Очікуваний_темп=("expected_progress", "mean"),
            Подано=("has_submission", "sum"),
        ).reset_index()
        .sort_values(["report_year", "report_quarter_num"])
    )
    result["Період"] = result["report_year"].astype(str) + " " + result["report_quarter"].astype(str)
    result["Виконання"] = result["Виконання"].round(1)
    result["Очікуваний_темп"] = result["Очікуваний_темп"].round(1)
    result["Відхилення"] = (result["Виконання"] - result["Очікуваний_темп"]).round(1)
    result["Покриття_%"] = (result["Подано"] / result["Заходів_періодів"] * 100).round(1)
    return result


# ============================================================
# Analytical text generator
# ============================================================

def generate_analytical_text(active, filters, metrics, goal_progress, dep_progress,
                              task_progress, product_progress, status_counts, period_dynamics):
    years_text = ", ".join(map(str, filters["years"]))
    quarters_text = ", ".join(filters["quarters"])
    selected_scope = (
        f"роки: {years_text}; квартали: {quarters_text}; "
        f"ССП: {filter_label(filters['ssp'], 'усі')}; "
        f"заступники: {filter_label(filters['deputies'], 'усі')}; "
        f"стратегічні цілі: {filter_label(filters['goal_labels'], 'усі')}; "
        f"типи продукту: {filter_label(filters['product_types'], 'усі')}"
    )

    completion = metrics["completion"]
    coverage = metrics["coverage"]
    deviation = metrics["deviation"]
    expected = metrics["expected"]

    if completion >= 70 and deviation >= -10:
        general_assessment = (
            "Загальний стан виконання можна оцінити як контрольований: середній фактичний рівень виконання "
            "близький до очікуваного квартального темпу або перевищує його."
        )
    elif completion >= 40:
        general_assessment = (
            "Стан виконання характеризується помірними відхиленнями: частина заходів рухається в межах "
            "очікуваного темпу, однак окремі напрями потребують додаткового управлінського контролю."
        )
    else:
        general_assessment = (
            "Стан виконання свідчить про суттєві відхилення від очікуваного темпу та потребує концентрації уваги "
            "на причинах затримок, неповного подання даних і недостатньої фактичної динаміки."
        )

    if coverage >= 80:
        coverage_assessment = "Інформаційна база є достатньою для формування узагальнених управлінських висновків."
    elif coverage >= 40:
        coverage_assessment = ("Інформаційна база є частковою, тому окремі висновки варто інтерпретувати "
                               "з урахуванням неповного покриття моніторингом.")
    else:
        coverage_assessment = ("Інформаційна база є недостатньою; це знижує точність оцінки та підвищує "
                               "ризик викривлення загальної картини виконання.")

    best_goal_text = worst_goal_text = attention_goal_text = "н/д"
    if not goal_progress.empty:
        best_goal = goal_progress.sort_values(["Виконання", "Покриття_%"], ascending=False).iloc[0]
        worst_goal = goal_progress.sort_values(["Виконання", "Покриття_%"], ascending=True).iloc[0]
        attention_goal = goal_progress.sort_values(["Відхилення", "Без_даних"], ascending=[True, False]).iloc[0]
        best_goal_text = f"СЦ {best_goal['goal_code']} — {round(best_goal['Виконання'],1)}%, покриття {round(best_goal['Покриття_%'],1)}%"
        worst_goal_text = f"СЦ {worst_goal['goal_code']} — {round(worst_goal['Виконання'],1)}%, покриття {round(worst_goal['Покриття_%'],1)}%"
        attention_goal_text = f"СЦ {attention_goal['goal_code']} — відхилення {deviation_label(attention_goal['Відхилення'])}, без даних {int(attention_goal['Без_даних'])}"

    best_dep_text = worst_dep_text = "н/д"
    if not dep_progress.empty:
        best_dep = dep_progress.sort_values(["Виконання", "Покриття_%"], ascending=False).iloc[0]
        worst_dep = dep_progress.sort_values(["Виконання", "Покриття_%"], ascending=True).iloc[0]
        best_dep_text = f"{best_dep['department']} — {round(best_dep['Виконання'],1)}%"
        worst_dep_text = f"{worst_dep['department']} — {round(worst_dep['Виконання'],1)}%, без даних {int(worst_dep['Без_даних'])}"

    task_attention_text = "н/д"
    if not task_progress.empty:
        task_attention = task_progress.sort_values(["Відхилення", "Без_даних"], ascending=[True, False]).head(3)
        task_attention_text = concise_list([
            f"{row['task_code']} — {round(row['Виконання'],1)}%, відхилення {deviation_label(row['Відхилення'])}"
            for _, row in task_attention.iterrows()
        ], limit=3)

    product_text = "н/д"
    if not product_progress.empty:
        product_text = concise_list([
            f"{row['product_type']} — {int(row['Унікальних_заходів'])} заходів, виконання {round(row['Виконання'],1)}%"
            for _, row in product_progress.head(4).iterrows()
        ], limit=4)

    status_text = "н/д"
    if not status_counts.empty:
        status_text = concise_list([
            f"{row['status']} — {int(row['Кількість'])}"
            for _, row in status_counts.iterrows()
        ], limit=6)

    dynamics_text = "н/д"
    if not period_dynamics.empty:
        dynamics_text = concise_list([
            f"{row['Період']}: виконання {round(row['Виконання'],1)}%, покриття {round(row['Покриття_%'],1)}%, відхилення {deviation_label(row['Відхилення'])}"
            for _, row in period_dynamics.iterrows()
        ], limit=6)

    return f"""
За результатами автоматизованого аналізу сформовано аналітичну довідку щодо стану виконання Стратегічного плану за обраним зрізом. Параметри аналізу: {selected_scope}. У межах відібраного масиву враховано {metrics['total_rows']} записів «захід-період», що відповідають {metrics['unique_measures']} унікальним заходам, {metrics['tasks']} завданням та {metrics['goals']} стратегічним цілям.

Середній розрахунковий рівень виконання Стратегічного плану в обраному періоді становить {completion}%. Очікуваний темп для відповідних кварталів становить {expected}%, тому відхилення у звітному періоді дорівнює {deviation_label(deviation)}. {general_assessment}

Покриття моніторингом становить {coverage}%: подано та погоджено дані за {metrics['submitted']} записами з {metrics['total_rows']}. Без поданих погоджених даних залишаються {metrics['no_data']} записів. {coverage_assessment}

У динаміці за періодами: {dynamics_text}. Цей блок показує, чи накопичується відставання протягом року, чи навпаки спостерігається поступове наближення до планового темпу виконання.

У розрізі стратегічних цілей найвищий рівень виконання зафіксовано за напрямом {best_goal_text}. Найнижчий — за напрямом {worst_goal_text}. Окремої уваги потребує {attention_goal_text}.

У розрізі завдань першочергової уваги потребують: {task_attention_text}.

У розрізі самостійних структурних підрозділів найкращий результат демонструє {best_dep_text}. Найнижчий показник — у {worst_dep_text}.

За типами продукту: {product_text}. За статусами виконання: {status_text}.

З огляду на результати аналізу доцільно зосередити подальшу роботу на трьох напрямах: забезпечити повноту подання даних за заходами без погодженого моніторингу; уточнити причини відхилень у стратегічних цілях та завданнях із найнижчим темпом виконання; підготувати пропозиції щодо коригування строків, відповідальних або змісту заходів там, де фактичний прогрес системно не відповідає очікуваному квартальному темпу.
""".strip()


# ============================================================
# Export functions
# ============================================================

def create_excel_report(active, period_requests, goal_progress, dep_progress, task_progress,
                        product_progress, status_counts, period_dynamics, metrics, filters):
    output = BytesIO()
    active_export = active.rename(columns={
        "report_year": "Рік", "report_quarter": "Квартал",
        "code": "Код заходу", "name": "Захід",
        "goal_code": "Код СЦ", "strategic_goal": "Стратегічна ціль",
        "task_code": "Код завдання", "task_name": "Завдання",
        "product_type": "Тип продукту", "department": "ССП",
        "deputy_minister": "Заступник Міністра",
        "indicator": "Індикатор", "unit": "Одиниця виміру",
        "selected_target": "Планове значення", "numeric_value": "Фактичне значення",
        "status": "Статус", "performance_score": "Виконання, %",
        "expected_progress": "Очікуваний темп, %", "period_deviation": "Відхилення, в.п.",
        "traffic_light": "Оцінка темпу", "progress_text": "Пояснення",
        "risks": "Ризики/відхилення",
    })
    active_cols = [
        "Рік", "Квартал", "Код заходу", "Захід", "Код СЦ", "Стратегічна ціль",
        "Код завдання", "Завдання", "Тип продукту", "ССП", "Заступник Міністра",
        "Індикатор", "Одиниця виміру", "Планове значення", "Фактичне значення",
        "Статус", "Виконання, %", "Очікуваний темп, %", "Відхилення, в.п.",
        "Оцінка темпу", "Пояснення", "Ризики/відхилення"
    ]
    summary_df = pd.DataFrame([
        ["Період", f"Роки: {', '.join(map(str, filters['years']))}; квартали: {', '.join(filters['quarters'])}"],
        ["ССП", filter_label(filters["ssp"], "Усі")],
        ["Заступники Міністра", filter_label(filters["deputies"], "Усі")],
        ["Стратегічні цілі", filter_label(filters["goal_labels"], "Усі")],
        ["Типи продукту", filter_label(filters["product_types"], "Усі")],
        ["Дата формування", datetime.now().strftime("%d.%m.%Y %H:%M")],
        ["Унікальних заходів", metrics["unique_measures"]],
        ["Покриття моніторингом", f"{metrics['coverage']}%"],
        ["Виконання СП", f"{metrics['completion']}%"],
        ["Відхилення", deviation_label(metrics["deviation"])],
    ], columns=["Показник", "Значення"])

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        summary_df.to_excel(writer, sheet_name="Пояснення", index=False)
        active_export[[c for c in active_cols if c in active_export.columns]].to_excel(writer, sheet_name="Аналітичний масив", index=False)
        goal_progress.to_excel(writer, sheet_name="Стратегічні цілі", index=False)
        task_progress.to_excel(writer, sheet_name="Завдання", index=False)
        dep_progress.to_excel(writer, sheet_name="ССП", index=False)
        product_progress.to_excel(writer, sheet_name="Типи продукту", index=False)
        period_dynamics.to_excel(writer, sheet_name="Динаміка", index=False)
        status_counts.to_excel(writer, sheet_name="Статуси", index=False)
        period_requests.to_excel(writer, sheet_name="Реєстр заявок", index=False)
        workbook = writer.book
        header_fmt = workbook.add_format({"bold": True, "bg_color": "#D9EAF7", "border": 1})
        sheet_map = {
            "Пояснення": summary_df,
            "Аналітичний масив": active_export[[c for c in active_cols if c in active_export.columns]],
            "Стратегічні цілі": goal_progress, "Завдання": task_progress,
            "ССП": dep_progress, "Типи продукту": product_progress,
            "Динаміка": period_dynamics, "Статуси": status_counts, "Реєстр заявок": period_requests,
        }
        for sheet_name, worksheet in writer.sheets.items():
            worksheet.freeze_panes(1, 0)
            worksheet.set_column(0, 30, 20)
            worksheet.autofilter(0, 0, 0, 20)
            try:
                ws_df = sheet_map[sheet_name]
                for col_num, value in enumerate(ws_df.columns.values):
                    worksheet.write(0, col_num, value, header_fmt)
            except Exception:
                pass
    output.seek(0)
    return output


def create_docx_report(text, metrics, filters, goal_progress, dep_progress, product_progress):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("АНАЛІТИЧНА ДОВІДКА")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("щодо стану виконання Стратегічного плану").font.size = Pt(12)

    scope = document.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.add_run(
        f"Роки: {', '.join(map(str, filters['years']))}; квартали: {', '.join(filters['quarters'])}"
    ).italic = True

    document.add_paragraph("Ключові показники").runs[0].bold = True
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Показник"
    hdr[1].text = "Значення"
    for key, value in {
        "Унікальних заходів": metrics["unique_measures"],
        "Покриття моніторингом": f"{metrics['coverage']}%",
        "Рівень виконання СП": f"{metrics['completion']}%",
        "Очікуваний темп": f"{metrics['expected']}%",
        "Відхилення": deviation_label(metrics["deviation"]),
        "Без погоджених даних": metrics["no_data"],
    }.items():
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    document.add_paragraph("Аналітичний висновок").runs[0].bold = True
    for paragraph in text.split("\n\n"):
        p = document.add_paragraph(paragraph.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.size = Pt(11)

    document.add_paragraph("Додаткова структура даних").runs[0].bold = True
    document.add_paragraph(f"Стратегічних цілей в аналізі: {len(goal_progress)}.")
    document.add_paragraph(f"ССП в аналізі: {len(dep_progress)}.")
    document.add_paragraph(f"Типів продукту: {len(product_progress)}.")

    document.add_paragraph(
        "Сформовано автоматизованою системою моніторингу стратегічного плану."
    ).italic = True
    document.add_paragraph(
        "Розроблено департаментом стратегічного планування та макроекономічного прогнозування."
    ).italic = True

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


# ============================================================
# ── Load data ──
# ============================================================

strat_df = load_strat_matrix()
requests_df = ensure_request_columns(load_requests())
measures_all = base_measures(strat_df)

if measures_all.empty:
    st.warning("У стратегічній матриці не знайдено заходів для аналізу.")
    st.stop()

# ============================================================
# ── Filter options ──
# ============================================================

ssp_options_df = (
    measures_all[["ssp_index", "department"]].dropna().drop_duplicates()
    .sort_values("department", key=lambda s: s.apply(ssp_sort_key))
)
ssp_options = [raw_value(x) for x in ssp_options_df["ssp_index"].tolist() if raw_value(x)]
ssp_labels = {
    raw_value(row["ssp_index"]): raw_value(row["department"])
    for _, row in ssp_options_df.iterrows()
    if raw_value(row["ssp_index"])
}

goal_rows = measures_all[["goal_code", "strategic_goal"]].drop_duplicates().sort_values("goal_code")
goal_options = {
    raw_value(row["goal_code"]): f"{raw_value(row['goal_code'])} {raw_value(row['strategic_goal'])}"
    for _, row in goal_rows.iterrows()
    if raw_value(row["goal_code"])
}

task_rows = measures_all[["task_code", "task_name"]].drop_duplicates().sort_values("task_code")
task_options = {
    raw_value(row["task_code"]): f"{raw_value(row['task_code'])} {raw_value(row['task_name'])}"
    for _, row in task_rows.iterrows()
    if raw_value(row["task_code"])
}

deputy_options = sorted([x for x in measures_all["deputy_minister"].dropna().astype(str).unique() if raw_value(x)])
product_type_options = sorted([x for x in measures_all["product_type"].dropna().astype(str).unique() if raw_value(x)])


# ============================================================
# ── HEADER ──
# ============================================================

st.markdown('<div class="ua-line"></div>', unsafe_allow_html=True)
st.markdown(
    '<div class="ministry-label">🇺🇦 Міністерство економіки, довкілля та сільського господарства України</div>',
    unsafe_allow_html=True
)

# Quick system-wide snapshot (before any filters)
snap_total = measures_all["code"].nunique()
snap_requests = len(requests_df) if not requests_df.empty else 0
snap_approved = int((requests_df["approval_status"] == "Погоджено").sum()) if not requests_df.empty and "approval_status" in requests_df.columns else 0
snap_coverage = round(snap_approved / snap_total * 100, 1) if snap_total else 0
snap_updated = "—"
if not requests_df.empty and "submitted_at" in requests_df.columns:
    try:
        last_ts = pd.to_datetime(requests_df["submitted_at"], errors="coerce").max()
        if pd.notna(last_ts):
            snap_updated = last_ts.strftime("%d.%m.%Y %H:%M")
    except Exception:
        pass

st.markdown(f"""
<div class="header-box">
    <div class="header-title">Аналітичні відомості</div>
    <div class="header-subtitle">
        Формування управлінських матеріалів на основі результатів моніторингу й оцінювання
        стратегічних результатів — для аналізу тенденцій, виявлення відхилень та підготовки
        пропозицій щодо коригування Стратегічного плану.
    </div>
    <div class="badge-wrap">
        <div class="badge">● Аналітична довідка</div>
        <div class="badge">● Динаміка виконання</div>
        <div class="badge">● СЦ / Завдання / ССП</div>
        <div class="badge badge-green">● DOCX та Excel</div>
        <div class="badge badge-yellow">● Оновлено: {snap_updated}</div>
    </div>
</div>
""", unsafe_allow_html=True)

# ── System-wide summary banner (always visible, before filters) ──
st.markdown(f"""
<div class="summary-banner">
    <div class="summary-banner-title">Загальний стан системи моніторингу (без фільтрів)</div>
    <div class="summary-grid">
        <div class="summary-cell blue">
            <div class="summary-cell-label">Заходів у СП</div>
            <div class="summary-cell-value">{snap_total}</div>
            <div class="summary-cell-note">унікальних кодів</div>
        </div>
        <div class="summary-cell gray">
            <div class="summary-cell-label">Заявок у системі</div>
            <div class="summary-cell-value">{snap_requests}</div>
            <div class="summary-cell-note">усіх записів</div>
        </div>
        <div class="summary-cell green">
            <div class="summary-cell-label">Погоджено</div>
            <div class="summary-cell-value">{snap_approved}</div>
            <div class="summary-cell-note">записів зі статусом «Погоджено»</div>
        </div>
        <div class="summary-cell yellow">
            <div class="summary-cell-label">Покриття системи</div>
            <div class="summary-cell-value">{snap_coverage}%</div>
            <div class="summary-cell-note">погоджено від загальної кількості заходів</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)


# ============================================================
# ── FILTERS ──
# ============================================================

st.markdown("""
<div class="filter-box">
    <div class="filter-box-title">Параметри звіту</div>
    <div class="filter-box-hint">
        Оберіть роки, квартали та організаційний зріз. Якщо рік або квартал не обрано —
        система автоматично бере повний 2026 рік. Фільтр заступника Міністра та стратегічної цілі
        дозволяє швидко сформувати звіт для конкретного напряму.
    </div>
    <div class="filter-section-title">Звітний період</div>
""", unsafe_allow_html=True)

p1, p2 = st.columns([1, 1])
with p1:
    selected_years_raw = st.multiselect("Рік", YEAR_OPTIONS, default=[], placeholder="Усі або обрані роки")
with p2:
    selected_quarters_raw = st.multiselect("Квартал", QUARTERS, default=[], placeholder="Усі або обрані квартали")

st.markdown('<div class="filter-section-title">Організаційний зріз</div>', unsafe_allow_html=True)

oa, ob = st.columns([1.2, 1.2])
with oa:
    selected_ssp_indices = st.multiselect(
        "Самостійний структурний підрозділ",
        ssp_options,
        format_func=lambda x: ssp_labels.get(x, x),
        placeholder="Оберіть один або декілька ССП"
    )
with ob:
    selected_deputies = st.multiselect(
        "Заступник Міністра",
        deputy_options,
        placeholder="Оберіть заступника Міністра"
    )

st.markdown('<div class="filter-section-title">Змістовий зріз</div>', unsafe_allow_html=True)

ca, cb, cc = st.columns([1.3, 1.3, 1])
with ca:
    selected_goal_labels = st.multiselect(
        "Стратегічна ціль",
        list(goal_options.values()),
        placeholder="Оберіть стратегічну ціль"
    )
with cb:
    selected_task_labels = st.multiselect(
        "Завдання",
        list(task_options.values()),
        placeholder="Оберіть завдання"
    )
with cc:
    selected_product_types = st.multiselect(
        "Тип продукту",
        product_type_options,
        placeholder="Оберіть тип"
    )

st.markdown('</div>', unsafe_allow_html=True)

# Defaults
selected_years = selected_years_raw if selected_years_raw else [2026]
selected_quarters = selected_quarters_raw if selected_quarters_raw else QUARTERS.copy()
selected_goal_codes = [code for code, label in goal_options.items() if label in selected_goal_labels]
selected_task_codes = [code for code, label in task_options.items() if label in selected_task_labels]
selected_ssp_labels_display = [ssp_labels.get(x, x) for x in selected_ssp_indices]


# ============================================================
# ── Compute data ──
# ============================================================

all_period_data = prepare_analysis_data(strat_df, requests_df, selected_years, selected_quarters)
active = apply_dimension_filters(
    all_period_data,
    selected_ssp_indices, selected_deputies,
    selected_goal_codes, selected_task_codes, selected_product_types,
)

if active.empty:
    st.warning("За обраними параметрами активних заходів не знайдено. Спробуйте розширити фільтри.")
    st.stop()

period_requests = requests_df.copy()
if not period_requests.empty:
    period_requests = period_requests[
        period_requests["year"].astype(str).isin([str(y) for y in selected_years]) &
        period_requests["quarter"].astype(str).isin([str(q) for q in selected_quarters])
    ].copy()
    if selected_ssp_indices:
        period_requests["department_index"] = period_requests["department"].apply(extract_ssp_index)
        period_requests = period_requests[
            period_requests["department_index"].astype(str).isin(set(selected_ssp_indices))
        ]

metrics = build_metrics(active)
goal_progress = aggregate_goal_progress(active)
dep_progress = aggregate_dep_progress(active)
task_progress = aggregate_task_progress(active)
product_progress = aggregate_product_progress(active)
status_counts = aggregate_status(active)
period_dynamics = aggregate_period_dynamics(active)

filters = {
    "years": selected_years, "quarters": selected_quarters,
    "ssp": selected_ssp_labels_display, "ssp_indices": selected_ssp_indices,
    "deputies": selected_deputies,
    "goal_labels": selected_goal_labels, "task_labels": selected_task_labels,
    "product_types": selected_product_types,
}

analytical_text = generate_analytical_text(
    active, filters, metrics, goal_progress, dep_progress,
    task_progress, product_progress, status_counts, period_dynamics,
)


# ============================================================
# ── SECTION 1: KPI cards for the selected slice ──
# ============================================================

dev_color = "green" if metrics["deviation"] >= 0 else ("yellow" if metrics["deviation"] >= -15 else "red")
dev_arrow = "▲" if metrics["deviation"] > 0 else ("▼" if metrics["deviation"] < 0 else "—")

st.markdown(f"""
<div class="section-card primary">
    <div class="section-card-title">Показники обраного зрізу</div>
    <div class="section-card-sub">
        {', '.join(map(str, selected_years))} рік · квартали {', '.join(selected_quarters)} ·
        ССП: {filter_label(selected_ssp_labels_display, 'усі')} ·
        Заступники: {filter_label(selected_deputies, 'усі')}
    </div>
    <div class="summary-grid">
        <div class="summary-cell blue">
            <div class="summary-cell-label">Виконання СП</div>
            <div class="summary-cell-value">{metrics['completion']}%</div>
            <div class="summary-cell-note">середній фактичний прогрес</div>
        </div>
        <div class="summary-cell {dev_color}">
            <div class="summary-cell-label">Відхилення {dev_arrow}</div>
            <div class="summary-cell-value">{deviation_label(metrics['deviation'])}</div>
            <div class="summary-cell-note">від очікуваного темпу {metrics['expected']}%</div>
        </div>
        <div class="summary-cell yellow">
            <div class="summary-cell-label">Покриття моніторингом</div>
            <div class="summary-cell-value">{metrics['coverage']}%</div>
            <div class="summary-cell-note">{metrics['submitted']} із {metrics['total_rows']} записів</div>
        </div>
        <div class="summary-cell blue">
            <div class="summary-cell-label">Унікальних заходів</div>
            <div class="summary-cell-value">{metrics['unique_measures']}</div>
            <div class="summary-cell-note">записів захід-період: {metrics['total_rows']}</div>
        </div>
        <div class="summary-cell gray">
            <div class="summary-cell-label">Стратегічних цілей</div>
            <div class="summary-cell-value">{metrics['goals']}</div>
            <div class="summary-cell-note">у відібраному масиві</div>
        </div>
        <div class="summary-cell gray">
            <div class="summary-cell-label">Завдань</div>
            <div class="summary-cell-value">{metrics['tasks']}</div>
            <div class="summary-cell-note">у відібраному масиві</div>
        </div>
        <div class="summary-cell red">
            <div class="summary-cell-label">Без даних</div>
            <div class="summary-cell-value">{metrics['no_data']}</div>
            <div class="summary-cell-note">записів без погодженого подання</div>
        </div>
        <div class="summary-cell {'red' if metrics['problem'] > 0 else 'green'}">
            <div class="summary-cell-label">Проблемних</div>
            <div class="summary-cell-value">{metrics['problem']}</div>
            <div class="summary-cell-note">заходів із відхиленням &gt;25 в.п.</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)


# ============================================================
# ── SECTION 2: Analytical note ──
# ============================================================

meta_chips = " ".join([
    f'<span class="report-meta-chip">📅 {", ".join(map(str, selected_years))}</span>',
    f'<span class="report-meta-chip">🗓 {", ".join(selected_quarters)}</span>',
    f'<span class="report-meta-chip">🏢 {filter_label(selected_ssp_labels_display, "усі ССП")}</span>',
    f'<span class="report-meta-chip">👤 {filter_label(selected_deputies, "усі заступники")}</span>',
    f'<span class="report-meta-chip">🕒 {datetime.now().strftime("%d.%m.%Y %H:%M")}</span>',
])

st.markdown(f"""
<div class="section-card primary">
    <div class="section-card-title">Автоматично сформована аналітична довідка</div>
    <div class="report-meta">{meta_chips}</div>
""", unsafe_allow_html=True)

for paragraph in analytical_text.split("\n\n"):
    st.markdown(f"<p class='report-text'>{clean(paragraph)}</p>", unsafe_allow_html=True)

st.markdown("</div>", unsafe_allow_html=True)


# ============================================================
# ── SECTION 3: Charts ──
# ============================================================

PLOTLY_COLORS = ["#2563eb", "#16a34a", "#d97706", "#dc2626", "#7c3aed",
                 "#0891b2", "#db2777", "#65a30d", "#ea580c", "#475569"]

st.markdown("""
<div class="section-card">
    <div class="section-card-title">Графіки</div>
    <div class="section-card-sub">Сформовані за тією самою вибіркою, що й текст довідки.</div>
""", unsafe_allow_html=True)

# Row 1: dynamics + goal progress
r1a, r1b = st.columns(2)

with r1a:
    st.markdown('<div class="chart-wrap"><div class="chart-title">Динаміка виконання проти очікуваного темпу</div><div class="chart-sub">Порівняння факт / план по кожному кварталу</div>', unsafe_allow_html=True)
    if not period_dynamics.empty:
        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=period_dynamics["Період"], y=period_dynamics["Виконання"],
            name="Виконання", mode="lines+markers",
            line=dict(color="#2563eb", width=2.5),
            marker=dict(size=8, color="#2563eb"),
        ))
        fig.add_trace(go.Scatter(
            x=period_dynamics["Період"], y=period_dynamics["Очікуваний_темп"],
            name="Очікуваний темп", mode="lines+markers",
            line=dict(color="#d97706", width=2, dash="dot"),
            marker=dict(size=7, color="#d97706"),
        ))
        fig.update_layout(**plotly_theme(), legend=dict(orientation="h", y=-0.15))
        fig.update_yaxes(range=[0, 110], ticksuffix="%")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Недостатньо даних для побудови динаміки.")
    st.markdown('</div>', unsafe_allow_html=True)

with r1b:
    st.markdown('<div class="chart-wrap"><div class="chart-title">Виконання за стратегічними цілями</div><div class="chart-sub">Порівняння факт vs очікуваний темп</div>', unsafe_allow_html=True)
    if not goal_progress.empty:
        gp = goal_progress.sort_values("goal_code").copy()
        fig = go.Figure()
        fig.add_trace(go.Bar(
            name="Виконання",
            x=gp["goal_code"], y=gp["Виконання"],
            marker_color="#2563eb",
            text=gp["Виконання"].apply(lambda v: f"{v}%"),
            textposition="outside",
            customdata=list(zip(gp["strategic_goal"], gp["Покриття_%"], gp["Відхилення"], gp["Без_даних"])),
            hovertemplate="<b>%{x}</b><br>%{customdata[0]}<br>Виконання: %{y}%<br>Покриття: %{customdata[1]}%<br>Відхилення: %{customdata[2]} в.п.<br>Без даних: %{customdata[3]}<extra></extra>",
        ))
        fig.add_trace(go.Bar(
            name="Очікуваний темп",
            x=gp["goal_code"], y=gp["Очікуваний_темп"],
            marker_color="rgba(217,119,6,0.35)",
            text=gp["Очікуваний_темп"].apply(lambda v: f"{v}%"),
            textposition="outside",
        ))
        fig.update_layout(**plotly_theme(), barmode="group", legend=dict(orientation="h", y=-0.15))
        fig.update_yaxes(range=[0, 120], ticksuffix="%")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Немає даних по стратегічних цілях.")
    st.markdown('</div>', unsafe_allow_html=True)

# Row 2: per-SSP + status donut
r2a, r2b = st.columns(2)

with r2a:
    st.markdown('<div class="chart-wrap"><div class="chart-title">Виконання за самостійними структурними підрозділами</div><div class="chart-sub">Відхилення від очікуваного темпу по кожному ССП</div>', unsafe_allow_html=True)
    if not dep_progress.empty:
        dp = dep_progress.copy()
        dp["label"] = dp["ssp_index"].astype(str) + " " + dp["department"].astype(str).str.slice(0, 25)
        dp["color"] = dp["Відхилення"].apply(lambda v: "#16a34a" if v >= 0 else ("#d97706" if v >= -15 else "#dc2626"))
        fig = go.Figure(go.Bar(
            x=dp["Відхилення"], y=dp["label"],
            orientation="h",
            marker_color=dp["color"],
            text=dp["Відхилення"].apply(lambda v: f"{'+' if v>0 else ''}{round(v,1)} в.п."),
            textposition="outside",
            customdata=list(zip(dp["Виконання"], dp["Очікуваний_темп"], dp["Покриття_%"], dp["Без_даних"])),
            hovertemplate="<b>%{y}</b><br>Виконання: %{customdata[0]}%<br>Очікуваний темп: %{customdata[1]}%<br>Покриття: %{customdata[2]}%<br>Без даних: %{customdata[3]}<extra></extra>",
        ))
        fig.add_vline(x=0, line_color="#475569", line_width=1.5, line_dash="dot")
        fig.update_layout(**plotly_theme(), margin=dict(l=10, r=60, t=40, b=10))
        fig.update_xaxes(ticksuffix=" в.п.")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Немає даних по ССП.")
    st.markdown('</div>', unsafe_allow_html=True)

with r2b:
    st.markdown('<div class="chart-wrap"><div class="chart-title">Структура статусів виконання</div><div class="chart-sub">Розподіл усіх записів за статусом</div>', unsafe_allow_html=True)
    if not status_counts.empty:
        STATUS_COLORS = {
            "Виконано": "#16a34a", "Виконано частково": "#65a30d",
            "Виконується": "#2563eb", "Потребує уваги": "#d97706",
            "Не розпочато": "#94a3b8", "Прострочено": "#dc2626", "Не подано": "#e2e8f0",
        }
        colors = [STATUS_COLORS.get(s, "#cbd5e1") for s in status_counts["status"]]
        fig = go.Figure(go.Pie(
            labels=status_counts["status"],
            values=status_counts["Кількість"],
            hole=0.48,
            marker_colors=colors,
            textinfo="label+percent",
            insidetextorientation="radial",
        ))
        fig.update_layout(**plotly_theme(), showlegend=False)
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Немає даних про статуси.")
    st.markdown('</div>', unsafe_allow_html=True)

# Row 3: tasks worst deviation + product structure
r3a, r3b = st.columns(2)

with r3a:
    st.markdown('<div class="chart-wrap"><div class="chart-title">Завдання з найбільшим відхиленням</div><div class="chart-sub">10 завдань, де факт найбільше відстає від темпу</div>', unsafe_allow_html=True)
    if not task_progress.empty:
        top_tasks = task_progress.sort_values("Відхилення", ascending=True).head(10).copy()
        top_tasks["label"] = top_tasks["task_code"].astype(str)
        top_tasks["color"] = top_tasks["Відхилення"].apply(
            lambda v: "#16a34a" if v >= 0 else ("#d97706" if v >= -15 else "#dc2626")
        )
        fig = go.Figure(go.Bar(
            x=top_tasks["label"], y=top_tasks["Відхилення"],
            marker_color=top_tasks["color"],
            text=top_tasks["Відхилення"].apply(lambda v: f"{'+' if v>0 else ''}{round(v,1)}"),
            textposition="outside",
            customdata=list(zip(top_tasks["task_name"], top_tasks["Виконання"], top_tasks["Покриття_%"])),
            hovertemplate="<b>%{x}</b><br>%{customdata[0]}<br>Виконання: %{customdata[1]}%<br>Покриття: %{customdata[2]}%<extra></extra>",
        ))
        fig.add_hline(y=0, line_color="#475569", line_width=1.5, line_dash="dot")
        fig.update_layout(**plotly_theme())
        fig.update_yaxes(ticksuffix=" в.п.")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Немає даних по завданнях.")
    st.markdown('</div>', unsafe_allow_html=True)

with r3b:
    st.markdown('<div class="chart-wrap"><div class="chart-title">Структура заходів за типами продукту</div><div class="chart-sub">Кількість заходів та рівень виконання</div>', unsafe_allow_html=True)
    if not product_progress.empty:
        pp = product_progress.head(10).copy()
        fig = go.Figure()
        fig.add_trace(go.Bar(
            name="Заходів",
            x=pp["product_type"], y=pp["Унікальних_заходів"],
            marker_color="#2563eb",
            text=pp["Унікальних_заходів"],
            textposition="outside",
            yaxis="y",
        ))
        fig.add_trace(go.Scatter(
            name="Виконання, %",
            x=pp["product_type"], y=pp["Виконання"],
            mode="lines+markers",
            line=dict(color="#d97706", width=2.5),
            marker=dict(size=8),
            yaxis="y2",
        ))
        fig.update_layout(
            **plotly_theme(),
            yaxis=dict(title="Заходів", showgrid=True),
            yaxis2=dict(title="Виконання, %", overlaying="y", side="right", range=[0, 120], ticksuffix="%"),
            legend=dict(orientation="h", y=-0.2),
            xaxis=dict(tickangle=-20),
        )
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Немає даних по типах продукту.")
    st.markdown('</div>', unsafe_allow_html=True)

# Row 4 (if multiple periods): period comparison cells
if not period_dynamics.empty and len(period_dynamics) >= 2:
    st.markdown("""
    <div class="section-card warning">
        <div class="section-card-title">Порівняння кварталів «пліч-о-пліч»</div>
        <div class="section-card-sub">Зміна виконання між суміжними звітними кварталами</div>
    """, unsafe_allow_html=True)
    pd_rows = period_dynamics.reset_index(drop=True)
    cells_html = ""
    for i in range(1, len(pd_rows)):
        prev = pd_rows.iloc[i - 1]
        curr = pd_rows.iloc[i]
        delta = round(curr["Виконання"] - prev["Виконання"], 1)
        delta_class = "delta-up" if delta > 0 else ("delta-down" if delta < 0 else "delta-flat")
        delta_sign = "▲ +" if delta > 0 else ("▼ " if delta < 0 else "— ")
        cells_html += f"""
        <div class="comparison-cell">
            <div class="comparison-period">{prev['Період']} → {curr['Період']}</div>
            <div class="comparison-value">{curr['Виконання']}%</div>
            <div class="comparison-delta {delta_class}">{delta_sign}{delta} в.п.</div>
        </div>"""
    st.markdown(f'<div class="comparison-grid">{cells_html}</div>', unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

st.markdown("</div>", unsafe_allow_html=True)  # close charts section-card


# ============================================================
# ── SECTION 4: Export ──
# ============================================================

excel_file = create_excel_report(
    active, period_requests, goal_progress, dep_progress, task_progress,
    product_progress, status_counts, period_dynamics, metrics, filters,
)
docx_file = create_docx_report(
    analytical_text, metrics, filters, goal_progress, dep_progress, product_progress,
)
requests_output = BytesIO()
period_requests.to_excel(requests_output, index=False, engine="openpyxl")
requests_output.seek(0)

fname_suffix = f"{'_'.join(map(str, selected_years))}_{'_'.join(selected_quarters)}"

st.markdown("""
<div class="export-box">
    <div class="export-title">Експорт матеріалів</div>
    <div class="export-sub">Завантаження формуються відповідно до всіх обраних фільтрів.</div>
""", unsafe_allow_html=True)

e1, e2, e3 = st.columns(3)
with e1:
    st.download_button(
        "📊 Excel-звіт (9 аркушів)",
        data=excel_file,
        file_name=f"analytics_{fname_suffix}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )
with e2:
    st.download_button(
        "📄 Аналітична довідка DOCX",
        data=docx_file,
        file_name=f"analytical_note_{fname_suffix}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
with e3:
    st.download_button(
        "📋 Реєстр заявок",
        data=requests_output,
        file_name=f"requests_{fname_suffix}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )

st.markdown("</div>", unsafe_allow_html=True)


# ============================================================
# ── SECTION 5: Data tables (tabs) ──
# ============================================================

st.markdown("""
<div class="tables-section">
    <div class="tables-title">Деталізовані таблиці</div>
    <div class="tables-sub">Повний аналітичний масив за обраним зрізом. Доступний для сортування та пошуку.</div>
""", unsafe_allow_html=True)

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "Аналітичний масив",
    "Стратегічні цілі",
    "Завдання",
    "ССП",
    "Типи продукту",
    "Реєстр заявок",
])

with tab1:
    show_active = active.rename(columns={
        "report_year": "Рік", "report_quarter": "Квартал",
        "code": "Код", "name": "Захід",
        "strategic_goal": "Стратегічна ціль", "task_name": "Завдання",
        "product_type": "Тип продукту", "department": "ССП",
        "deputy_minister": "Заступник Міністра",
        "selected_target": "Планове значення", "numeric_value": "Фактичне значення",
        "status": "Статус", "performance_score": "Виконання, %",
        "expected_progress": "Очікуваний темп, %",
        "period_deviation": "Відхилення, в.п.", "traffic_light": "Оцінка темпу",
    })
    disp_cols = [
        "Рік", "Квартал", "Код", "Захід", "Стратегічна ціль", "Завдання",
        "Тип продукту", "ССП", "Заступник Міністра",
        "Планове значення", "Фактичне значення", "Статус",
        "Виконання, %", "Очікуваний темп, %", "Відхилення, в.п.", "Оцінка темпу",
    ]
    st.dataframe(
        show_active[[c for c in disp_cols if c in show_active.columns]],
        use_container_width=True, hide_index=True,
    )
with tab2:
    st.dataframe(goal_progress, use_container_width=True, hide_index=True)
with tab3:
    st.dataframe(task_progress, use_container_width=True, hide_index=True)
with tab4:
    st.dataframe(dep_progress, use_container_width=True, hide_index=True)
with tab5:
    st.dataframe(product_progress, use_container_width=True, hide_index=True)
with tab6:
    st.dataframe(period_requests, use_container_width=True, hide_index=True)

st.markdown("</div>", unsafe_allow_html=True)


# ============================================================
# ── Footer ──
# ============================================================

st.markdown("""
<div class="footer">
    Розроблено департаментом стратегічного планування та макроекономічного прогнозування<br>
    Версія DEMO 1.5 | 2026 | Внутрішня система моніторингу стратегічного плану
</div>
""", unsafe_allow_html=True)
