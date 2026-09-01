import os
import re
import hashlib
from io import BytesIO
from html import escape

import pandas as pd
import plotly.express as px
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.period_locks import apply_locked_status
from core.db import fetch_all
from core.deputies import DEPUTY_MINISTER_BY_SSP
from core.ui import render_readonly_table
from core import exports as core_exports
from core.exports import fig_png_bytes
import plotly.express as _px_rep
from core.page_setup import page_setup, render_footer
from core.access import is_super_admin_user
from core.strategic_data import load_strat_matrix as core_load_strat_matrix
from core import monitoring_data, operational, mio_shared, analytics_calculations
from core import periods as core_periods
from core.closeouts import append_confirmed_closeout_facts
from core.errors import log_exception, show_warning
from core.analytics_text import build_context as build_analytics_text_context, generate_analytics_note
from core.stage4 import build_approval_speed_analytics, build_return_analytics, kyiv_now


current_user = page_setup("Аналітика", page_name="Аналітика")
if not is_super_admin_user(current_user):
    st.error("У вас немає доступу до розділу «Аналітика».")
    st.stop()

st.markdown(
    """
<style>
header[data-testid="stHeader"] {
    background: transparent !important;
}

.stApp {
    background: #F7F9FC;
}
.main .block-container {
    max-width: 1550px;
    padding-top: 1.2rem;
    position: relative;
    z-index: 1;
}

.ua-line {
    height: 7px;
    border-radius: 999px;
    background: linear-gradient(90deg, #005BBB 0%, #005BBB 50%, #FFD500 50%, #FFD500 100%);
    margin-bottom: 14px;
}

.ministry-label {
    text-align: right;
    color: #61708A;
    font-size: 14px;
    font-weight: 700;
    margin-bottom: 8px;
}

.header-box,
.card,
.report-box,
.export-box,
.table-box {
    background: rgba(255,255,255,0.96);
    border: 1px solid #DCE4F0;
    border-radius: 16px;
    box-shadow: 0 8px 24px rgba(15,23,42,0.06);
}

.header-box {
    padding: 22px 26px;
    margin-bottom: 18px;
    backdrop-filter: blur(8px);
}

.header-title {
    font-size: 32px;
    font-weight: 950;
    color: #132238;
    margin-bottom: 8px;
}

.header-subtitle,
.card-subtitle {
    font-size: 15px;
    color: #61708A;
    line-height: 1.55;
}

.badge-wrap {
    display: flex;
    flex-wrap: wrap;
    gap: 8px;
    margin-top: 14px;
}

.badge {
    background: #EAF1FF;
    border: 1px solid #BFD3F2;
    color: #005BBB;
    border-radius: 999px;
    padding: 7px 11px;
    font-size: 13px;
    font-weight: 850;
}

.badge-green {
    background: #E4F5EC;
    border-color: #1E9E57;
    color: #0C713A;
}

.badge-yellow {
    background: #FDF3D8;
    border-color: #F4B400;
    color: #8A6400;
}

.card,
.export-box,
.table-box {
    padding: 22px 24px;
    margin-bottom: 18px;
}

.card-title,
.report-title {
    font-size: 21px;
    font-weight: 950;
    color: #132238;
    margin-bottom: 8px;
}

/* Filter fields inherit the single system template from assets/app.css, identical to Dashboard. */

.alert-grid {
    display: grid;
    grid-template-columns: repeat(4, minmax(0, 1fr));
    gap: 10px;
    margin: 10px 0 14px 0;
}

.alert-card {
    border-radius: 12px;
    padding: 11px 14px;
    border: 1px solid #DCE4F0;
    box-shadow: 0 8px 20px rgba(15,23,42,0.06);
}

.alert-title {
    font-size: 13px;
    color: #61708A;
    font-weight: 900;
    line-height: 1.35;
    min-height: 30px;
}

.alert-value {
    font-size: 24px;
    color: #132238;
    font-weight: 950;
    margin-top: 4px;
}

.alert-note {
    font-size: 12px;
    color: #61708A;
    margin-top: 3px;
    line-height: 1.25;
}

.alert-blue { background: #EAF1FF; border-color: #BFD3F2; }
.alert-green { background: #E4F5EC; border-color: #1E9E57; }
.alert-yellow { background: #FDF3D8; border-color: #F4B400; }
.alert-red { background: #FBE5E5; border-color: #DC4A4A; }

.mio-summary-box {
    background: #FFFFFF;
    border: 1px solid #DCE4F0;
    border-radius: 12px;
    padding: 12px 14px;
    margin: 4px 0 14px 0;
    box-shadow: 0 2px 10px rgba(15,23,42,0.04);
}
.mio-summary-title { font-size: 13px; font-weight: 900; color:#032A63; margin-bottom:8px; }
.mio-summary-grid { display:grid; grid-template-columns:repeat(5,minmax(0,1fr)); gap:8px; }
.mio-mini { background:#F7F9FC; border:1px solid #E3E9F2; border-radius:9px; padding:8px 10px; min-height:58px; }
.mio-mini span { display:block; color:#61708A; font-size:11px; font-weight:750; line-height:1.2; }
.mio-mini b { display:block; color:#132238; font-size:20px; margin-top:4px; }

.report-box {
    border-left: 7px solid #005BBB;
    padding: 24px 28px;
    margin: 18px 0;
}
.report-title { font-size:24px; margin-bottom:12px; }
.report-text { font-size:15px; line-height:1.75; color:#34445C; text-align:left; }
.report-meta { background:#F7F9FC; border:1px solid #DCE4F0; border-radius:12px; padding:12px 14px; color:#61708A; font-size:13px; font-weight:700; margin-bottom:12px; }

[data-testid="stMain"] div.stDownloadButton > button,
[data-testid="stMain"] div.stButton > button {
    border-radius: 12px;
    padding: 12px 18px;
    font-weight: 850;
}

.footer {
    text-align:center; color:#61708A; font-size:13px; margin-top:50px; padding:22px 0 12px 0; border-top:1px solid #DCE4F0;
}

@media (max-width:1100px) {
    .mio-summary-grid { grid-template-columns:repeat(2,minmax(0,1fr)); }
    .alert-grid { grid-template-columns:repeat(2,minmax(0,1fr)); }
}
</style>
""",
    unsafe_allow_html=True
)

QUARTERS = ["I", "II", "III", "IV"]
YEAR_OPTIONS = [2026, 2027, 2028]


def raw_value(value):
    try:
        if value is None or pd.isna(value) or str(value) == "None": return ""
    except (TypeError, ValueError):
        if value is None: return ""
    return str(value).strip()


def clean(value): return escape(raw_value(value))


def extract_ssp_index(value):
    match = re.search(r"\d+", raw_value(value)); return match.group(0) if match else ""


def get_goal_code(code):
    parts = raw_value(code).split("."); return parts[0] + "." if parts and parts[0] else ""


def get_task_code(code):
    parts = raw_value(code).split("."); return f"{parts[0]}.{parts[1]}." if len(parts) >= 2 else ""


def parse_period(value):
    text = raw_value(value).lower().replace("і", "i")
    if not text or text in {"nan", "none", "н.д.", "нд"}: return None
    q = None
    for token, n in (("iv квартал",4),("4 квартал",4),("iii квартал",3),("3 квартал",3),("ii квартал",2),("2 квартал",2),("i квартал",1),("1 квартал",1)):
        if token in text: q = n; break
    year = re.search(r"20\d{2}", text)
    return int(year.group()) * 10 + q if year and q else None


def format_pct(value):
    try:
        if value is None or pd.isna(value): return "—"
        number = float(value)
    except (TypeError, ValueError): return "—"
    return f"{int(number)}%" if number.is_integer() else f"{number:.1f}%"


def format_number_2(value):
    try:
        if value is None or pd.isna(value): return "—"
        number = float(value)
        return f"{number:.2f}".rstrip("0").rstrip(".")
    except (TypeError, ValueError): return raw_value(value) or "—"


def load_requests():
    return monitoring_data.measures_only(monitoring_data.load_monitoring_requests())


def load_workflow_logs():
    return pd.DataFrame(fetch_all("monitoring_logs", "*", order=("changed_at", False)))


def ensure_request_columns(requests_df):
    required = ["id","year","quarter","department","strat_code","status","numeric_value","risks","progress_text","approval_status","submitted_at","responsible_person","phone","email","file_names","file_urls","admin_comment","start_date","end_date"]
    out = requests_df.copy()
    for col in required:
        if col not in out.columns: out[col] = ""
    return out


def base_measures(strat_df):
    measures = strat_df[strat_df["object_type"] == "measure"].copy()
    goals = strat_df[strat_df["object_type"] == "goal"].copy(); tasks = strat_df[strat_df["object_type"] == "task"].copy()
    if measures.empty: return measures
    goal_names = goals.set_index("code")["name"].to_dict() if not goals.empty else {}
    task_names = tasks.set_index("code")["name"].to_dict() if not tasks.empty else {}
    parent_goal = measures.get("parent_goal_code", pd.Series("", index=measures.index)).fillna("").astype(str)
    parent_task = measures.get("parent_task_code", pd.Series("", index=measures.index)).fillna("").astype(str)
    measures["goal_code"] = parent_goal.where(parent_goal.str.strip().ne(""), measures["code"].apply(get_goal_code))
    measures["task_code"] = parent_task.where(parent_task.str.strip().ne(""), measures["code"].apply(get_task_code))
    measures["strategic_goal"] = measures["goal_code"].map(goal_names).fillna(measures.get("parent_goal_name", ""))
    measures["task_name"] = measures["task_code"].map(task_names).fillna(measures.get("parent_task_name", ""))
    measures["start_num"] = measures.get("start_period", pd.Series(index=measures.index)).apply(parse_period)
    measures["end_num"] = measures.get("end_period", pd.Series(index=measures.index)).apply(parse_period)
    measures["ssp_index"] = measures.get("resp_main", pd.Series("", index=measures.index)).apply(extract_ssp_index)
    measures["deputy_minister"] = measures["ssp_index"].map(DEPUTY_MINISTER_BY_SSP).fillna("")
    return measures


def filter_period_requests_to_active_cohort(requests, active, years, quarters):
    if requests is None or requests.empty: return pd.DataFrame()
    codes = set(active.get("code", pd.Series(dtype=object)).fillna("").astype(str).str.strip())
    data = requests[requests["year"].astype(str).isin(map(str, years)) & requests["quarter"].astype(str).isin(map(str, quarters))].copy()
    return data[data["strat_code"].fillna("").astype(str).str.strip().isin(codes)].copy()


def aggregate_status(active):
    if active.empty or "status" not in active.columns: return pd.DataFrame(columns=["status","Кількість"])
    return active.groupby("status", dropna=False).size().reset_index(name="Кількість").sort_values("Кількість", ascending=False)


def _excel_safe_value(value):
    """Convert structured/missing Python values to Excel-safe scalar values."""
    if isinstance(value, dict):
        return "; ".join(f"{key}: {item}" for key, item in value.items())
    if isinstance(value, (list, tuple, set)):
        return "; ".join(str(item) for item in value)
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except (TypeError, ValueError):
        pass
    return value


def _excel_safe_frame(df):
    if df is None:
        return pd.DataFrame()
    result = df.copy()
    for column in result.columns:
        result[column] = result[column].map(_excel_safe_value)
    return result


def create_excel_report(
    active,
    period_requests,
    goal_progress,
    dep_progress,
    task_progress,
    product_progress,
    status_counts,
    period_dynamics,
    metrics,
    filters,
):
    """Existing Analytics workbook, updated only for the new metric contract."""
    active_export = active.rename(columns={
        "report_year": "Рік",
        "report_quarter": "Квартал",
        "code": "Код заходу",
        "name": "Захід",
        "goal_code": "Код СЦ",
        "strategic_goal": "Стратегічна ціль",
        "task_code": "Код завдання",
        "task_name": "Завдання",
        "product_type": "Тип продукту",
        "department": "Самостійний структурний підрозділ",
        "deputy_minister": "Заступник Міністра",
        "indicator": "Індикатор",
        "unit": "Одиниця виміру",
        "selected_target": "Планове значення",
        "numeric_value": "Фактичне значення",
        "status": "Статус",
        "execution_score": "Виконання, %",
        "risk_level": "Рівень ризику",
        "progress_text": "Пояснення",
        "risks": "Ризики/відхилення",
    })
    active_cols = [
        "Рік", "Квартал", "Код заходу", "Захід", "Код СЦ", "Стратегічна ціль",
        "Код завдання", "Завдання", "Тип продукту", "Самостійний структурний підрозділ",
        "Заступник Міністра", "Індикатор", "Одиниця виміру", "Планове значення",
        "Фактичне значення", "Статус", "Виконання, %", "Рівень ризику", "Пояснення", "Ризики/відхилення",
    ]
    latest = metrics.get("latest_period")
    latest_label = f"{latest[1]} кв. {latest[0]}" if isinstance(latest, tuple) and len(latest) == 2 else "—"
    summary_df = pd.DataFrame([
        ["Період", f"Роки: {', '.join(map(str, filters['years']))}; квартали: {', '.join(filters['quarters'])}"],
        ["ССП", "; ".join(map(str, filters.get("ssp", []))) if filters.get("ssp") else "Усі"],
        ["Заступники Міністра", "; ".join(map(str, filters.get("deputies", []))) if filters.get("deputies") else "Усі"],
        ["Стратегічні цілі", "; ".join(map(str, filters.get("goal_labels", []))) if filters.get("goal_labels") else "Усі"],
        ["Завдання", "; ".join(map(str, filters.get("task_labels", []))) if filters.get("task_labels") else "Усі"],
        ["Типи продукту", "; ".join(map(str, filters.get("product_types", []))) if filters.get("product_types") else "Усі"],
        ["Дата формування", kyiv_now().strftime("%d.%m.%Y %H:%M")],
        ["Останній обраний період", latest_label],
        ["Унікальних заходів", metrics.get("unique_measures", 0)],
        ["Рівень виконання — останній обраний період", format_pct(metrics.get("completion"))],
        ["Покриття — середнє за вибраний діапазон", format_pct(metrics.get("coverage"))],
        ["Покриття — останній обраний період", format_pct(metrics.get("coverage_latest"))],
        [metrics.get("attention_label") or "Управлінська увага", metrics.get("attention_count", 0)],
    ], columns=["Показник", "Значення"])
    sheets = {
        "Пояснення": summary_df,
        "Підсумок": summary_df,
        "Аналітичний масив": active_export[[c for c in active_cols if c in active_export.columns]],
        "Стратегічні цілі": goal_progress,
        "Завдання": task_progress,
        "ССП": dep_progress,
        "Типи продукту": product_progress,
        "Динаміка": period_dynamics,
        "Статуси": status_counts,
        "Реєстр заявок": period_requests,
    }
    sheets = {name: _excel_safe_frame(frame) for name, frame in sheets.items()}
    output = BytesIO(core_exports.write_styled_excel(sheets, freeze_first_col=1))
    output.seek(0)
    return output


def build_report_charts(goal_progress, dep_progress, status_counts, period_dynamics):
    """Preserve the existing DOCX graphical materials using the updated Analytics columns."""
    charts = []
    brand = ["#005BBB", "#4D8DFF", "#BFD3F2", "#FFD500", "#FF7A45", "#DC4A4A"]

    def _style(fig, h=430):
        fig.update_layout(
            plot_bgcolor="white", paper_bgcolor="white",
            font=dict(family="Arial", size=13, color="#132238"),
            margin=dict(l=40, r=20, t=50, b=40), height=h,
        )
        return fig

    try:
        if period_dynamics is not None and not period_dynamics.empty:
            fig = _px_rep.line(
                period_dynamics, x="Період", y="Виконання", markers=True,
                color_discrete_sequence=[brand[0], brand[4]],
                title="Динаміка оціненого виконання, %",
            )
            fig.update_layout(legend_title_text="")
            png = fig_png_bytes(_style(fig), scale=2, width=1000, height=430)
            if png:
                charts.append(("Рис. Динаміка рівня виконання СП у розрізі звітних періодів", png))

        if status_counts is not None and not status_counts.empty:
            fig = _px_rep.pie(
                status_counts, names="status", values="Кількість", hole=0.45,
                color_discrete_sequence=brand,
                title="Розподіл заходів за статусами виконання",
            )
            png = fig_png_bytes(_style(fig), scale=2, width=900, height=430)
            if png:
                charts.append(("Рис. Структура портфеля заходів за статусами виконання", png))

        if goal_progress is not None and not goal_progress.empty and "Виконання" in goal_progress.columns:
            frame = goal_progress.sort_values("Виконання", ascending=True).copy()
            frame["Виконання"] = pd.to_numeric(frame["Виконання"], errors="coerce").round(2)
            fig = _px_rep.bar(
                frame, x="Виконання", y=frame["goal_code"].astype(str), orientation="h",
                color_discrete_sequence=[brand[0]], title="Рівень виконання за стратегічними цілями, %",
            )
            fig.update_layout(yaxis_title="СЦ", xaxis_title="Виконання, %")
            png = fig_png_bytes(_style(fig, h=max(360, 34 * len(frame) + 120)), scale=2, width=1000)
            if png:
                charts.append(("Рис. Порівняння рівня виконання у розрізі стратегічних цілей", png))

        if dep_progress is not None and not dep_progress.empty and "Виконання" in dep_progress.columns:
            frame = dep_progress.sort_values("Виконання", ascending=False).head(15).sort_values("Виконання", ascending=True)
            fig = _px_rep.bar(
                frame, x="Виконання", y=frame["ssp_index"].astype(str).apply(lambda v: f"ССП {v}"),
                orientation="h", color_discrete_sequence=[brand[1]],
                title="Рівень виконання за ССП (топ-15), %",
            )
            fig.update_layout(yaxis_title="", xaxis_title="Виконання, %")
            png = fig_png_bytes(_style(fig, h=max(360, 30 * len(frame) + 120)), scale=2, width=1000)
            if png:
                charts.append(("Рис. Рівень виконання у розрізі самостійних структурних підрозділів", png))
    except Exception as exc:
        show_warning(
            "Частину графіків аналітичного звіту не сформовано.",
            exc,
            "Підготовка графіків аналітичного звіту",
        )
    return charts


def create_docx_report(
    text,
    metrics,
    filters,
    goal_progress=None,
    dep_progress=None,
    product_progress=None,
    status_counts=None,
    period_dynamics=None,
    flex_note="",
):
    """Preserve the existing DOCX report sections; only replace obsolete metrics."""
    goal_progress = pd.DataFrame() if goal_progress is None else goal_progress
    dep_progress = pd.DataFrame() if dep_progress is None else dep_progress
    product_progress = pd.DataFrame() if product_progress is None else product_progress
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("АНАЛІТИЧНА ДОВІДКА")
    run.bold = True
    run.font.size = Pt(16)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("щодо стану виконання Стратегічного плану").font.size = Pt(12)
    scope = document.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.add_run(
        f"Роки: {', '.join(map(str, filters['years']))}; квартали: {', '.join(filters['quarters'])}"
    ).italic = True

    document.add_paragraph("Ключові показники").runs[0].bold = True
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Показник"
    table.rows[0].cells[1].text = "Значення"
    metric_rows = {
        "Унікальних заходів": metrics.get("unique_measures", 0),
        "Записів захід-період": metrics.get("total_rows", 0),
        "Рівень виконання — останній обраний період": format_pct(metrics.get("completion")),
        "Покриття за діапазон": format_pct(metrics.get("coverage")),
        "Покриття — останній період": format_pct(metrics.get("coverage_latest")),
        metrics.get("attention_label") or "Управлінська увага": metrics.get("attention_count", 0),
        "Без поточного подання": metrics.get("no_data", 0),
    }
    for key, value in metric_rows.items():
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    document.add_paragraph("Аналітичний висновок").runs[0].bold = True
    for paragraph in (text or "").split("\n\n"):
        if not paragraph.strip():
            continue
        p = document.add_paragraph(paragraph.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs:
            r.font.size = Pt(11)

    charts = []
    if any(frame is not None for frame in (goal_progress, dep_progress, status_counts, period_dynamics)):
        charts = build_report_charts(goal_progress, dep_progress, status_counts, period_dynamics)
    if charts:
        document.add_paragraph("Графічні матеріали").runs[0].bold = True
        intro = document.add_paragraph(
            "Наведені нижче рисунки ілюструють викладені вище висновки: динаміку оціненого виконання, "
            "структуру портфеля за статусами та порівняльний рівень виконання у розрізі стратегічних цілей і ССП."
        )
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for caption, png in charts:
            document.add_picture(BytesIO(png), width=Inches(6.3))
            cap = document.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap.runs:
                r.italic = True
                r.font.size = Pt(10)
    elif goal_progress is not None or dep_progress is not None or period_dynamics is not None:
        note = document.add_paragraph(
            "Графічні матеріали не додано: у середовищі відсутня можливість сформувати PNG-графіки."
        )
        note.italic = True

    document.add_paragraph("Додаткова структура даних").runs[0].bold = True
    document.add_paragraph(f"Стратегічні цілі в аналізі: {len(goal_progress)}.")
    document.add_paragraph(f"Самостійні структурні підрозділи в аналізі: {len(dep_progress)}.")
    document.add_paragraph(f"Типи продукту в аналізі: {len(product_progress)}.")
    document.add_paragraph("Сформовано автоматизованою системою моніторингу стратегічного плану.").italic = True
    document.add_paragraph(
        "Розроблено департаментом стратегічного планування та макроекономічного прогнозування."
    ).italic = True

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


st.markdown('<div class="ua-line"></div>', unsafe_allow_html=True)
st.markdown('<div class="ministry-label">🇺🇦 Міністерство економіки, довкілля та сільського господарства України</div>', unsafe_allow_html=True)
st.markdown('<div class="header-box"><div class="header-title">Аналітика</div><div class="header-subtitle">Автоматизований аналіз виконання Стратегічного плану, динаміки, структурних відхилень, результатів МіО та фінансової складової.</div></div>', unsafe_allow_html=True)

strat_df = core_load_strat_matrix()
_mio_requests_all = monitoring_data.load_monitoring_requests()
requests_df = ensure_request_columns(monitoring_data.measures_only(_mio_requests_all))
mio_requests_df = _mio_requests_all.copy() if isinstance(_mio_requests_all, pd.DataFrame) else pd.DataFrame()
workflow_logs = load_workflow_logs(); analytics_read_at = kyiv_now()
measures_all = base_measures(strat_df)
if measures_all.empty:
    st.warning("У стратегічній матриці не знайдено заходів для аналізу."); render_footer(); st.stop()

ssp_options_df = measures_all[["ssp_index", "department"]].dropna().drop_duplicates().sort_values("ssp_index")
ssp_options = [raw_value(x) for x in ssp_options_df["ssp_index"] if raw_value(x)]
ssp_labels = {raw_value(r["ssp_index"]): raw_value(r["department"]) for _, r in ssp_options_df.iterrows() if raw_value(r["ssp_index"])}
goal_rows = measures_all[["goal_code","strategic_goal"]].drop_duplicates().sort_values("goal_code")
goal_options = {raw_value(r["goal_code"]): f"{raw_value(r['goal_code'])} {raw_value(r['strategic_goal'])}" for _, r in goal_rows.iterrows() if raw_value(r["goal_code"])}
task_rows = measures_all[["task_code","task_name"]].drop_duplicates().sort_values("task_code")
task_options = {raw_value(r["task_code"]): f"{raw_value(r['task_code'])} {raw_value(r['task_name'])}" for _, r in task_rows.iterrows() if raw_value(r["task_code"])}
deputy_options = sorted(x for x in measures_all["deputy_minister"].dropna().astype(str).unique() if raw_value(x))
product_type_options = sorted(x for x in measures_all.get("product_type", pd.Series(dtype=object)).dropna().astype(str).unique() if raw_value(x))

_an_defaults = {"data_mode":operational.MODE_CONFIRMED,"years":[],"quarters":[],"ssp":[],"deputies":[],"goals":[],"tasks":[],"product_types":[]}
if "analytics_filters_applied_v20" not in st.session_state: st.session_state["analytics_filters_applied_v20"] = _an_defaults.copy()
_applied_seed = st.session_state["analytics_filters_applied_v20"]
for key, value in {
    "analytics_pending_data_mode":_applied_seed.get("data_mode",operational.MODE_CONFIRMED),
    "analytics_pending_years":list(_applied_seed.get("years",[])), "analytics_pending_quarters":list(_applied_seed.get("quarters",[])),
    "analytics_pending_ssp":list(_applied_seed.get("ssp",[])), "analytics_pending_deputies":list(_applied_seed.get("deputies",[])),
    "analytics_pending_goals":list(_applied_seed.get("goals",[])), "analytics_pending_tasks":list(_applied_seed.get("tasks",[])),
    "analytics_pending_products":list(_applied_seed.get("product_types",[])),
}.items():
    if key not in st.session_state: st.session_state[key] = value


def _apply_filters():
    st.session_state["analytics_filters_applied_v20"] = {
        "data_mode":st.session_state.get("analytics_pending_data_mode",operational.MODE_CONFIRMED),
        "years":list(st.session_state.get("analytics_pending_years",[]) or []), "quarters":list(st.session_state.get("analytics_pending_quarters",[]) or []),
        "ssp":list(st.session_state.get("analytics_pending_ssp",[]) or []), "deputies":list(st.session_state.get("analytics_pending_deputies",[]) or []),
        "goals":list(st.session_state.get("analytics_pending_goals",[]) or []), "tasks":list(st.session_state.get("analytics_pending_tasks",[]) or []),
        "product_types":list(st.session_state.get("analytics_pending_products",[]) or []),
    }


def _reset_filters():
    st.session_state["analytics_filters_applied_v20"] = _an_defaults.copy()
    st.session_state["analytics_pending_data_mode"] = operational.MODE_CONFIRMED
    for key in ("analytics_pending_years","analytics_pending_quarters","analytics_pending_ssp","analytics_pending_deputies","analytics_pending_goals","analytics_pending_tasks","analytics_pending_products"): st.session_state[key] = []

with st.form("analytics_filters_form_v20"):
    st.markdown('<div class="filter-title">Параметри відбору</div>', unsafe_allow_html=True)
    a0,a1,a2 = st.columns([1.45,.75,.9])
    with a0:
        st.markdown('<div class="filter-field-label">Джерело даних</div>', unsafe_allow_html=True); st.radio("Джерело даних", operational.MODE_OPTIONS, horizontal=True, key="analytics_pending_data_mode", label_visibility="collapsed")
    with a1:
        st.markdown('<div class="filter-field-label">Рік</div>', unsafe_allow_html=True); st.multiselect("Рік", YEAR_OPTIONS, key="analytics_pending_years", placeholder="Усі роки", label_visibility="collapsed")
    with a2:
        st.markdown('<div class="filter-field-label">Квартал</div>', unsafe_allow_html=True); st.multiselect("Квартал", QUARTERS, key="analytics_pending_quarters", placeholder="Усі квартали", label_visibility="collapsed")
    f1,f2,f3 = st.columns([1.2,1.15,1.3])
    with f1:
        st.markdown('<div class="filter-field-label">Самостійний структурний підрозділ</div>', unsafe_allow_html=True); st.multiselect("Самостійний структурний підрозділ", ssp_options, format_func=lambda x:ssp_labels.get(x,x), key="analytics_pending_ssp", placeholder="Усі підрозділи", label_visibility="collapsed")
    with f2:
        st.markdown('<div class="filter-field-label">Заступник Міністра</div>', unsafe_allow_html=True); st.multiselect("Заступник Міністра", deputy_options, key="analytics_pending_deputies", placeholder="Усі заступники", label_visibility="collapsed")
    with f3:
        st.markdown('<div class="filter-field-label">Стратегічна ціль</div>', unsafe_allow_html=True); st.multiselect("Стратегічна ціль", list(goal_options.values()), key="analytics_pending_goals", placeholder="Усі стратегічні цілі", label_visibility="collapsed")
    f4,f5 = st.columns([1.4,1])
    with f4:
        st.markdown('<div class="filter-field-label">Завдання</div>', unsafe_allow_html=True); st.multiselect("Завдання", list(task_options.values()), key="analytics_pending_tasks", placeholder="Усі завдання", label_visibility="collapsed")
    with f5:
        st.markdown('<div class="filter-field-label">Тип продукту</div>', unsafe_allow_html=True); st.multiselect("Тип продукту", product_type_options, key="analytics_pending_products", placeholder="Усі типи продукту", label_visibility="collapsed")
    b1,b2 = st.columns(2)
    with b1: st.form_submit_button("Застосувати обрані параметри", type="primary", use_container_width=True, on_click=_apply_filters)
    with b2: st.form_submit_button("Скинути параметри", use_container_width=True, on_click=_reset_filters)

applied = st.session_state.get("analytics_filters_applied_v20", _an_defaults.copy())
data_mode = applied.get("data_mode", operational.MODE_CONFIRMED)
selected_years = list(applied.get("years",[]) or []) or [2026]; selected_quarters = list(applied.get("quarters",[]) or []) or QUARTERS.copy()
selected_ssp = list(applied.get("ssp",[]) or []); selected_deputies = list(applied.get("deputies",[]) or [])
selected_goal_labels = list(applied.get("goals",[]) or []); selected_task_labels = list(applied.get("tasks",[]) or []); selected_products = list(applied.get("product_types",[]) or [])
selected_goals = [code for code,label in goal_options.items() if label in selected_goal_labels]; selected_tasks = [code for code,label in task_options.items() if label in selected_task_labels]

if not mio_requests_df.empty: mio_requests_df = apply_locked_status(mio_requests_df, status_col="status")
if data_mode == operational.MODE_OPERATIONAL:
    targets = operational.build_target_map(strat_df)
    if not requests_df.empty: requests_df, _ = operational.apply_operational_mode(requests_df, targets)
    if not mio_requests_df.empty:
        mio_requests_df, _ = operational.apply_operational_mode(mio_requests_df, targets); mio_requests_df = apply_locked_status(mio_requests_df, status_col="status")
requests_df = ensure_request_columns(append_confirmed_closeout_facts(requests_df))
mio_requests_df = append_confirmed_closeout_facts(mio_requests_df, include_incomplete=False)

base_results, _ = analytics_calculations.prepare_analysis_context(strat_df, requests_df, selected_years, selected_quarters)
ssp_base_results, period_results = analytics_calculations.build_analytics_result_context(base_results, selected_ssp, selected_deputies, selected_goals, selected_tasks, selected_products)
active = analytics_calculations.snapshot_rows_from_period_results(period_results)
if active.empty:
    st.warning("За обраними параметрами активних заходів не знайдено."); render_footer(); st.stop()
period_requests = filter_period_requests_to_active_cohort(requests_df, active, selected_years, selected_quarters)

plan = analytics_calculations.build_analytics_plan_summary(period_results)
metrics = analytics_calculations.build_metrics(active, period_results)
metrics.update({
    "completion":plan.get("execution_by_measures"), "goal_completion":plan.get("execution_by_goals"),
    "completion_change":plan.get("execution_by_measures_change"), "goal_completion_change":plan.get("execution_by_goals_change"),
    "coverage":plan.get("coverage_average"), "coverage_latest":plan.get("coverage_latest"), "coverage_change":plan.get("coverage_change"),
    "latest_period":plan.get("latest_period"), "latest_risk_summary":plan.get("latest_risk_summary") or {},
})
attention = plan.get("management_attention") or analytics_calculations.management_attention_info(period_results)
metrics.update({"attention_count":attention.get("count",0),"attention_type":attention.get("type"),"attention_label":attention.get("label"),"attention_assessed_count":attention.get("assessed_count",0)})

goal_progress = analytics_calculations.build_analytics_goal_summary(period_results, active)
task_progress = analytics_calculations.build_analytics_task_summary(period_results, active)
dep_progress = analytics_calculations.build_analytics_ssp_summary(period_results, active, base_results=ssp_base_results)
product_progress = analytics_calculations.aggregate_product_progress(period_results, active)
status_counts = aggregate_status(active); period_dynamics = analytics_calculations.build_analytics_dynamics(period_results)

# Reuse the existing MіO compatibility logic from the production Analytics page.
# Integral evaluation remains annual; content filters may narrow compatible MіO
# tables, while organisational/product filters never trigger a synthetic
# recalculation. Measure-level MіO and financing may follow the active cohort.
mio_goal_evaluation = pd.DataFrame()
mio_goal_task_evaluation = pd.DataFrame()
mio_measure_evaluation = pd.DataFrame()
mio_financing = pd.DataFrame()
mio_filter_limited = False
try:
    _mio_years = [int(y) for y in selected_years if int(y) in (2026, 2027, 2028)]
    _mio_outputs = mio_shared.build_mio_analytics(strat_df, mio_requests_df, _mio_years or [2026])
    _all_mio_goals = _mio_outputs.get("goals", pd.DataFrame()).copy()
    _all_mio_goal_tasks = _mio_outputs.get("goals_tasks", pd.DataFrame()).copy()
    _all_mio_measures = _mio_outputs.get("measures", pd.DataFrame()).copy()
    _all_mio_financing = _mio_outputs.get("financing", pd.DataFrame()).copy()

    _annual_scope = set(selected_quarters) == set(QUARTERS)
    _integral_compatible = _annual_scope and not any(
        (selected_ssp, selected_deputies, selected_tasks, selected_products)
    )
    _mio_indicator_compatible = _annual_scope and not any(
        (selected_ssp, selected_deputies, selected_products)
    )

    if _integral_compatible:
        mio_goal_evaluation = _all_mio_goals
        if selected_goals and not mio_goal_evaluation.empty and "Код" in mio_goal_evaluation.columns:
            mio_goal_evaluation = mio_goal_evaluation[
                mio_goal_evaluation["Код"].astype(str).isin(set(map(str, selected_goals)))
            ].copy()
    else:
        mio_filter_limited = not _all_mio_goals.empty

    if _mio_indicator_compatible and not _all_mio_goal_tasks.empty:
        mio_goal_task_evaluation = _all_mio_goal_tasks.copy()
        if selected_goals and "Код" in mio_goal_task_evaluation.columns:
            _goal_prefixes = tuple(str(x) for x in selected_goals)
            mio_goal_task_evaluation = mio_goal_task_evaluation[
                mio_goal_task_evaluation["Код"].astype(str).apply(
                    lambda x: any(x == g or x.startswith(g + ".") for g in _goal_prefixes)
                )
            ].copy()
        if selected_tasks and "Код" in mio_goal_task_evaluation.columns:
            _task_set = set(map(str, selected_tasks))
            _level = mio_goal_task_evaluation.get(
                "Рівень", pd.Series("", index=mio_goal_task_evaluation.index)
            ).astype(str)
            mio_goal_task_evaluation = mio_goal_task_evaluation[
                _level.eq("goal") | mio_goal_task_evaluation["Код"].astype(str).isin(_task_set)
            ].copy()

    _active_codes_for_mio = set(
        active.get("code", pd.Series(dtype=object)).astype(str).str.strip()
    )
    if _annual_scope and not _all_mio_measures.empty and "Захід" in _all_mio_measures.columns:
        mio_measure_evaluation = _all_mio_measures[
            _all_mio_measures["Захід"].astype(str).str.strip().isin(_active_codes_for_mio)
        ].copy()
    if not _all_mio_financing.empty and "Захід" in _all_mio_financing.columns:
        mio_financing = _all_mio_financing[
            _all_mio_financing["Захід"].astype(str).str.strip().isin(_active_codes_for_mio)
        ].copy()
except Exception as exc:
    log_exception("Analytics reusable MіO outputs", exc)
    mio_goal_evaluation = pd.DataFrame()
    mio_goal_task_evaluation = pd.DataFrame()
    mio_measure_evaluation = pd.DataFrame()
    mio_financing = pd.DataFrame()

selected_ssp_labels = [ssp_labels.get(x, x) for x in selected_ssp]
filters = {"years":selected_years,"quarters":selected_quarters,"ssp":selected_ssp_labels,"ssp_indices":selected_ssp,"deputies":selected_deputies,"goal_labels":selected_goal_labels,"task_labels":selected_task_labels,"product_types":selected_products,"data_mode":data_mode}
analytics_text_context = build_analytics_text_context(
    filters=filters, metrics=metrics, goal_progress=goal_progress, task_progress=task_progress,
    department_progress=dep_progress, product_progress=product_progress, status_counts=status_counts,
    period_dynamics=period_dynamics, active=active, mio_goal_evaluation=mio_goal_evaluation,
    mio_goal_task_evaluation=mio_goal_task_evaluation, mio_measure_evaluation=mio_measure_evaluation, mio_financing=mio_financing,
)

ANALYTICS_TEXT_DEBUG = str(os.getenv("ANALYTICS_TEXT_DEBUG","")).strip().lower() in {"1","true","yes","on"}
analytics_text_available=True; analytics_text_engine_incident=""
try:
    analytical_text = generate_analytics_note(context=analytics_text_context)
except Exception as exc:
    incident_seed=f"{type(exc).__name__}:{exc}".encode("utf-8",errors="replace"); analytics_text_engine_incident="AN-"+hashlib.sha256(incident_seed).hexdigest()[:10].upper()
    _validation_warnings=list(getattr(exc,"validation_warnings",()) or ())
    log_exception(
        "Analytics rule-based text generator", exc,
        incident_code=analytics_text_engine_incident,
        diagnostics={"validation_warnings":_validation_warnings,"filters":filters},
    )
    analytics_text_available=False; analytical_text=""
    if ANALYTICS_TEXT_DEBUG and _validation_warnings:
        st.code("\n".join(_validation_warnings), language="text")
    if ANALYTICS_TEXT_DEBUG:
        raise
    st.error(f"Аналітичну довідку не сформовано через технічну помилку. Код інциденту: {analytics_text_engine_incident}.")

st.caption(f"Режим даних: {data_mode} · станом на {analytics_read_at:%d.%m.%Y %H:%M}")
latest_period = metrics.get("latest_period"); latest_caption = f"{latest_period[1]} кв. {latest_period[0]}" if latest_period else "—"
st.markdown(
    f'''<div class="alert-grid">
<div class="alert-card"><div class="alert-title">Рівень виконання Стратегічного плану · {latest_caption}</div><div class="alert-value">{format_pct(metrics.get('completion'))}</div><div class="alert-note">Стан саме останнього обраного періоду</div></div>
<div class="alert-card alert-red"><div class="alert-title">{clean(metrics.get('attention_label') or 'Потребують управлінської уваги')}</div><div class="alert-value">{int(metrics.get('attention_count') or 0)}</div><div class="alert-note">Унікальні заходи в актуальному зрізі</div></div>
<div class="alert-card alert-yellow"><div class="alert-title">Покриття моніторингом</div><div class="alert-value">{format_pct(metrics.get('coverage'))}</div><div class="alert-note">Середнє за діапазон · останній період {format_pct(metrics.get('coverage_latest'))}</div></div>
<div class="alert-card alert-green"><div class="alert-title">Заходів у вибірці</div><div class="alert-value">{metrics.get('unique_measures',0)}</div><div class="alert-note">Унікальні коди за вибраним діапазоном</div></div>
</div>''', unsafe_allow_html=True,
)

if not mio_goal_evaluation.empty:
    mio_year=max([int(y) for y in selected_years if int(y) in (2026,2027,2028)] or [2026]); summary=mio_shared.summarize_integral_goals(mio_goal_evaluation,mio_year)
    st.markdown(f'''<div class="mio-summary-box"><div class="mio-summary-title">Оцінка МіО · {mio_year}</div><div class="mio-summary-grid">
<div class="mio-mini"><span>Інтегральна оцінка</span><b>{format_pct(summary.average_integral)}</b></div>
<div class="mio-mini"><span>Виконання заходів</span><b>{format_pct(summary.average_measure_execution)}</b></div>
<div class="mio-mini"><span>Оцінка завдань</span><b>{format_pct(summary.average_task_score)}</b></div>
<div class="mio-mini"><span>Прогрес індикаторів цілей</span><b>{format_pct(summary.average_strategic_progress)}</b></div>
<div class="mio-mini"><span>Фінансове виконання</span><b>{format_pct(analytics_text_context.factual_value('mio.fin.avg_financial_execution'))}</b></div>
</div></div>''', unsafe_allow_html=True)
elif mio_filter_limited:
    st.caption("Оцінка МіО не перераховується за організаційним, продуктовим або неповним квартальним зрізом; методологія МіО не змінюється.")

if analytics_text_available:
    st.markdown('<div class="report-box"><div class="report-title">Автоматично сформована аналітична довідка</div>', unsafe_allow_html=True)
    for paragraph in analytical_text.split("\n\n"): st.markdown(f"<p class='report-text'>{clean(paragraph)}</p>", unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="card"><div class="card-title">Графіки до довідки</div>', unsafe_allow_html=True)
g1,g2=st.columns(2)
with g1:
    if not period_dynamics.empty:
        st.plotly_chart(px.line(period_dynamics,x="Період",y="Виконання",markers=True,title="Динаміка оціненого виконання",labels={"Виконання":"Виконання, %"}),use_container_width=True)
with g2:
    if not goal_progress.empty:
        hover=[c for c in ["strategic_goal","Унікальних_заходів","Покриття_середнє_%","Покриття_останній_%","Актуальна_увага","Без_даних","Тип_уваги"] if c in goal_progress.columns]
        fig=px.bar(goal_progress.sort_values("goal_code"),x="goal_code",y="Виконання",text="Виконання",hover_data=hover,title="Виконання за стратегічними цілями",labels={"goal_code":"Стратегічна ціль","Виконання":"Виконання, %"}); st.plotly_chart(fig,use_container_width=True)
g3,g4=st.columns(2)
with g3:
    if not dep_progress.empty:
        chart=dep_progress.copy(); chart["ССП"]=chart["ssp_index"].astype(str)+" — "+chart.get("department",pd.Series("",index=chart.index)).astype(str)
        hover=[c for c in ["deputy_minister","Унікальних_заходів","Покриття_середнє_%","Покриття_останній_%","Актуальна_увага","Без_даних","risk_high_critical_latest"] if c in chart.columns]
        fig=px.bar(chart,x="ССП",y="Виконання",text="Виконання",hover_data=hover,title="Виконання за самостійними структурними підрозділами"); fig.update_layout(xaxis_tickangle=-35); st.plotly_chart(fig,use_container_width=True)
with g4:
    if not product_progress.empty:
        hover=[c for c in ["Виконання","Покриття_середнє_%","Покриття_останній_%","Актуальна_увага","Без_даних"] if c in product_progress.columns]
        fig=px.bar(product_progress,x="product_type",y="Унікальних_заходів",text="Унікальних_заходів",hover_data=hover,title="Структура заходів за типами продукту"); fig.update_layout(xaxis_tickangle=-25); st.plotly_chart(fig,use_container_width=True)
g5,g6=st.columns(2)
with g5:
    if not status_counts.empty: st.plotly_chart(px.pie(status_counts,names="status",values="Кількість",hole=.45,title="Структура статусів виконання"),use_container_width=True)
with g6:
    if not task_progress.empty and "Актуальна_увага" in task_progress.columns:
        top=task_progress.sort_values(["Актуальна_увага","Без_даних"],ascending=False).head(10).copy(); top["Завдання"]=top["task_code"].astype(str)
        hover=[c for c in ["task_name","Виконання","Покриття_останній_%","Без_даних","Тип_уваги"] if c in top.columns]
        st.plotly_chart(px.bar(top,x="Завдання",y="Актуальна_увага",text="Актуальна_увага",hover_data=hover,title="Завдання з найбільшою кількістю актуальних сигналів управлінської уваги",labels={"Актуальна_увага":"Унікальних заходів"}),use_container_width=True)
st.markdown('</div>', unsafe_allow_html=True)

# Workflow analytics remains unchanged in methodology and uses the same filtered request cohort.
active_codes=set(active.get("code",pd.Series(dtype=object)).fillna("").astype(str).str.strip()); workflow_requests=period_requests.copy()
if not workflow_requests.empty and active_codes: workflow_requests=workflow_requests[workflow_requests["strat_code"].fillna("").astype(str).str.strip().isin(active_codes)].copy()
return_analytics=build_return_analytics(workflow_logs,workflow_requests); approval_speed=build_approval_speed_analytics(workflow_logs,workflow_requests,now=analytics_read_at)
st.markdown('<div class="card"><div class="card-title">Аналіз повернень на доопрацювання <span style="font-size:11px;color:#8A6400">тест</span></div></div>',unsafe_allow_html=True)
w1,w2=st.columns(2); w1.metric("Кількість повернень",return_analytics["total_returns"]); w2.metric("Середня кількість повернень на одну заявку",return_analytics["average_per_request"])
if not return_analytics["by_department"].empty: st.plotly_chart(px.bar(return_analytics["by_department"],x="ССП",y="Кількість повернень",title="Рейтинг ССП за кількістю повернень"),use_container_width=True)
if not return_analytics["by_stage"].empty: st.plotly_chart(px.bar(return_analytics["by_stage"],x="Ланка, що повернула",y="Кількість повернень",title="Розподіл за ланками, які повертають"),use_container_width=True)
if not return_analytics["top_requests"].empty: render_readonly_table(return_analytics["top_requests"].head(20),visual_style="signal",variant="ranking")

st.markdown(
    '<div class="card"><div class="card-title">Швидкість погодження '
    '<span style="font-size:11px;color:#8A6400;background:#FDF3D8;border:1px solid #F4B400;'
    'border-radius:999px;padding:3px 8px;">тест</span></div>'
    '<div class="card-subtitle">Час розраховано за датою подання і послідовністю подій у журналі.</div>',
    unsafe_allow_html=True,
)
_speed_c1, _speed_c2 = st.columns(2)
with _speed_c1:
    st.metric(
        "Середній час від подання до фінального погодження",
        f"{approval_speed['average_total_days']} дн.",
    )
with _speed_c2:
    st.metric("Фінально погоджених заявок у розрахунку", approval_speed["completed_requests"])

_speed_left, _speed_right = st.columns(2)
with _speed_left:
    st.markdown("**Середній час на кожній ланці**")
    if approval_speed["stage_average"].empty:
        st.info("Для розрахунку часу на ланках недостатньо завершених переходів.")
    else:
        render_readonly_table(approval_speed["stage_average"], visual_style="signal", variant="analytics")
with _speed_right:
    st.markdown("**Заявки, що зараз очікують найдовше**")
    if approval_speed["hanging"].empty:
        st.info("У поточному зрізі немає заявок на активних ланках погодження.")
    else:
        render_readonly_table(approval_speed["hanging"].head(20), visual_style="signal", variant="analytics")
st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="export-box"><div class="card-title">Експорт</div>',unsafe_allow_html=True)
e1,e2,e3=st.columns(3)
with e1:
    excel=create_excel_report(active,period_requests,goal_progress,dep_progress,task_progress,product_progress,status_counts,period_dynamics,metrics,filters); excel.seek(0)
    st.download_button("Завантажити Excel-звіт",data=excel,file_name=f"analytics_report_{'_'.join(map(str,selected_years))}_{'_'.join(selected_quarters)}.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",use_container_width=True)
with e2:
    if analytics_text_available:
        docx=create_docx_report(analytical_text,metrics,filters,goal_progress,dep_progress,product_progress,status_counts=status_counts,period_dynamics=period_dynamics,flex_note=""); st.download_button("Аналітична довідка DOCX",data=docx,file_name=f"analytical_note_{'_'.join(map(str,selected_years))}_{'_'.join(selected_quarters)}.docx",mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",use_container_width=True)
with e3:
    registry=BytesIO(core_exports.write_styled_excel({"Реєстр заявок":period_requests})); registry.seek(0); st.download_button("Реєстр заявок",data=registry,file_name=f"requests_registry_{'_'.join(map(str,selected_years))}_{'_'.join(selected_quarters)}.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",use_container_width=True)
st.markdown('</div>',unsafe_allow_html=True)

st.markdown('<div class="table-box"><div class="card-title">Таблиці для перевірки</div>',unsafe_allow_html=True)
tab1,tab2,tab3,tab4,tab5,tab6=st.tabs(["Аналітичний масив","Стратегічні цілі","Завдання","ССП","Типи продукту","Реєстр заявок"])
with tab1:
    cols=["report_year","report_quarter","code","name","strategic_goal","task_name","product_type","department","deputy_minister","status","execution_score","risk_level","missing_required_submission"]
    render_readonly_table(active[[c for c in cols if c in active.columns]],visual_style="signal",variant="wide")
with tab2:
    show=goal_progress.rename(columns={"goal_code":"Код стратегічної цілі","strategic_goal":"Стратегічна ціль","Виконання":"Виконання, %","Зміна":"Зміна, в.п.","Покриття_середнє_%":"Середнє покриття, %","Покриття_останній_%":"Покриття останнього періоду, %","Актуальна_увага":"Актуальна управлінська увага","Без_даних":"Без поточного подання","Заходів_періодів":"Записів захід-період","Унікальних_заходів":"Унікальних заходів"}); render_readonly_table(show,visual_style="signal",variant="analytics")
with tab3:
    show=task_progress.rename(columns={"task_code":"Код завдання","task_name":"Завдання","Виконання":"Виконання, %","Зміна":"Зміна, в.п.","Покриття_середнє_%":"Середнє покриття, %","Покриття_останній_%":"Покриття останнього періоду, %","Актуальна_увага":"Актуальна управлінська увага","Без_даних":"Без поточного подання"}); render_readonly_table(show,visual_style="signal",variant="analytics")
with tab4:
    show=dep_progress.rename(columns={"ssp_index":"ССП","Виконання":"Виконання, %","Зміна":"Зміна, в.п.","Покриття_середнє_%":"Середнє покриття, %","Покриття_останній_%":"Покриття останнього періоду, %","Актуальна_увага":"Актуальна управлінська увага","Без_даних":"Без поточного подання"}); render_readonly_table(show,visual_style="signal",variant="ranking")
with tab5:
    show=product_progress.rename(columns={"product_type":"Тип продукту","Виконання":"Виконання, %","Покриття_середнє_%":"Середнє покриття, %","Покриття_останній_%":"Покриття останнього періоду, %","Актуальна_увага":"Актуальна управлінська увага","Без_даних":"Без поточного подання"}); render_readonly_table(show,visual_style="signal",variant="analytics")
with tab6:
    render_readonly_table(period_requests,visual_style="signal",variant="wide")
st.markdown('</div>',unsafe_allow_html=True)

render_footer()
