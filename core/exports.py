# core/exports.py

"""
Спільні інструменти вивантаження (правка №15).

- fig_png_bytes / render_png_download — PNG будь-якого plotly-графіка
  (роздільна «для друку», scale=2) через kaleido;
- register_cyrillic_font / build_presentation_pdf — PDF презентаційного
  режиму Dashboard у фірмовому стилі (reportlab + PNG графіків);
- усі функції деградують коректно: якщо kaleido/reportlab недоступні,
  показується зрозуміла підказка, а сторінка працює далі.
"""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path

from core.timeutils import now_kyiv
from core.errors import log_cosmetic_error, show_warning


# ------------------------------------------------------------
# PNG для plotly-графіків
# ------------------------------------------------------------

def fig_png_bytes(fig, scale: int = 2, width: int | None = None,
                  height: int | None = None) -> bytes | None:
    """PNG-байти plotly-фігури через kaleido; None — якщо kaleido недоступний."""
    try:
        kwargs = {"format": "png", "scale": scale}
        if width:
            kwargs["width"] = width
        if height:
            kwargs["height"] = height
        return fig.to_image(**kwargs)
    except Exception:
        return None


def render_png_download(fig, filename: str, key: str,
                        label: str = "🖼 Завантажити графік (PNG)") -> None:
    """Кнопка завантаження PNG під графіком. Не ламає сторінку без kaleido."""
    import streamlit as st
    png = fig_png_bytes(fig)
    if png is None:
        st.caption(
            "PNG-експорт недоступний: додайте пакет `kaleido` у requirements.txt "
            "та перезапустіть застосунок."
        )
        return
    st.download_button(
        label, data=png,
        file_name=f"{filename}.png", mime="image/png",
        key=key, use_container_width=False,
    )


# ------------------------------------------------------------
# PDF презентаційного режиму
# ------------------------------------------------------------

_CYRILLIC_FONT_CANDIDATES = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
]

_PRES_BG = "#032A63"
_PRES_WHITE = "#FFFFFF"
_PRES_BLUE = "#005BBB"
_PRES_YELLOW = "#FFD500"
_PRES_METRIC_BLUE = "#4D8DFF"
_PRES_TEAL = "#00A8A8"
_PRES_ORANGE = "#FF7A45"
_PRES_AMBER = "#F4B400"
_PRES_GREY = "#8A96A8"
_PRES_RED = "#DC4A4A"
_PRES_GREEN = "#118847"
_PRES_BRIGHT_GREEN = "#1E9E57"

# Legacy generic renderer colors retained for the existing MIO consumer.
BRAND_BLUE = (0 / 255, 91 / 255, 187 / 255)
BRAND_YELLOW = (255 / 255, 213 / 255, 0 / 255)
DARK = (15 / 255, 23 / 255, 42 / 255)
GREY = (71 / 255, 85 / 255, 105 / 255)


def _register_fonts():
    """Register bundled/system DejaVu fonts so Ukrainian text is deterministic."""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        candidates = [
            (Path("assets/fonts/DejaVuSans.ttf"), Path("assets/fonts/DejaVuSans-Bold.ttf")),
            (Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
             Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")),
        ]
        for regular_path, bold_path in candidates:
            if not regular_path.exists():
                continue
            pdfmetrics.registerFont(TTFont("DejaVu", str(regular_path)))
            if bold_path.exists():
                pdfmetrics.registerFont(TTFont("DejaVu-Bold", str(bold_path)))
                return "DejaVu", "DejaVu-Bold"
            return "DejaVu", "DejaVu"
        return None
    except Exception:
        return None


def _register_presentation_fonts():
    """Use the same Linux fontconfig fallback as browser Helvetica Neue/Arial.

    Dashboard HTML declares Helvetica Neue, Arial, sans-serif. Chromium resolves
    that stack through fontconfig on Linux. Ask fontconfig for the same physical
    regular/bold font files so ReportLab uses matching metrics. Legacy MIO keeps
    the existing DejaVu resolver.
    """
    try:
        import subprocess
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        def fc_file(pattern):
            try:
                result = subprocess.run(
                    ["fc-match", "-f", "%{file}", pattern],
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=2,
                )
                path = Path(result.stdout.strip())
                return path if path.exists() else None
            except Exception:
                return None

        dynamic_pairs = [
            (fc_file("Helvetica Neue"), fc_file("Helvetica Neue:style=Bold")),
            (fc_file("Arial"), fc_file("Arial:style=Bold")),
        ]
        static_pairs = [
            (
                Path("/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
                Path("/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf"),
            ),
            (
                Path("/usr/share/fonts/truetype/croscore/Arimo-Regular.ttf"),
                Path("/usr/share/fonts/truetype/croscore/Arimo-Bold.ttf"),
            ),
            (
                Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
                Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
            ),
            (
                Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
                Path("/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"),
            ),
        ]
        for regular_path, bold_path in [*dynamic_pairs, *static_pairs]:
            if regular_path is None or not regular_path.exists():
                continue
            pdfmetrics.registerFont(TTFont("PresentationSans", str(regular_path)))
            if bold_path is not None and bold_path.exists():
                pdfmetrics.registerFont(TTFont("PresentationSans-Bold", str(bold_path)))
                return "PresentationSans", "PresentationSans-Bold"
            pdfmetrics.registerFont(TTFont("PresentationSans-Bold", str(regular_path)))
            return "PresentationSans", "PresentationSans-Bold"
    except Exception:
        pass
    return _register_fonts()


def build_legacy_presentation_pdf(
    title: str,
    period_text: str,
    kpi_items: list[tuple[str, str]],
    verdict_text: str,
    verdict_level: str,
    insight_lines: list[str],
    figures: list[tuple[str, object]],
    logo_path: str = "assets/Мінекономіки.png",
) -> bytes | None:
    """
    Legacy generic PDF renderer retained only for non-Dashboard consumers:
    титульний слайд → слайд KPI → по слайду на кожен графік → висновок.

    figures: [(назва слайда, plotly_fig), ...]
    verdict_level: "high" | "medium" | "low" (колір плашки висновку).
    Повертає bytes або None (якщо reportlab недоступний).
    """
    try:
        from reportlab.lib.pagesizes import landscape
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas as pdf_canvas
    except Exception:
        return None

    fonts = _register_fonts()
    if fonts is None:
        font_r, font_b = "Helvetica", "Helvetica-Bold"  # запасний варіант
    else:
        font_r, font_b = fonts

    page_w, page_h = 338.7 * mm / 1.27, 190.5 * mm / 1.27  # ~16:9
    page_size = (960, 540)
    page_w, page_h = page_size

    buffer = io.BytesIO()
    c = pdf_canvas.Canvas(buffer, pagesize=page_size)

    def slide_bg():
        c.setFillColorRGB(0.97, 0.98, 0.99)
        c.rect(0, 0, page_w, page_h, fill=1, stroke=0)
        c.setFillColorRGB(*BRAND_BLUE)
        c.rect(0, page_h - 8, page_w, 8, fill=1, stroke=0)
        c.setFillColorRGB(*BRAND_YELLOW)
        c.rect(0, page_h - 12, page_w, 4, fill=1, stroke=0)

    def footer(page_no: int):
        c.setFont(font_r, 9)
        c.setFillColorRGB(*GREY)
        c.drawString(30, 16, "Мінекономіки · Система моніторингу стратегічного плану")
        c.drawRightString(page_w - 30, 16, f"Слайд {page_no}")

    def slide_title(text: str):
        c.setFont(font_b, 24)
        c.setFillColorRGB(*DARK)
        c.drawString(40, page_h - 58, text[:80])

    page_no = 1

    # ── Титульний слайд ──
    slide_bg()
    try:
        if Path(logo_path).exists():
            c.drawImage(logo_path, 40, page_h - 170, width=420,
                        preserveAspectRatio=True, anchor="nw", mask="auto")
    except Exception as exc:
        log_cosmetic_error("Додавання логотипа до PDF-презентації", exc)
    c.setFont(font_b, 34)
    c.setFillColorRGB(*DARK)
    c.drawString(40, page_h / 2 + 10, title[:60])
    c.setFont(font_r, 18)
    c.setFillColorRGB(*GREY)
    c.drawString(40, page_h / 2 - 26, period_text[:90])
    c.setFont(font_r, 12)
    c.drawString(40, 60, f"Сформовано: {now_kyiv().strftime('%d.%m.%Y %H:%M')}")
    footer(page_no)
    c.showPage()

    # ── Слайд KPI ──
    page_no += 1
    slide_bg()
    slide_title("Ключові показники")
    cols, card_w, card_h, gap = 4, 205, 120, 18
    x0, y0 = 40, page_h - 110 - card_h
    for i, (label, value) in enumerate(kpi_items[:8]):
        row, col = divmod(i, cols)
        x = x0 + col * (card_w + gap)
        y = y0 - row * (card_h + gap)
        c.setFillColorRGB(1, 1, 1)
        c.roundRect(x, y, card_w, card_h, 12, fill=1, stroke=0)
        c.setFillColorRGB(*BRAND_BLUE)
        c.rect(x, y + card_h - 6, card_w, 6, fill=1, stroke=0)
        c.setFont(font_b, 30)
        c.setFillColorRGB(*DARK)
        c.drawString(x + 16, y + card_h - 58, str(value)[:12])
        c.setFont(font_r, 11)
        c.setFillColorRGB(*GREY)
        # перенос підпису на 2 рядки
        words, line1, line2 = str(label).split(), "", ""
        for w in words:
            if len(line1) + len(w) < 30:
                line1 += (" " if line1 else "") + w
            else:
                line2 += (" " if line2 else "") + w
        c.drawString(x + 16, y + 34, line1[:34])
        if line2:
            c.drawString(x + 16, y + 20, line2[:34])
    footer(page_no)
    c.showPage()

    # ── Слайди-графіки ──
    for fig_title, fig in figures:
        png = fig_png_bytes(fig, scale=2, width=1100, height=560)
        if png is None:
            continue
        page_no += 1
        slide_bg()
        slide_title(fig_title)
        try:
            from reportlab.lib.utils import ImageReader
            img = ImageReader(io.BytesIO(png))
            c.drawImage(img, 40, 50, width=page_w - 80, height=page_h - 140,
                        preserveAspectRatio=True, anchor="c")
        except Exception as exc:
            log_cosmetic_error("Додавання графіка до PDF-презентації", exc)
        footer(page_no)
        c.showPage()

    # ── Слайд висновку ──
    page_no += 1
    slide_bg()
    slide_title("Висновок")
    verdict_colors = {
        "high": (220 / 255, 38 / 255, 38 / 255),
        "medium": (180 / 255, 83 / 255, 9 / 255),
        "low": (21 / 255, 128 / 255, 61 / 255),
    }
    c.setFillColorRGB(*verdict_colors.get(verdict_level, GREY))
    c.roundRect(40, page_h - 150, page_w - 80, 56, 12, fill=1, stroke=0)
    c.setFont(font_b, 16)
    c.setFillColorRGB(1, 1, 1)
    c.drawString(58, page_h - 128, verdict_text[:110])

    c.setFont(font_r, 13)
    c.setFillColorRGB(*DARK)
    y = page_h - 190
    for line in insight_lines[:9]:
        c.drawString(48, y, ("• " + str(line))[:130])
        y -= 24
    footer(page_no)
    c.showPage()

    c.save()
    return buffer.getvalue()


def build_presentation_pdf(presentation_payload: dict) -> bytes | None:
    """Render the canonical seven-slide Dashboard Presentation payload.

    Geometry follows the production CSS box model at the deterministic
    1366x768 reference viewport. This function does not recalculate analytics,
    risk, finance, source, status, filter, Top-5, or KPI semantics.
    """
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfgen import canvas as pdf_canvas
    except Exception:
        return None

    from core.presentation import (
        REFERENCE_HEIGHT,
        REFERENCE_WIDTH,
        validate_presentation_payload,
    )

    validate_presentation_payload(presentation_payload)
    fonts = _register_presentation_fonts()
    font_r, font_b = fonts if fonts is not None else ("Helvetica", "Helvetica-Bold")

    page_w = float(REFERENCE_WIDTH)
    page_h = float(REFERENCE_HEIGHT)
    content_w = page_w - 128.0
    buffer = io.BytesIO()
    c = pdf_canvas.Canvas(buffer, pagesize=(page_w, page_h))

    TITLE_SIZE = min(max(page_w * 0.04, 32.0), 56.0)
    TITLE_LH = TITLE_SIZE * 1.1
    TITLE_SUB_SIZE = min(max(page_w * 0.014, 14.0), 18.0)
    TITLE_SUB_LH = TITLE_SUB_SIZE * 1.6
    H2_SIZE = min(max(page_w * 0.028, 24.0), 38.0)
    H2_LH = H2_SIZE * 1.15
    HSUB_SIZE = min(max(page_w * 0.011, 12.0), 15.0)
    HSUB_LH = HSUB_SIZE * 1.2
    KPI_VALUE_SIZE = min(max(page_w * 0.04, 36.0), 56.0)
    RISK_VALUE_SIZE = min(max(page_w * 0.05, 40.0), 64.0)
    VERDICT_BADGE_SIZE = min(max(page_w * 0.02, 18.0), 26.0)
    VERDICT_BODY_SIZE = min(max(page_w * 0.012, 13.0), 16.0)

    EMOJI_TOKENS = (
        "🇺🇦", "🔴", "🟡", "🟢", "ℹ️", "📅", "🗓", "🏢", "📌", "🕐",
        "📋", "📊", "🎯", "⬆", "⛶", "✕",
    )

    def rgb(value):
        value = str(value or "#000000").lstrip("#")
        return tuple(int(value[i:i + 2], 16) / 255.0 for i in (0, 2, 4))

    def blend(fg, alpha, bg=_PRES_BG):
        fr, fg_, fb = rgb(fg)
        br, bg_, bb = rgb(bg)
        return (
            br * (1 - alpha) + fr * alpha,
            bg_ * (1 - alpha) + fg_ * alpha,
            bb * (1 - alpha) + fb * alpha,
        )

    def set_fill(value):
        c.setFillColorRGB(*rgb(value))

    def set_stroke(value):
        c.setStrokeColorRGB(*rgb(value))

    def plain(value):
        text = str(value if value is not None else "")
        for token in EMOJI_TOKENS:
            text = text.replace(token, "")
        return text.replace("\ufe0f", "").strip()

    def leading_marker(value):
        text = str(value if value is not None else "").strip()
        for token in EMOJI_TOKENS:
            if text.startswith(token):
                return token, text[len(token):].replace("\ufe0f", "").strip()
        return "", plain(text)

    def text_width(text, font_name, font_size, char_space=0.0):
        value = str(text or "")
        width = pdfmetrics.stringWidth(value, font_name, font_size)
        if char_space and len(value) > 1:
            width += char_space * (len(value) - 1)
        return width

    def wrap_lines(text, font_name, font_size, max_width, char_space=0.0):
        paragraphs = plain(text).splitlines() or [""]
        lines = []
        for paragraph in paragraphs:
            words = paragraph.split()
            if not words:
                lines.append("")
                continue
            current = ""
            for word in words:
                candidate = word if not current else f"{current} {word}"
                if text_width(candidate, font_name, font_size, char_space) <= max_width:
                    current = candidate
                    continue
                if current:
                    lines.append(current)
                    current = ""
                if text_width(word, font_name, font_size, char_space) <= max_width:
                    current = word
                    continue
                part = ""
                for ch in word:
                    candidate = part + ch
                    if text_width(candidate, font_name, font_size, char_space) <= max_width:
                        part = candidate
                    else:
                        if part:
                            lines.append(part)
                        part = ch
                current = part
            if current:
                lines.append(current)
        return lines or [""]

    def ellipsize(text, font_name, font_size, max_width, char_space=0.0):
        value = plain(text)
        if text_width(value, font_name, font_size, char_space) <= max_width:
            return value
        value = value.rstrip("… ")
        while value and text_width(value + "…", font_name, font_size, char_space) > max_width:
            value = value[:-1]
        return value.rstrip() + "…"

    def baseline_for_top(top, font_name, font_size, line_height):
        ascent = pdfmetrics.getAscent(font_name) * font_size / 1000.0
        return page_h - (top + (line_height - font_size) / 2.0 + ascent)

    def draw_line(
        text, x, top, *, font=font_r, size=13.0, line_height=None,
        color=_PRES_WHITE, char_space=0.0, align="left", width=None,
    ):
        line_height = float(line_height or size * 1.2)
        value = plain(text)
        c.setFillColorRGB(*rgb(color) if isinstance(color, str) else color)
        text_obj = c.beginText()
        text_obj.setTextOrigin(x, baseline_for_top(top, font, size, line_height))
        text_obj.setFont(font, size)
        if char_space:
            text_obj.setCharSpace(char_space)
        if align == "right" and width is not None:
            text_obj.setTextOrigin(
                x + width - text_width(value, font, size, char_space),
                baseline_for_top(top, font, size, line_height),
            )
        text_obj.textOut(value)
        c.drawText(text_obj)

    def draw_text_block(
        text, x, top, max_width, *, font=font_r, size=13.0,
        line_height=None, color=_PRES_WHITE, char_space=0.0,
        max_lines=None, ellipsis=False,
    ):
        line_height = float(line_height or size * 1.2)
        lines = wrap_lines(text, font, size, max_width, char_space)
        if max_lines is not None and len(lines) > max_lines:
            lines = lines[:max_lines]
            if ellipsis and lines:
                lines[-1] = ellipsize(lines[-1], font, size, max_width, char_space)
        for idx, line in enumerate(lines):
            draw_line(
                line, x, top + idx * line_height,
                font=font, size=size, line_height=line_height,
                color=color, char_space=char_space,
            )
        return len(lines) * line_height

    def text_block_height(text, font, size, line_height, max_width, char_space=0.0):
        return len(wrap_lines(text, font, size, max_width, char_space)) * line_height

    def to_rl_y(top, height):
        return page_h - top - height

    def round_rect_top(x, top, w, h, radius, *, fill=None, stroke=None, line_width=1.0):
        if fill is not None:
            c.setFillColorRGB(*(rgb(fill) if isinstance(fill, str) else fill))
        if stroke is not None:
            c.setStrokeColorRGB(*(rgb(stroke) if isinstance(stroke, str) else stroke))
        c.setLineWidth(line_width)
        c.roundRect(x, to_rl_y(top, h), w, h, radius, fill=1 if fill is not None else 0, stroke=1 if stroke is not None else 0)

    def rect_top(x, top, w, h, *, fill):
        c.setFillColorRGB(*(rgb(fill) if isinstance(fill, str) else fill))
        c.rect(x, to_rl_y(top, h), w, h, fill=1, stroke=0)

    def line_top(x1, top, x2, *, color):
        c.setStrokeColorRGB(*(rgb(color) if isinstance(color, str) else color))
        c.line(x1, page_h - top, x2, page_h - top)

    def circle_top(cx, cy_top, radius, *, fill=None, stroke=None, line_width=1.0):
        if fill is not None:
            c.setFillColorRGB(*(rgb(fill) if isinstance(fill, str) else fill))
        if stroke is not None:
            c.setStrokeColorRGB(*(rgb(stroke) if isinstance(stroke, str) else stroke))
        c.setLineWidth(line_width)
        c.circle(cx, page_h - cy_top, radius, fill=1 if fill is not None else 0, stroke=1 if stroke is not None else 0)

    def slide_bg():
        rect_top(0, 0, page_w, page_h, fill=_PRES_BG)

    def slide_num(page_no):
        draw_line(
            f"{page_no:02d} / 07", page_w - 140, 24,
            font=font_b, size=11, line_height=13.2,
            color=blend(_PRES_WHITE, 0.20), align="right", width=100,
            char_space=1.1,
        )

    def centered_top(content_height):
        return (page_h - content_height) / 2.0

    def section_metrics(slide):
        section_h = 13.2 + 24.0
        h2_h = text_block_height(slide.get("title", ""), font_b, H2_SIZE, H2_LH, content_w) + 4.0
        hsub_h = text_block_height(slide.get("subtitle", ""), font_r, HSUB_SIZE, HSUB_LH, content_w)
        return section_h, h2_h, hsub_h

    def draw_section_header(slide, top):
        draw_line(
            str(slide.get("section", "")).upper(), 64, top,
            font=font_b, size=11, line_height=13.2,
            color=blend(_PRES_WHITE, 0.35), char_space=11 * 0.18,
        )
        top += 13.2 + 24.0
        h2_h = draw_text_block(
            slide.get("title", ""), 64, top, content_w,
            font=font_b, size=H2_SIZE, line_height=H2_LH, color=_PRES_WHITE,
        )
        top += h2_h + 4.0
        hsub_h = draw_text_block(
            slide.get("subtitle", ""), 64, top, content_w,
            font=font_r, size=HSUB_SIZE, line_height=HSUB_LH,
            color=blend(_PRES_WHITE, 0.40),
        )
        return top + hsub_h

    def draw_flag_icon(x, top):
        rect_top(x, top, 20, 6, fill=_PRES_BLUE)
        rect_top(x, top + 6, 20, 6, fill=_PRES_YELLOW)

    def draw_calendar_icon(x, top, *, quarter=False):
        color = blend(_PRES_WHITE, 0.55)
        c.setStrokeColorRGB(*color)
        c.setLineWidth(1.2)
        c.roundRect(x, to_rl_y(top + 1, 12), 13, 12, 2, fill=0, stroke=1)
        line_top(x + 1, top + 5, x + 12, color=color)
        line_top(x + 4, top, x + 4, color=color)
        line_top(x + 9, top, x + 9, color=color)
        if quarter:
            circle_top(x + 9.5, top + 9, 1.5, fill=color)

    def draw_building_icon(x, top):
        color = blend(_PRES_WHITE, 0.55)
        c.setStrokeColorRGB(*color)
        c.setLineWidth(1.2)
        c.rect(x, to_rl_y(top + 1, 12), 14, 12, fill=0, stroke=1)
        for dx in (3, 7, 11):
            line_top(x + dx, top + 4, x + dx, color=color)
            line_top(x + dx, top + 10, x + dx, color=color)

    def draw_pin_icon(x, top):
        color = blend(_PRES_WHITE, 0.55)
        circle_top(x + 6, top + 5, 4.5, stroke=color, line_width=1.2)
        line_top(x + 6, top + 9, x + 6, color=color)
        line_top(x + 6, top + 9, x + 3, color=color)
        line_top(x + 6, top + 9, x + 9, color=color)

    def draw_clock_icon(x, top):
        color = blend(_PRES_WHITE, 0.55)
        circle_top(x + 6.5, top + 6.5, 6, stroke=color, line_width=1.2)
        line_top(x + 6.5, top + 6.5, x + 6.5, color=color)
        c.line(x + 6.5, page_h - (top + 6.5), x + 10, page_h - (top + 8.5))

    def draw_doc_icon(x, top, color):
        c.setStrokeColorRGB(*color)
        c.setLineWidth(1.0)
        c.rect(x, to_rl_y(top, 10), 8, 10, fill=0, stroke=1)
        line_top(x + 2, top + 3, x + 6, color=color)
        line_top(x + 2, top + 6, x + 6, color=color)

    def draw_chart_icon(x, top, color):
        for idx, height in enumerate((4, 7, 10)):
            rect_top(x + idx * 3.5, top + 10 - height, 2.2, height, fill=color)

    def draw_target_icon(x, top, color):
        circle_top(x + 5, top + 5, 4.5, stroke=color, line_width=1.0)
        circle_top(x + 5, top + 5, 1.6, stroke=color, line_width=1.0)

    def filter_icon(marker, index, x, top):
        if marker == "📅" or index == 0:
            draw_calendar_icon(x, top)
        elif marker == "🗓" or index == 1:
            draw_calendar_icon(x, top, quarter=True)
        elif marker == "🏢" or index == 2:
            draw_building_icon(x, top)
        elif marker == "📌" or index == 3:
            draw_pin_icon(x, top)
        else:
            draw_clock_icon(x, top)

    def pill_layout(pills):
        rows = []
        row = []
        row_w = 0.0
        for index, raw in enumerate(pills):
            marker, text = leading_marker(raw)
            width = 16.0 + 14.0 + 6.0 + text_width(text, font_b, 12) + 16.0
            if row and row_w + 10.0 + width > content_w:
                rows.append(row)
                row = []
                row_w = 0.0
            if row:
                row_w += 10.0
            row.append((index, marker, text, width, row_w))
            row_w += width
        if row:
            rows.append(row)
        return rows

    def render_title(slide):
        eyebrow_h = 13.2
        title_h = text_block_height(slide.get("title", ""), font_b, TITLE_SIZE, TITLE_LH, 800.0)
        sub_h = text_block_height(slide.get("subtitle", ""), font_r, TITLE_SUB_SIZE, TITLE_SUB_LH, 600.0)
        pill_rows = pill_layout(list(slide.get("filter_pills") or []))
        pill_h = len(pill_rows) * 26.4 + max(0, len(pill_rows) - 1) * 10.0
        total_h = eyebrow_h + 20.0 + title_h + 16.0 + sub_h + 40.0 + 8.0 + pill_h
        top = centered_top(total_h)

        marker, eyebrow = leading_marker(slide.get("eyebrow", ""))
        draw_flag_icon(64, top + 0.6)
        draw_line(
            eyebrow.upper(), 92, top,
            font=font_b, size=11, line_height=13.2,
            color=_PRES_YELLOW, char_space=11 * 0.20,
        )
        top += eyebrow_h + 20.0
        top += draw_text_block(
            slide.get("title", ""), 64, top, 800.0,
            font=font_b, size=TITLE_SIZE, line_height=TITLE_LH, color=_PRES_WHITE,
        ) + 16.0
        top += draw_text_block(
            slide.get("subtitle", ""), 64, top, 600.0,
            font=font_r, size=TITLE_SUB_SIZE, line_height=TITLE_SUB_LH,
            color=blend(_PRES_WHITE, 0.50),
        ) + 40.0 + 8.0

        for row_index, row in enumerate(pill_rows):
            row_top = top + row_index * (26.4 + 10.0)
            x = 64.0
            for index, marker, text, width, _ in row:
                round_rect_top(
                    x, row_top, width, 26.4, 13.2,
                    fill=blend(_PRES_WHITE, 0.06),
                    stroke=blend(_PRES_WHITE, 0.12),
                )
                filter_icon(marker, index, x + 16.0, row_top + 6.2)
                draw_line(
                    text, x + 36.0, row_top + 6.0,
                    font=font_b, size=12, line_height=14.4,
                    color=blend(_PRES_WHITE, 0.70),
                )
                x += width + 10.0

    def verdict_card_height(item, card_w):
        inner_w = card_w - 36.0
        label_h = text_block_height(item.get("label", ""), font_b, 11, 13.2, inner_w, 11 * 0.10)
        sub_h = text_block_height(item.get("subtitle", ""), font_r, 12, 14.4, inner_w)
        return 20.0 + label_h + 8.0 + 44.0 + 4.0 + sub_h + 20.0

    def render_verdict(slide):
        section_h = 13.2 + 24.0
        badge_h = 20.0 + VERDICT_BADGE_SIZE * 1.2
        verdict_text_h = text_block_height(
            slide.get("text", ""), font_r, VERDICT_BODY_SIZE,
            VERDICT_BODY_SIZE * 1.7, 680.0,
        )
        cards = list(slide.get("cards") or [])[:3]
        card_w = (680.0 - 40.0) / 3.0
        grid_h = max([verdict_card_height(item, card_w) for item in cards] or [0.0])
        total_h = section_h + badge_h + 20.0 + verdict_text_h + 40.0 + grid_h
        top = centered_top(total_h)

        draw_line(
            str(slide.get("section", "")).upper(), 64, top,
            font=font_b, size=11, line_height=13.2,
            color=blend(_PRES_WHITE, 0.35), char_space=11 * 0.18,
        )
        top += section_h

        severity = str(slide.get("severity") or "medium")
        accent = {"high": _PRES_RED, "medium": _PRES_AMBER, "low": _PRES_GREEN}.get(severity, _PRES_AMBER)
        _, verdict_title = leading_marker(slide.get("title", ""))
        title_w = text_width(verdict_title, font_b, VERDICT_BADGE_SIZE)
        badge_w = 24.0 + 12.0 + 10.0 + title_w + 24.0
        round_rect_top(
            64, top, badge_w, badge_h, 10,
            fill=blend(accent, 0.18), stroke=accent, line_width=1.5,
        )
        circle_top(64 + 24 + 6, top + badge_h / 2.0, 6, fill=accent)
        draw_line(
            verdict_title, 64 + 24 + 12 + 10, top + 10.0,
            font=font_b, size=VERDICT_BADGE_SIZE,
            line_height=VERDICT_BADGE_SIZE * 1.2, color=accent,
        )
        top += badge_h + 20.0

        top += draw_text_block(
            slide.get("text", ""), 64, top, 680.0,
            font=font_r, size=VERDICT_BODY_SIZE,
            line_height=VERDICT_BODY_SIZE * 1.7,
            color=blend(_PRES_WHITE, 0.55),
        ) + 40.0

        for idx, item in enumerate(cards):
            x = 64.0 + idx * (card_w + 20.0)
            round_rect_top(
                x, top, card_w, grid_h, 12,
                fill=blend(_PRES_WHITE, 0.04),
                stroke=blend(_PRES_WHITE, 0.08),
            )
            inner_x = x + 18.0
            inner_w = card_w - 36.0
            cursor = top + 20.0
            label_h = draw_text_block(
                str(item.get("label", "")).upper(), inner_x, cursor, inner_w,
                font=font_b, size=11, line_height=13.2,
                color=blend(_PRES_WHITE, 0.35), char_space=11 * 0.10,
            )
            cursor += label_h + 8.0
            draw_line(
                item.get("value_text", ""), inner_x, cursor,
                font=font_b, size=44, line_height=44, color=item.get("color") or _PRES_WHITE,
            )
            cursor += 44.0 + 4.0
            draw_text_block(
                item.get("subtitle", ""), inner_x, cursor, inner_w,
                font=font_r, size=12, line_height=14.4,
                color=blend(_PRES_WHITE, 0.35),
            )

    def kpi_card_height(item, card_w):
        inner_w = card_w - 48.0
        label_h = text_block_height(item.get("label", ""), font_b, 11, 13.2, inner_w, 11 * 0.08)
        return 28.0 + label_h + 6.0 + KPI_VALUE_SIZE + 6.0 + 15.6 + 28.0

    def render_key_metrics(slide):
        section_h, h2_h, hsub_h = section_metrics(slide)
        cards = list(slide.get("cards") or [])[:8]
        card_w = (content_w - 60.0) / 4.0
        row1 = cards[:4]
        row2 = cards[4:8]
        row_heights = []
        for row in (row1, row2):
            if row:
                row_heights.append(max(kpi_card_height(item, card_w) for item in row))
        grid_h = sum(row_heights) + 20.0 * max(0, len(row_heights) - 1)
        metric_row_h = 19.2
        bars = list(slide.get("bars") or [])[:4]
        bars_h = len(bars) * metric_row_h + 20.0 * max(0, len(bars) - 1)
        total_h = section_h + h2_h + hsub_h + 32.0 + grid_h + 40.0 + bars_h
        top = centered_top(total_h)
        top = draw_section_header(slide, top)
        top += 32.0

        row_top = top
        card_index = 0
        for row_height, row in zip(row_heights, (row1, row2)):
            for col, item in enumerate(row):
                x = 64.0 + col * (card_w + 20.0)
                round_rect_top(
                    x, row_top, card_w, row_height, 14,
                    fill=blend(_PRES_WHITE, 0.04),
                    stroke=blend(_PRES_WHITE, 0.08),
                )
                rect_top(x, row_top, card_w, 3.0, fill=item.get("color") or _PRES_METRIC_BLUE)
                cursor = row_top + 28.0
                inner_x = x + 24.0
                inner_w = card_w - 48.0
                label_h = draw_text_block(
                    str(item.get("label", "")).upper(), inner_x, cursor, inner_w,
                    font=font_b, size=11, line_height=13.2,
                    color=blend(_PRES_WHITE, 0.40), char_space=11 * 0.08,
                )
                cursor += label_h + 6.0
                draw_line(
                    item.get("value_text", ""), inner_x, cursor,
                    font=font_b, size=KPI_VALUE_SIZE,
                    line_height=KPI_VALUE_SIZE, color=_PRES_WHITE,
                )
                cursor += KPI_VALUE_SIZE + 6.0
                draw_line(
                    item.get("sub_text", ""), inner_x, cursor,
                    font=font_b, size=13, line_height=15.6,
                    color=blend(_PRES_WHITE, 0.35),
                )
                card_index += 1
            row_top += row_height + 20.0

        top += grid_h + 40.0
        metric_w = 680.0
        label_w = 220.0
        value_w = 56.0
        track_w = metric_w - label_w - value_w - 40.0
        for idx, item in enumerate(bars):
            row_y = top + idx * (metric_row_h + 20.0)
            draw_line(
                item.get("label", ""), 64, row_y + (metric_row_h - 15.6) / 2.0,
                font=font_b, size=13, line_height=15.6,
                color=blend(_PRES_WHITE, 0.55),
            )
            track_x = 64.0 + label_w + 20.0
            track_top = row_y + (metric_row_h - 12.0) / 2.0
            round_rect_top(
                track_x, track_top, track_w, 12.0, 6.0,
                fill=blend(_PRES_WHITE, 0.06),
            )
            try:
                pct = min(max(float(item.get("value") or 0), 0), 100)
            except (TypeError, ValueError):
                pct = 0.0
            if pct > 0:
                round_rect_top(
                    track_x, track_top, track_w * pct / 100.0, 12.0, 6.0,
                    fill=item.get("color") or _PRES_BLUE,
                )
            draw_line(
                item.get("value_text", ""),
                track_x + track_w + 20.0, row_y,
                font=font_b, size=16, line_height=19.2,
                color=_PRES_WHITE, align="right", width=value_w,
            )

    def render_goals(slide):
        section_h, h2_h, hsub_h = section_metrics(slide)
        rows = list(slide.get("rows") or [])
        row_h = 15.6
        rows_h = len(rows) * row_h + 14.0 * max(0, len(rows) - 1)
        if not rows:
            rows_h = 15.6
        total_h = section_h + h2_h + hsub_h + 28.0 + rows_h
        top = centered_top(total_h)
        top = draw_section_header(slide, top)
        top += 28.0
        if not rows:
            draw_line(
                slide.get("empty_text", "Дані відсутні за обраними фільтрами"),
                64, top, font=font_r, size=13, line_height=15.6,
                color=blend(_PRES_WHITE, 0.30),
            )
            return

        code_w = 36.0
        name_w = 320.0
        pct_w = 44.0
        bar_w = content_w - code_w - name_w - pct_w - 48.0
        code_x = 64.0
        name_x = code_x + code_w + 16.0
        bar_x = name_x + name_w + 16.0
        pct_x = bar_x + bar_w + 16.0
        for idx, row in enumerate(rows):
            row_top = top + idx * (row_h + 14.0)
            draw_line(
                row.get("code", ""), code_x, row_top,
                font=font_b, size=11, line_height=13.2,
                color=blend(_PRES_WHITE, 0.40), align="right", width=code_w,
            )
            name = ellipsize(row.get("name", ""), font_r, 13, name_w)
            draw_line(
                name, name_x, row_top,
                font=font_r, size=13, line_height=15.6,
                color=blend(_PRES_WHITE, 0.70),
            )
            track_top = row_top + (row_h - 10.0) / 2.0
            round_rect_top(
                bar_x, track_top, bar_w, 10.0, 5.0,
                fill=blend(_PRES_WHITE, 0.06),
            )
            try:
                pct = min(max(float(row.get("value") or 0), 0), 100)
            except (TypeError, ValueError):
                pct = 0.0
            if pct > 0:
                round_rect_top(
                    bar_x, track_top, bar_w * pct / 100.0, 10.0, 5.0,
                    fill=row.get("color") or _PRES_GREY,
                )
            draw_line(
                row.get("value_text", ""), pct_x, row_top,
                font=font_b, size=13, line_height=15.6,
                color=_PRES_WHITE, align="right", width=pct_w,
            )

    def risk_card_height(item, card_w):
        inner_w = card_w - 48.0
        _, label = leading_marker(item.get("label", ""))
        label_h = text_block_height(label, font_b, 11, 13.2, inner_w - 15.0, 11 * 0.10)
        return 28.0 + label_h + 8.0 + RISK_VALUE_SIZE + 8.0 + 15.6 + 28.0

    def tag_layout(tags, max_width):
        result = []
        x = 0.0
        row = 0
        for tag in tags:
            text = plain(tag)
            width = text_width(text, font_b, 11) + 24.0
            if x and x + 12.0 + width > max_width:
                row += 1
                x = 0.0
            if x:
                x += 12.0
            result.append((row, x, width, text))
            x += width
        rows = (max((item[0] for item in result), default=-1) + 1)
        height = rows * 23.2 + max(0, rows - 1) * 12.0
        return result, height

    def render_risks(slide):
        section_h, h2_h, hsub_h = section_metrics(slide)
        cards = list(slide.get("cards") or [])[:3]
        card_w = (content_w - 40.0) / 3.0
        grid_h = max([risk_card_height(item, card_w) for item in cards] or [0.0])
        summary_inner_w = 640.0 - 56.0
        summary_text_h = text_block_height(
            slide.get("summary_text", ""), font_r, 15, 25.5, summary_inner_w,
        )
        tags = list(slide.get("tags") or [])[:2]
        tag_positions, tags_h = tag_layout(tags, summary_inner_w)
        summary_h = 24.0 + 13.2 + 12.0 + summary_text_h + 16.0 + tags_h + 24.0
        total_h = section_h + h2_h + hsub_h + 32.0 + grid_h + 48.0 + summary_h
        top = centered_top(total_h)
        top = draw_section_header(slide, top)
        top += 32.0

        for idx, item in enumerate(cards):
            x = 64.0 + idx * (card_w + 20.0)
            accent = item.get("color") or (_PRES_RED, _PRES_AMBER, _PRES_BRIGHT_GREEN)[idx]
            round_rect_top(
                x, top, card_w, grid_h, 14,
                fill=blend(accent, 0.10),
                stroke=blend(accent, 0.30), line_width=1.5,
            )
            cursor = top + 28.0
            marker, label = leading_marker(item.get("label", ""))
            circle_top(x + 24.0 + 5.0, cursor + 6.6, 4.5, fill=accent)
            label_h = draw_text_block(
                label.upper(), x + 24.0 + 15.0, cursor,
                card_w - 48.0 - 15.0,
                font=font_b, size=11, line_height=13.2,
                color=accent, char_space=11 * 0.10,
            )
            cursor += label_h + 8.0
            draw_line(
                item.get("value_text", ""), x + 24.0, cursor,
                font=font_b, size=RISK_VALUE_SIZE,
                line_height=RISK_VALUE_SIZE, color=accent,
            )
            cursor += RISK_VALUE_SIZE + 8.0
            draw_line(
                item.get("sub_text", ""), x + 24.0, cursor,
                font=font_b, size=13, line_height=15.6,
                color=blend(_PRES_WHITE, 0.40),
            )

        top += grid_h + 48.0
        round_rect_top(
            64.0, top, 640.0, summary_h, 14,
            fill=blend(_PRES_WHITE, 0.03),
            stroke=blend(_PRES_WHITE, 0.07),
        )
        cursor = top + 24.0
        draw_line(
            str(slide.get("summary_label", "")).upper(), 92.0, cursor,
            font=font_b, size=11, line_height=13.2,
            color=blend(_PRES_WHITE, 0.30), char_space=11 * 0.12,
        )
        cursor += 13.2 + 12.0
        cursor += draw_text_block(
            slide.get("summary_text", ""), 92.0, cursor, summary_inner_w,
            font=font_r, size=15, line_height=25.5,
            color=blend(_PRES_WHITE, 0.70),
        ) + 16.0
        for row, x_off, width, text in tag_positions:
            tag_top = cursor + row * (23.2 + 12.0)
            round_rect_top(
                92.0 + x_off, tag_top, width, 23.2, 6,
                fill=blend(_PRES_WHITE, 0.06),
                stroke=blend(_PRES_WHITE, 0.10),
            )
            draw_line(
                text, 92.0 + x_off + 12.0, tag_top + 5.0,
                font=font_b, size=11, line_height=13.2,
                color=blend(_PRES_WHITE, 0.50),
            )

    def top5_meta_layout(row, available_width):
        items = [
            ("doc", plain(row.get("code", ""))),
            ("building", plain(row.get("department", ""))),
            ("chart", plain(row.get("status", ""))),
            ("target", f"Виконання: {plain(row.get('performance_text', ''))}"),
        ]
        result = []
        x = 0.0
        line = 0
        for kind, text in items:
            width = 10.0 + 5.0 + text_width(text, font_r, 10)
            if x and x + 10.0 + width > available_width:
                line += 1
                x = 0.0
            if x:
                x += 10.0
            result.append((line, x, width, kind, text))
            x += width
        lines = max((item[0] for item in result), default=-1) + 1
        height = lines * 12.0 + max(0, lines - 1) * 10.0
        return result, height

    def top5_row_metrics(row):
        badge_text = plain(row.get("risk_label", ""))
        badge_w = text_width(badge_text, font_b, 10) + 16.0
        right_w = 860.0 - badge_w - 14.0
        name_h = text_block_height(row.get("name", ""), font_b, 13, 18.2, right_w)
        meta, meta_h = top5_meta_layout(row, right_w)
        content_h = name_h + 5.0 + meta_h
        row_h = 28.0 + max(20.0, content_h)
        return badge_w, right_w, name_h, meta, meta_h, row_h

    def draw_meta_icon(kind, x, top, color):
        if kind == "doc":
            draw_doc_icon(x, top + 1.0, color)
        elif kind == "building":
            draw_building_icon(x, top - 1.0)
        elif kind == "chart":
            draw_chart_icon(x, top + 1.0, color)
        else:
            draw_target_icon(x, top + 1.0, color)

    def render_top5(slide):
        section_h, h2_h, hsub_h = section_metrics(slide)
        rows = list(slide.get("rows") or [])[:5]
        metrics = [top5_row_metrics(row) for row in rows]
        list_h = sum(item[-1] for item in metrics)
        if not rows:
            list_h = 15.6
        total_h = section_h + h2_h + hsub_h + 28.0 + list_h
        top = centered_top(total_h)
        top = draw_section_header(slide, top)
        top += 28.0
        if not rows:
            draw_line(
                slide.get("empty_text", "Критичних заходів не виявлено"), 64, top,
                font=font_r, size=13, line_height=15.6,
                color=blend(_PRES_WHITE, 0.30),
            )
            return

        row_top = top
        meta_color = blend(_PRES_WHITE, 0.35)
        for row, metric in zip(rows, metrics):
            badge_w, right_w, name_h, meta, meta_h, row_h = metric
            badge_text = plain(row.get("risk_label", ""))
            badge_top = row_top + 14.0 + 2.0
            round_rect_top(
                64.0, badge_top, badge_w, 18.0, 6,
                fill=row.get("risk_color") or _PRES_GREY,
            )
            draw_line(
                badge_text, 72.0, badge_top + 3.0,
                font=font_b, size=10, line_height=12.0, color=_PRES_BG,
            )
            right_x = 64.0 + badge_w + 14.0
            cursor = row_top + 14.0
            name_h_drawn = draw_text_block(
                row.get("name", ""), right_x, cursor, right_w,
                font=font_b, size=13, line_height=18.2,
                color=blend(_PRES_WHITE, 0.85),
            )
            cursor += name_h_drawn + 5.0
            for line_idx, x_off, width, kind, text in meta:
                item_top = cursor + line_idx * (12.0 + 10.0)
                draw_meta_icon(kind, right_x + x_off, item_top, meta_color)
                draw_line(
                    text, right_x + x_off + 15.0, item_top,
                    font=font_r, size=10, line_height=12.0,
                    color=meta_color,
                )
            line_top(64.0, row_top + row_h, 924.0, color=blend(_PRES_WHITE, 0.06))
            row_top += row_h

    def render_finance(slide):
        section_h, h2_h, hsub_h = section_metrics(slide)
        groups = list(slide.get("groups") or [])[:4]
        source_header_h = 13.2 + 20.0
        source_item_h = 15.6 + 5.0 + 10.0 + 16.0
        budget_h = 20.0 + 13.2 + 8.0 + 36.0 + 4.0 + 14.4 + 20.0
        left_h = source_header_h + len(groups) * source_item_h + 24.0 + budget_h
        kpkvk_rows = list(slide.get("kpkvk_rows") or [])[:6]
        right_header_h = 13.2 + 20.0
        kpkvk_row_h = 20.0 + 16.8
        right_h = right_header_h + max(15.6, len(kpkvk_rows) * kpkvk_row_h)
        grid_h = max(left_h, right_h)
        total_h = section_h + h2_h + hsub_h + 36.0 + grid_h
        top = centered_top(total_h)
        top = draw_section_header(slide, top)
        top += 36.0

        left_x = 64.0
        right_x = 64.0 + 430.0 + 40.0
        col_w = 430.0

        draw_line(
            str(slide.get("sources_label", "")).upper(), left_x, top,
            font=font_b, size=11, line_height=13.2,
            color=blend(_PRES_WHITE, 0.30), char_space=11 * 0.12,
        )
        cursor = top + source_header_h
        for item in groups:
            draw_line(
                item.get("label", ""), left_x, cursor,
                font=font_b, size=13, line_height=15.6,
                color=blend(_PRES_WHITE, 0.70),
            )
            display = str(item.get("display", ""))
            draw_line(
                display, left_x, cursor,
                font=font_b, size=13, line_height=15.6,
                color=_PRES_WHITE, align="right", width=col_w,
            )
            bar_top = cursor + 15.6 + 5.0
            round_rect_top(
                left_x, bar_top, col_w, 10.0, 5.0,
                fill=blend(_PRES_WHITE, 0.07),
            )
            try:
                pct = min(max(float(item.get("percent") or 0), 0), 100)
            except (TypeError, ValueError):
                pct = 0.0
            if pct > 0:
                round_rect_top(
                    left_x, bar_top, col_w * pct / 100.0, 10.0, 5.0,
                    fill=item.get("color") or _PRES_GREY,
                )
            cursor += source_item_h

        cursor += 24.0
        budget = dict(slide.get("budget") or {})
        round_rect_top(
            left_x, cursor, col_w, budget_h, 12,
            fill=blend(_PRES_BLUE, 0.12),
            stroke=blend(_PRES_BLUE, 0.25),
        )
        bcur = cursor + 20.0
        draw_line(
            str(budget.get("label", "")).upper(), left_x + 22.0, bcur,
            font=font_b, size=11, line_height=13.2,
            color=blend(_PRES_WHITE, 0.30), char_space=11 * 0.10,
        )
        bcur += 13.2 + 8.0
        draw_line(
            budget.get("value_text", ""), left_x + 22.0, bcur,
            font=font_b, size=36, line_height=36, color=_PRES_WHITE,
        )
        bcur += 36.0 + 4.0
        draw_line(
            budget.get("subtitle", ""), left_x + 22.0, bcur,
            font=font_r, size=12, line_height=14.4,
            color=blend(_PRES_WHITE, 0.30),
        )

        draw_line(
            str(slide.get("kpkvk_label", "")).upper(), right_x, top,
            font=font_b, size=11, line_height=13.2,
            color=blend(_PRES_WHITE, 0.30), char_space=11 * 0.12,
        )
        rcur = top + right_header_h
        if not kpkvk_rows:
            draw_line(
                slide.get("kpkvk_empty_text", "КПКВК не визначено"), right_x, rcur,
                font=font_r, size=12, line_height=14.4,
                color=blend(_PRES_WHITE, 0.30),
            )
            return
        for row in kpkvk_rows:
            row_top = rcur + 10.0
            draw_line(
                row.get("code", ""), right_x, row_top,
                font=font_b, size=14, line_height=16.8, color=_PRES_YELLOW,
            )
            draw_line(
                row.get("count_text", ""), right_x + 150.0, row_top,
                font=font_r, size=12, line_height=14.4,
                color=blend(_PRES_WHITE, 0.50),
            )
            draw_line(
                row.get("budget_text", ""), right_x, row_top,
                font=font_b, size=12, line_height=14.4,
                color=blend(_PRES_WHITE, 0.70), align="right", width=col_w,
            )
            rcur += kpkvk_row_h
            line_top(right_x, rcur, right_x + col_w, color=blend(_PRES_WHITE, 0.06))

    renderers = {
        "title": render_title,
        "verdict": render_verdict,
        "key_metrics": render_key_metrics,
        "strategic_goals": render_goals,
        "risks": render_risks,
        "top5": render_top5,
        "finance": render_finance,
    }

    for page_no, slide in enumerate(presentation_payload["slides"], start=1):
        slide_bg()
        slide_num(page_no)
        renderers[slide["key"]](slide)
        c.showPage()

    c.save()
    return buffer.getvalue()


# ------------------------------------------------------------
# Табличні вивантаження заходів (правка К9 — перенесено з app.py)
# ------------------------------------------------------------

def build_measures_export_df(measures, quarter_data, quarter_columns):
    """
    Формує DataFrame вивантаження заходів.

    quarter_columns: список (year, quarter, label) — колонки квартальних
    значень (сторінка передає власний перелік, напр. з get_quarter_columns).
    """
    import pandas as pd
    from core.text_utils import raw_value, strip_leading_code

    rows = []
    for _, measure in measures.iterrows():
        code = raw_value(measure.get("code", ""))
        row = {
            "Код": code,
            "Захід": strip_leading_code(measure.get("name", ""), code),
            "Тип продукту": raw_value(measure.get("product_type", "")),
            "Індикатор": raw_value(measure.get("indicator", "")),
            "Одиниці виміру": raw_value(measure.get("unit", "")),
            "Головний виконавець": raw_value(measure.get("resp_main", "")),
            "Співвиконавець": raw_value(measure.get("resp_co_1", "")),
        }
        for year, quarter, label in quarter_columns:
            key = f"{year}_{quarter}"
            item = quarter_data.get(code, {}).get(key, {})
            row[label] = item.get("value", "")
        rows.append(row)

    return pd.DataFrame(rows)


def dataframe_to_excel_bytes(df, sheet_name: str = "Заходи") -> bytes:
    """DataFrame → байти .xlsx (одна вкладка), з охайним оформленням.

    Тонкий зручний виклик поверх write_styled_excel — лишений для
    сумісності з існуючими викликами (app.py тощо), але тепер теж
    видає красиво оформлений файл, а не голий df.to_excel().
    """
    return write_styled_excel({sheet_name: df})


# ------------------------------------------------------------
# Охайне оформлення Excel-вивантажень (пункт запиту:
# "Зроби нормальні експорти в ексель... щоб воно вивантажувалося
# так само гарно як і ексель Під моніторинг СП")
# ------------------------------------------------------------
#
# Раніше практично кожне вивантаження в системі (app.py, Аналітика,
# Журнал дій) робило голий df.to_excel(writer, ...) — без жодного
# форматування: ані ширини колонок, ані закріпленої шапки, ані меж
# клітинок. write_styled_excel() — єдина спільна функція, яка тепер
# застосовується всюди: жирна темна шапка з білим текстом (той самий
# візуальний стиль, що й у "Під моніторинг СП.xlsx" — колір #032A63),
# закріplena шапка, автофільтр, розумна ширина колонок за вмістом,
# перенос тексту в довгих колонках, тонкі межі й легке "зебра"-заливання
# рядків для зручного читання великих вивантажень.

_HEADER_BG = "#032A63"
_HEADER_FG = "#FFFFFF"
_BORDER_COLOR = "#DCE4F0"
_BAND_BG = "#F7F9FC"

_MIN_COL_WIDTH = 10
_MAX_COL_WIDTH = 60
_WRAP_WIDTH_THRESHOLD = 38


def _estimate_column_width(series) -> int:
    """Приблизна ширина колонки за вмістом (символи), з розумними межами."""
    try:
        header_len = len(str(series.name))
    except Exception:
        header_len = _MIN_COL_WIDTH

    try:
        non_empty = series.dropna().astype(str)
        max_len = int(non_empty.map(len).max()) if len(non_empty) else header_len
    except Exception:
        max_len = header_len

    width = max(header_len, max_len) + 2
    return max(_MIN_COL_WIDTH, min(width, _MAX_COL_WIDTH))


def write_styled_excel(
    sheets: dict,
    *,
    freeze_first_col: int = 0,
    extra_sheets_no_style: dict | None = None,
) -> bytes:
    """
    Формує .xlsx з кількох аркушів із єдиним охайним оформленням.

    sheets: {назва_аркуша: DataFrame} — кожен аркуш отримує:
      - жирну темну шапку з білим текстом, по центру, з переносом;
      - закріплений перший рядок (і, за потреби, перші freeze_first_col
        колонок — зручно для широких таблиць на кшталт "Заходів", де
        хочеться завжди бачити код/назву під час горизонтальної прокрутки);
      - автофільтр на шапці;
      - ширину колонок за вмістом (у розумних межах);
      - тонкі межі й перенос тексту в довгих текстових колонках;
      - легке почергове заливання рядків для зручності читання.

    extra_sheets_no_style: {назва_аркуша: DataFrame} — додаткові аркуші
    (напр. "Параметри вивантаження"), які пишуться як є, без цього
    оформлення (короткі службові таблиці "параметр — значення").
    """
    import io
    import pandas as pd

    buffer = io.BytesIO()

    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        workbook = writer.book

        header_fmt = workbook.add_format({
            "bold": True, "font_color": _HEADER_FG, "bg_color": _HEADER_BG,
            "align": "center", "valign": "vcenter", "text_wrap": True,
            "border": 1, "border_color": "#FFFFFF",
        })
        cell_fmt = workbook.add_format({
            "valign": "top", "border": 1, "border_color": _BORDER_COLOR,
        })
        cell_wrap_fmt = workbook.add_format({
            "valign": "top", "text_wrap": True, "border": 1, "border_color": _BORDER_COLOR,
        })
        band_fmt = workbook.add_format({
            "valign": "top", "border": 1, "border_color": _BORDER_COLOR, "bg_color": _BAND_BG,
        })
        band_wrap_fmt = workbook.add_format({
            "valign": "top", "text_wrap": True, "border": 1,
            "border_color": _BORDER_COLOR, "bg_color": _BAND_BG,
        })

        def _fmt_for(row_idx: int, wrap: bool):
            banded = (row_idx % 2 == 0)  # парні рядки даних — легке заливання
            if wrap:
                return band_wrap_fmt if banded else cell_wrap_fmt
            return band_fmt if banded else cell_fmt

        for sheet_name, df in (sheets or {}).items():
            safe_name = str(sheet_name)[:31] or "Аркуш"
            df = df if df is not None else pd.DataFrame()
            n_rows, n_cols = df.shape

            worksheet = workbook.add_worksheet(safe_name)
            writer.sheets[safe_name] = worksheet

            if n_cols == 0:
                worksheet.write(0, 0, "Немає даних для вивантаження", header_fmt)
                continue

            wrap_flags = []
            for col_idx, col_name in enumerate(df.columns):
                width = _estimate_column_width(df[col_name])
                should_wrap = width >= _WRAP_WIDTH_THRESHOLD
                wrap_flags.append(should_wrap)
                worksheet.set_column(col_idx, col_idx, width)
                worksheet.write(0, col_idx, str(col_name), header_fmt)

            worksheet.set_row(0, 32)

            for row_idx in range(n_rows):
                row_values = df.iloc[row_idx]
                for col_idx, col_name in enumerate(df.columns):
                    value = row_values[col_name]
                    if value is None or (isinstance(value, float) and pd.isna(value)):
                        value = ""
                    fmt = _fmt_for(row_idx, wrap_flags[col_idx])
                    worksheet.write(row_idx + 1, col_idx, value, fmt)

            worksheet.freeze_panes(1, max(0, min(freeze_first_col, n_cols)))
            worksheet.autofilter(0, 0, n_rows, n_cols - 1)

        for sheet_name, df in (extra_sheets_no_style or {}).items():
            safe_name = str(sheet_name)[:31] or "Аркуш"
            df = df if df is not None else pd.DataFrame()
            df.to_excel(writer, index=False, sheet_name=safe_name)

    return buffer.getvalue()

# ------------------------------------------------------------
# DOCX-вивантаження таблиць у альбомному форматі (DEMO 1.9)
# ------------------------------------------------------------

def dataframes_to_docx_landscape(
    sheets: dict,
    *,
    title: str = "Табличний звіт",
    subtitle: str = "",
) -> bytes | None:
    """Формує DOCX з кількома таблицями в альбомній орієнтації.

    Повертає bytes або None, якщо python-docx недоступний. Таблиці навмисно
    компактні: повторювана шапка у Word не гарантується, але файл придатний
    для друку й службового долучення до матеріалів.
    """
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return None

    import io
    import pandas as pd

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(title)
    r.bold = True
    r.font.size = Pt(14)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)

    for sheet_name, df in sheets.items():
        if df is None:
            continue
        try:
            frame = pd.DataFrame(df).fillna("")
        except Exception as exc:
            show_warning(
                f"Аркуш «{sheet_name}» не додано до Word-файлу.",
                exc,
                "Перетворення аркуша для Word-експорту",
            )
            continue
        doc.add_paragraph()
        ph = doc.add_paragraph(str(sheet_name))
        ph.runs[0].bold = True
        ph.runs[0].font.size = Pt(10)
        if frame.empty:
            doc.add_paragraph("Даних немає.")
            continue

        max_rows = min(len(frame), 250)
        max_cols = min(len(frame.columns), 12)
        frame = frame.iloc[:max_rows, :max_cols]

        table = doc.add_table(rows=1, cols=len(frame.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col in enumerate(frame.columns):
            hdr[i].text = str(col)
            for paragraph in hdr[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(7)
            tc_pr = hdr[i]._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "434343")
            tc_pr.append(shd)

        for _, row in frame.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(frame.columns):
                value = str(row.get(col, ""))
                cells[i].text = value[:900]
                for paragraph in cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(6.5)

        if len(pd.DataFrame(df)) > max_rows or len(pd.DataFrame(df).columns) > max_cols:
            doc.add_paragraph(
                f"Примітка: для друкованої версії показано перші {max_rows} рядків і {max_cols} колонок. "
                "Повний масив доступний в Excel-вивантаженні."
            )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ------------------------------------------------------------
# Головний Excel-експорт сторінки app.py
# ------------------------------------------------------------

_MATRIX_DARK_BLUE = "#032A63"
_MATRIX_DARK_ORANGE = "#C65911"
_MATRIX_GOAL_BLUE = "#5B9BD5"
_MATRIX_TASK_BLUE = "#D9EAF7"
_MATRIX_WHITE = "#FFFFFF"
_MATRIX_GREY = "#F2F2F2"
_MATRIX_TEXT = "#132238"


def _export_clean(value) -> str:
    if value is None:
        return ""
    try:
        import pandas as pd
        if pd.isna(value):
            return ""
    except (TypeError, ValueError):
        pass
    text = str(value).strip()
    return "" if text.lower() in {"nan", "none", "nat"} else text


def _quarter_roman(value) -> str:
    from core.periods import quarter_key
    return quarter_key(value)


def _latest_period_records(monitoring_df):
    """Індекс (code, year, quarter) -> найактуальніший запис."""
    import pandas as pd
    result = {}
    if monitoring_df is None or monitoring_df.empty:
        return result
    data = monitoring_df.copy()
    for col in ["strat_code", "year", "quarter", "submitted_at", "id"]:
        if col not in data.columns:
            data[col] = ""
    data["_submitted_sort"] = pd.to_datetime(data["submitted_at"], errors="coerce", utc=True)
    data["_id_sort"] = pd.to_numeric(data["id"], errors="coerce").fillna(-1)
    data = data.sort_values(["_submitted_sort", "_id_sort"], na_position="first")
    for _, row in data.iterrows():
        key = (
            _export_clean(row.get("strat_code")),
            _export_clean(row.get("year")),
            _quarter_roman(row.get("quarter")),
        )
        if all(key):
            result[key] = row
    return result


def _indicator_latest_value(indicator_df, row, year) -> str:
    """Останнє подане значення конкретного індикатора за роком."""
    import pandas as pd
    if indicator_df is None or indicator_df.empty:
        return ""
    data = indicator_df.copy()
    code = _export_clean(row.get("code"))
    indicator_name = _export_clean(row.get("indicator"))
    if "strat_code" not in data.columns:
        return ""
    data = data[data["strat_code"].astype(str).str.strip() == code]
    if "year" in data.columns:
        data = data[data["year"].astype(str).str.strip() == str(year)]
    if indicator_name and "indicator_name" in data.columns:
        exact = data[data["indicator_name"].astype(str).str.strip() == indicator_name]
        if not exact.empty:
            data = exact
    if data.empty:
        return ""
    if "submitted_at" not in data.columns:
        data["submitted_at"] = ""
    if "id" not in data.columns:
        data["id"] = ""
    data["_submitted_sort"] = pd.to_datetime(data["submitted_at"], errors="coerce", utc=True)
    data["_id_sort"] = pd.to_numeric(data["id"], errors="coerce").fillna(-1)
    data = data.sort_values(["_submitted_sort", "_id_sort"], ascending=[False, False])
    return _export_clean(data.iloc[0].get("numeric_value", ""))


def _indicator_matches_export_filters(row, selected_ssp_indices=None, selected_product_types=None, search_query="") -> bool:
    from core.text_utils import extract_ssp_index
    selected_ssp_indices = {str(x).strip() for x in (selected_ssp_indices or []) if str(x).strip()}
    if selected_ssp_indices:
        values = [row.get("resp_main", ""), row.get("resp_co_1", ""), row.get("resp_co_2", "")]
        indexes = {extract_ssp_index(v) for v in values if _export_clean(v)}
        if not (indexes & selected_ssp_indices):
            return False
    selected_product_types = {_export_clean(x) for x in (selected_product_types or []) if _export_clean(x)}
    if selected_product_types and _export_clean(row.get("product_type")) not in selected_product_types:
        return False
    query = _export_clean(search_query).lower()
    if query:
        haystack = " ".join(_export_clean(row.get(col, "")) for col in [
            "code", "name", "indicator", "product_type", "resp_main", "resp_co_1", "resp_co_2"
        ]).lower()
        if query not in haystack:
            return False
    return True


def _goal_number(code: str) -> str:
    import re
    match = re.search(r"\d+", _export_clean(code))
    return match.group(0) if match else _export_clean(code)


def _approval_route_and_current(record, logs=None) -> tuple[str, str, str]:
    """(координатор, накопичувальна схема, поточна ланка/стан погодження)."""
    from core import approval_schemes as schemes

    if record is None:
        return "", "", ""
    chain = schemes.parse_chain(record.get("approval_chain", ""))
    coordinator = ""
    for stage in chain:
        if _export_clean(stage.get("role")) == "admin":
            coordinator = _export_clean(stage.get("name")) or _export_clean(stage.get("email"))
            break

    approval = _export_clean(record.get("approval_status"))
    stage_idx = schemes.parse_stage(record.get("chain_stage"))
    scheme_text = schemes.approval_scheme_text(chain, stage_idx, approval, logs)

    if approval == schemes.APPROVED_STATUS:
        current = schemes.APPROVED_STATUS
    elif approval in schemes.ALL_RETURNED_STATUSES:
        current = approval
    elif approval == schemes.STATUS_WAITING_MANAGER_SELECTION:
        current = schemes.STATUS_WAITING_MANAGER_SELECTION
    else:
        stage = schemes.current_stage(chain, stage_idx)
        if stage:
            who = _export_clean(stage.get("name")) or _export_clean(stage.get("email"))
            label = _export_clean(stage.get("label"))
            current = f"{label} ({who})" if who else label
        else:
            current = approval
    return coordinator, scheme_text, current


def _load_export_logs():
    import pandas as pd
    try:
        from core.db import fetch_all
        return pd.DataFrame(fetch_all("monitoring_logs", "*", order=("changed_at", False)))
    except Exception:
        return pd.DataFrame()


def _request_export_logs(logs_df, request_id):
    """Журнал однієї заявки у хронологічному порядку для схеми експорту."""
    import pandas as pd

    if logs_df is None or logs_df.empty or request_id in (None, ""):
        return pd.DataFrame()
    if "request_id" not in logs_df.columns:
        return pd.DataFrame()
    try:
        rid = int(float(str(request_id)))
    except (TypeError, ValueError):
        return pd.DataFrame()
    mask = pd.to_numeric(logs_df["request_id"], errors="coerce") == rid
    data = logs_df[mask].copy()
    if data.empty or "changed_at" not in data.columns:
        return data
    data["_scheme_sort"] = pd.to_datetime(data["changed_at"], errors="coerce", utc=True)
    return data.sort_values("_scheme_sort", na_position="last")


def _approval_history(logs_df, request_id) -> str:
    import pandas as pd
    if logs_df is None or logs_df.empty or request_id in (None, "") or "request_id" not in logs_df.columns:
        return ""
    try:
        rid = int(float(str(request_id)))
        data = logs_df[pd.to_numeric(logs_df["request_id"], errors="coerce") == rid].copy()
    except Exception:
        return ""
    if data.empty:
        return ""
    if "changed_at" not in data.columns:
        data["changed_at"] = ""
    data["_sort"] = pd.to_datetime(data["changed_at"], errors="coerce", utc=True)
    data = data.sort_values("_sort", na_position="last")
    parts = []
    for _, row in data.iterrows():
        ts = pd.to_datetime(row.get("changed_at"), errors="coerce", utc=True)
        if pd.notna(ts):
            try:
                stamp = ts.tz_convert("Europe/Kyiv").strftime("%d.%m.%Y %H:%M")
            except Exception:
                stamp = str(ts)
        else:
            stamp = ""
        actor = _export_clean(row.get("actor_name")) or _export_clean(row.get("changed_by")) or _export_clean(row.get("actor_email"))
        action = _export_clean(row.get("action"))
        old_status = _export_clean(row.get("old_status"))
        new_status = _export_clean(row.get("new_status"))
        transition = " → ".join(x for x in [old_status, new_status] if x)
        comment = _export_clean(row.get("admin_comment"))
        item = " · ".join(x for x in [stamp, actor, action, transition, comment] if x)
        if item:
            parts.append(item)
    return "\n".join(parts)


def _write_matrix_sheet(
    workbook,
    strat_df,
    filtered_measures,
    monitoring_df,
    indicator_monitoring_df,
    selected_years,
    selected_quarters,
    selected_ssp_indices,
    selected_product_types,
    search_query,
):
    import pandas as pd
    from core.period_locks import is_period_locked
    from core.text_utils import strip_leading_code

    ws = workbook.add_worksheet("Матриця стратег. результатів")

    title_fmt = workbook.add_format({"bold": True, "font_color": _MATRIX_DARK_BLUE, "font_size": 16, "align": "left", "valign": "vcenter"})
    header_blue = workbook.add_format({"bold": True, "font_color": "#FFFFFF", "bg_color": _MATRIX_DARK_BLUE, "align": "center", "valign": "vcenter", "text_wrap": True, "border": 1, "border_color": "#FFFFFF"})
    header_orange = workbook.add_format({"bold": True, "font_color": "#FFFFFF", "bg_color": _MATRIX_DARK_ORANGE, "align": "center", "valign": "vcenter", "text_wrap": True, "border": 1, "border_color": "#FFFFFF"})
    number_fmt = workbook.add_format({"bold": True, "font_color": _MATRIX_TEXT, "align": "center", "valign": "vcenter", "border": 1, "border_color": _BORDER_COLOR})
    goal_fmt = workbook.add_format({"bg_color": _MATRIX_GOAL_BLUE, "font_color": "#FFFFFF", "bold": True, "valign": "top", "text_wrap": True, "border": 1, "border_color": _BORDER_COLOR})
    task_fmt = workbook.add_format({"bg_color": _MATRIX_TASK_BLUE, "font_color": _MATRIX_TEXT, "bold": True, "valign": "top", "text_wrap": True, "border": 1, "border_color": _BORDER_COLOR})
    measure_white = workbook.add_format({"bg_color": _MATRIX_WHITE, "font_color": _MATRIX_TEXT, "valign": "top", "text_wrap": True, "border": 1, "border_color": _BORDER_COLOR})
    measure_grey = workbook.add_format({"bg_color": _MATRIX_GREY, "font_color": _MATRIX_TEXT, "valign": "top", "text_wrap": True, "border": 1, "border_color": _BORDER_COLOR})

    selected_years = [int(y) for y in (selected_years or [])]
    selected_quarters = [_quarter_roman(q) for q in (selected_quarters or []) if _quarter_roman(q)]
    q_label = {"I": "Q1", "II": "Q2", "III": "Q3", "IV": "Q4"}

    # Структура колонок. A — службова порожня; нумерація починається з B=1.
    columns = [
        {"key": "_a", "kind": "blue"},
        {"key": "_group", "kind": "blue"},
        {"key": "_code", "kind": "blue"},
        {"key": "name", "kind": "blue", "single": "Стратегічні цілі, завдання та заходи до них"},
        {"key": "product_type", "kind": "blue", "single": "Тип продукту"},
        {"key": "indicator", "kind": "blue", "single": "Індикатор"},
        {"key": "unit", "kind": "blue", "single": "Одиниці виміру"},
        {"key": "base_2021", "kind": "orange", "year": "2021", "sub": "базовий рівень (факт)"},
        {"key": "fact_2024", "kind": "orange", "year": "2024", "sub": "звіт"},
        {"key": "fact_2025", "kind": "orange", "year": "2025", "sub": "факт"},
    ]
    for year in [2026, 2027, 2028]:
        columns.append({"key": f"target_{year}", "kind": "orange", "year": str(year), "sub": "цільовий орієнтир для заходів на рік"})
        if year in selected_years:
            for quarter in selected_quarters:
                columns.append({"key": f"q_{year}_{quarter}", "kind": "orange", "single": f"{year} {q_label.get(quarter, quarter)}", "quarter": (year, quarter)})

    # Роки 2029–2033 не мають окремих планових колонок у затвердженій матриці;
    # їхні обрані квартальні зрізи ставимо після планового блоку 2028.
    for year in [y for y in selected_years if y in {2029, 2030, 2031, 2032, 2033}]:
        for quarter in selected_quarters:
            columns.append({"key": f"q_{year}_{quarter}", "kind": "orange", "single": f"{year} {q_label.get(quarter, quarter)}", "quarter": (year, quarter)})

    columns.append({"key": "strategic_target_2028", "kind": "orange", "single": "Проміжний цільовий орієнтир на кінець 2028 року"})
    columns.append({"key": "strategic_target_2034", "kind": "orange", "single": "Цільовий орієнтир на кінець 2034 року для цілей і завдань (відповідає цілі, визначеній в НЕС-2030, ЦСР-2030 для показників, де це зазначено. Ціль перенесена на 2034 рік через «втрату» 4-х років — 2022–2025 внаслідок повномасштабної війни. Інші індикативні значення мають встановлюватися такими, що є кількісно узгодженими з цілями НЕС і ЦСР)"})
    if 2034 in selected_years:
        for quarter in selected_quarters:
            columns.append({"key": f"q_2034_{quarter}", "kind": "orange", "single": f"2034 {q_label.get(quarter, quarter)}", "quarter": (2034, quarter)})

    # Службові маркери для багаторівневої шапки.
    source_global_idx = len(columns)
    columns += [
        {"key": "source_global", "kind": "blue"},
        {"key": "source_national", "kind": "blue"},
    ]
    resp_idx = len(columns)
    columns += [
        {"key": "resp_main", "kind": "blue"},
        {"key": "resp_co", "kind": "blue"},
        {"key": "deputy_minister_raw", "kind": "blue", "single": "Заступник Міністра"},
        {"key": "measure_period_years", "kind": "blue", "single": "Період дії заходу в межах планового періоду, років"},
        {"key": "measure_start_date", "kind": "blue", "single": "Початкова дата"},
        {"key": "measure_end_date", "kind": "blue", "single": "Кінцева дата"},
    ]
    finance_idx = len(columns)
    columns += [
        {"key": "budget_kpkvk", "kind": "blue"},
        {"key": "budget_2026_approved", "kind": "blue"},
        {"key": "budget_2027_forecast", "kind": "blue"},
        {"key": "budget_2028_forecast", "kind": "blue"},
        {"key": "other_source", "kind": "blue"},
        {"key": "other_2026_plan", "kind": "blue"},
        {"key": "other_2027_forecast", "kind": "blue"},
        {"key": "other_2028_forecast", "kind": "blue"},
    ]

    last_col = len(columns) - 1
    ws.merge_range(0, 0, 0, 3, "Матриця стратегічних результатів", title_fmt)
    ws.set_row(0, 24)
    ws.set_row(1, 8)

    # Рядки 3–5 (xlsxwriter zero-based 2–4).
    for idx, col in enumerate(columns):
        fmt = header_orange if col.get("kind") == "orange" else header_blue
        if col.get("single"):
            ws.merge_range(2, idx, 4, idx, col["single"], fmt)
        elif col.get("year"):
            ws.write(2, idx, col["year"], fmt)
            ws.merge_range(3, idx, 4, idx, col["sub"], fmt)
        elif idx < 3:
            ws.merge_range(2, idx, 4, idx, "", fmt)

    source_title = "Джерело даних (для індикаторів)/Підстава для включення (для заходів)"
    ws.merge_range(2, source_global_idx, 3, source_global_idx, source_title, header_blue)
    ws.write(4, source_global_idx, "Глобальний рівень", header_blue)
    ws.merge_range(2, source_global_idx + 1, 3, source_global_idx + 1, source_title, header_blue)
    ws.write(4, source_global_idx + 1, "Національний рівень", header_blue)

    ws.merge_range(2, resp_idx, 3, resp_idx + 1, "Відповідальні самостійні структурні підрозділи", header_blue)
    ws.write(4, resp_idx, "Головний виконавець", header_blue)
    ws.write(4, resp_idx + 1, "Співвиконавець", header_blue)

    ws.merge_range(2, finance_idx, 2, finance_idx + 7, "Фінансування", header_blue)
    ws.merge_range(3, finance_idx, 3, finance_idx + 3, "Державний бюджет України", header_blue)
    ws.merge_range(3, finance_idx + 4, 3, finance_idx + 7, "Інші джерела", header_blue)
    finance_headers = [
        "КПКВК", "2026 затверджено, млрд грн", "2027 прогноз, млрд грн", "2028 прогноз, млрд грн",
        "Джерело (МТД, кошти партнерів, інші небюджетні джерела)", "2026 план", "2027 прогноз", "2028 прогноз",
    ]
    for offset, label in enumerate(finance_headers):
        ws.write(4, finance_idx + offset, label, header_blue)

    # Нумерація рядка 6: B=1, A порожня.
    ws.write(5, 0, "", number_fmt)
    for idx in range(1, last_col + 1):
        ws.write(5, idx, idx, number_fmt)

    ws.freeze_panes(6, 0)
    ws.set_row(2, 42)
    ws.set_row(3, 42)
    ws.set_row(4, 58)

    # Ширини.
    widths = {0: 3, 1: 18, 2: 20, 3: 48, 4: 22, 5: 42, 6: 18}
    for idx in range(len(columns)):
        width = widths.get(idx, 18)
        key = columns[idx]["key"]
        if key.startswith("q_"):
            width = 14
        elif key in {"source_global", "source_national"}:
            width = 34
        elif key in {"resp_main", "resp_co", "deputy_minister_raw"}:
            width = 24
        elif key in {"other_source"}:
            width = 34
        ws.set_column(idx, idx, width)

    record_index = _latest_period_records(monitoring_df)

    measures = filtered_measures.copy() if filtered_measures is not None else pd.DataFrame()
    if measures.empty:
        return ws
    goal_codes = list(dict.fromkeys(measures["parent_goal_code"].astype(str).str.strip().tolist()))
    task_codes = list(dict.fromkeys(measures["parent_task_code"].astype(str).str.strip().tolist()))
    goals = strat_df[(strat_df["object_type"] == "goal") & strat_df["code"].astype(str).str.strip().isin(goal_codes)].copy()
    tasks = strat_df[(strat_df["object_type"] == "task") & strat_df["code"].astype(str).str.strip().isin(task_codes)].copy()

    def indicator_rows_for(kind: str, code: str):
        if kind == "goal":
            rows = strat_df[
                (strat_df["object_type"].isin(["goal", "goal_indicator"]))
                & (strat_df["parent_goal_code"].astype(str).str.strip().eq(code) | strat_df["code"].astype(str).str.strip().eq(code))
                & (strat_df["parent_task_code"].astype(str).str.strip() == "")
            ].copy()
        else:
            rows = strat_df[
                (strat_df["object_type"].isin(["task", "task_indicator"]))
                & (strat_df["parent_task_code"].astype(str).str.strip().eq(code) | strat_df["code"].astype(str).str.strip().eq(code))
            ].copy()
        return [
            row for _, row in rows.iterrows()
            if _export_clean(row.get("indicator"))
            and _indicator_matches_export_filters(
                row, selected_ssp_indices, selected_product_types, search_query
            )
        ]

    def row_values(source, row_type: str, group="", code_value=""):
        values = {col["key"]: "" for col in columns}
        values["_group"] = group
        values["_code"] = code_value
        if source is not None:
            for col in columns:
                key = col["key"]
                if key == "resp_co":
                    values[key] = "; ".join(x for x in [_export_clean(source.get("resp_co_1")), _export_clean(source.get("resp_co_2"))] if x)
                elif key.startswith("q_"):
                    year, quarter = col["quarter"]
                    if row_type == "measure":
                        record = record_index.get((_export_clean(source.get("code")), str(year), quarter))
                        actual = _export_clean(record.get("numeric_value")) if record is not None else ""
                        if is_period_locked(year, quarter) and not actual:
                            actual = "Не настав час"
                        values[key] = actual
                elif key in source:
                    values[key] = _export_clean(source.get(key))
            if row_type == "indicator":
                for key, year in [("base_2021", 2021), ("fact_2024", 2024), ("fact_2025", 2025), ("target_2026", 2026), ("target_2027", 2027), ("target_2028", 2028)]:
                    current = _export_clean(values.get(key))
                    if current.lower() in {"x", "х"}:
                        latest = _indicator_latest_value(indicator_monitoring_df, source, year)
                        if latest:
                            values[key] = latest
        if row_type == "indicator":
            values["name"] = ""
        return [values[col["key"]] for col in columns]

    out_row = 6
    zebra = 0
    for _, goal in goals.iterrows():
        goal_code = _export_clean(goal.get("code"))
        goal_measures = measures[measures["parent_goal_code"].astype(str).str.strip() == goal_code]
        if goal_measures.empty:
            continue
        goal_name = strip_leading_code(goal.get("name", ""), goal_code)
        values = row_values(None, "goal", f"Стратегічна ціль {_goal_number(goal_code)}", goal_name)
        for col_idx, value in enumerate(values):
            ws.write(out_row, col_idx, value, goal_fmt)
        out_row += 1
        for indicator_row in indicator_rows_for("goal", goal_code):
            values = row_values(indicator_row, "indicator")
            for col_idx, value in enumerate(values):
                ws.write(out_row, col_idx, value, goal_fmt)
            out_row += 1

        goal_task_codes = list(dict.fromkeys(goal_measures["parent_task_code"].astype(str).str.strip().tolist()))
        goal_tasks = tasks[tasks["code"].astype(str).str.strip().isin(goal_task_codes)]
        for _, task in goal_tasks.iterrows():
            task_code = _export_clean(task.get("code"))
            task_name = strip_leading_code(task.get("name", ""), task_code)
            values = row_values(None, "task", "Завдання:", "")
            # За ТЗ назва завдання — у колонці D.
            values[3] = task_name
            for col_idx, value in enumerate(values):
                ws.write(out_row, col_idx, value, task_fmt)
            out_row += 1
            for indicator_row in indicator_rows_for("task", task_code):
                values = row_values(indicator_row, "indicator")
                for col_idx, value in enumerate(values):
                    ws.write(out_row, col_idx, value, task_fmt)
                out_row += 1

            task_measures = goal_measures[goal_measures["parent_task_code"].astype(str).str.strip() == task_code]
            for measure_idx, (_, measure) in enumerate(task_measures.iterrows()):
                code = _export_clean(measure.get("code"))
                values = row_values(measure, "measure", "Заходи:" if measure_idx == 0 else "", code)
                values[3] = strip_leading_code(measure.get("name", ""), code)
                fmt = measure_grey if zebra % 2 else measure_white
                zebra += 1
                for col_idx, value in enumerate(values):
                    ws.write(out_row, col_idx, value, fmt)
                out_row += 1

    return ws


def _build_detailed_export_df(filtered_measures, monitoring_df, selected_years, selected_quarters, manual_closeouts):
    import pandas as pd
    from core.period_locks import is_period_locked
    from core.text_utils import strip_leading_code

    q_label = {"I": "Q1", "II": "Q2", "III": "Q3", "IV": "Q4"}
    records = _latest_period_records(monitoring_df)
    logs_df = _load_export_logs()
    manual_closeouts = set(manual_closeouts or set())
    rows = []
    for _, measure in filtered_measures.iterrows():
        code = _export_clean(measure.get("code"))
        row = {
            "Код": code,
            "Захід": strip_leading_code(measure.get("name", ""), code),
            "Тип продукту": _export_clean(measure.get("product_type")),
            "Індикатор": _export_clean(measure.get("indicator")),
            "Одиниці виміру": _export_clean(measure.get("unit")),
            "Головний виконавець": _export_clean(measure.get("resp_main")),
            "Співвиконавець": "; ".join(x for x in [_export_clean(measure.get("resp_co_1")), _export_clean(measure.get("resp_co_2"))] if x),
        }
        for year in selected_years or []:
            for quarter in selected_quarters or []:
                q = _quarter_roman(quarter)
                prefix = f"{year} {q_label.get(q, q)}"
                record = records.get((code, str(year), q))
                locked = is_period_locked(year, q)
                actual = _export_clean(record.get("numeric_value")) if record is not None else ""
                if locked and not actual:
                    actual = "—"
                row[prefix] = actual
                if locked:
                    status = "Не настав час"
                elif record is not None:
                    status = _export_clean(record.get("status"))
                elif (code, str(year), q) in manual_closeouts:
                    status = "Виконано"
                else:
                    status = ""
                record_logs = (
                    _request_export_logs(logs_df, record.get("id"))
                    if record is not None else None
                )
                coordinator, scheme_history, current = _approval_route_and_current(
                    record,
                    record_logs,
                )
                row[f"{prefix} — Статус заходу"] = status
                row[f"{prefix} — Відповідальна особа від ССП"] = _export_clean(record.get("responsible_person")) if record is not None else ""
                row[f"{prefix} — Координатор"] = coordinator
                row[f"{prefix} — Опис прогресу виконання"] = _export_clean(record.get("progress_text")) if record is not None else ""
                row[f"{prefix} — Ризики/проблеми/відхилення"] = _export_clean(record.get("risks")) if record is not None else ""
                row[f"{prefix} — Посилання на НПА"] = _export_clean(record.get("npa_link")) if record is not None else ""
                row[f"{prefix} — Діюча схема погодження"] = scheme_history
                row[f"{prefix} — Поточний статус погодження"] = current
        rows.append(row)
    return pd.DataFrame(rows)


def build_main_monitoring_export(
    *,
    strat_df,
    filtered_measures,
    monitoring_df,
    indicator_monitoring_df,
    selected_years,
    selected_quarters,
    selected_ssp_indices=None,
    selected_product_types=None,
    search_query="",
    manual_closeouts=None,
) -> bytes:
    """Єдиний двовкладковий Excel головної сторінки за застосованими фільтрами."""
    import io
    import pandas as pd

    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        workbook = writer.book
        matrix_ws = _write_matrix_sheet(
            workbook,
            strat_df,
            filtered_measures,
            monitoring_df,
            indicator_monitoring_df,
            selected_years,
            selected_quarters,
            selected_ssp_indices,
            selected_product_types,
            search_query,
        )
        writer.sheets["Матриця стратег. результатів"] = matrix_ws

        detail_df = _build_detailed_export_df(
            filtered_measures,
            monitoring_df,
            selected_years,
            selected_quarters,
            manual_closeouts,
        )
        sheet_name = "Детальні моніторингові дані"
        ws = workbook.add_worksheet(sheet_name)
        writer.sheets[sheet_name] = ws
        header_fmt = workbook.add_format({
            "bold": True, "font_color": _HEADER_FG, "bg_color": _HEADER_BG,
            "align": "center", "valign": "vcenter", "text_wrap": True,
            "border": 1, "border_color": "#FFFFFF",
        })
        cell_fmt = workbook.add_format({"valign": "top", "text_wrap": True, "border": 1, "border_color": _BORDER_COLOR})
        band_fmt = workbook.add_format({"valign": "top", "text_wrap": True, "border": 1, "border_color": _BORDER_COLOR, "bg_color": _BAND_BG})
        for col_idx, col_name in enumerate(detail_df.columns):
            ws.write(0, col_idx, col_name, header_fmt)
            width = min(max(len(str(col_name)) + 2, 14), 42)
            if "Опис прогресу" in str(col_name) or "Ризики" in str(col_name) or "Діюча схема" in str(col_name):
                width = 45
            ws.set_column(col_idx, col_idx, width)
        for row_idx in range(len(detail_df)):
            fmt = band_fmt if row_idx % 2 else cell_fmt
            for col_idx, col_name in enumerate(detail_df.columns):
                value = detail_df.iloc[row_idx][col_name]
                if value is None or (isinstance(value, float) and pd.isna(value)):
                    value = ""
                ws.write(row_idx + 1, col_idx, value, fmt)
        if len(detail_df.columns):
            ws.freeze_panes(1, 2)
            ws.autofilter(0, 0, len(detail_df), len(detail_df.columns) - 1)
            ws.set_row(0, 42)

    return buffer.getvalue()
