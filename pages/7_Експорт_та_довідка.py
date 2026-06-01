import streamlit as st
import pandas as pd
import plotly.express as px
from supabase import create_client
from datetime import datetime
from io import BytesIO
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


st.set_page_config(page_title="Експорт та довідка", layout="wide")

FILE_PATH = "Під моніторинг СП.xlsx"
SHEET_NAME = "Страт_матриця"

supabase = create_client(
    st.secrets["SUPABASE_URL"],
    st.secrets["SUPABASE_KEY"]
)

st.markdown("""
<style>
.stApp {
    background:
        radial-gradient(circle at top right, rgba(37,99,235,0.08), transparent 28%),
        radial-gradient(circle at bottom left, rgba(22,163,74,0.07), transparent 30%),
        linear-gradient(180deg, #f6f8fb 0%, #eef2f7 100%);
}

.stApp::before {
    content: "";
    position: fixed;
    top: -160px;
    right: -120px;
    width: 460px;
    height: 460px;
    border-radius: 50%;
    background: rgba(37, 99, 235, 0.045);
    z-index: 0;
}

.stApp::after {
    content: "";
    position: fixed;
    bottom: -180px;
    left: -120px;
    width: 390px;
    height: 390px;
    border-radius: 50%;
    background: rgba(22, 163, 74, 0.045);
    z-index: 0;
}

.main .block-container {
    max-width: 1550px;
    padding-top: 1.2rem;
    position: relative;
    z-index: 1;
}

.ua-line {
    height: 7px;
    border-radius: 999px;
    background: linear-gradient(90deg, #005BBB 0%, #005BBB 50%, #FFD500 50%, #FFD500 100%);
    margin-bottom: 14px;
}

.ministry-label {
    text-align: right;
    color: #475569;
    font-size: 14px;
    font-weight: 700;
    margin-bottom: 8px;
}

.header-box, .card {
    background: rgba(255,255,255,0.94);
    border: 1px solid #d8dee9;
    border-radius: 16px;
    padding: 22px 26px;
    margin-bottom: 18px;
    box-shadow: 0 8px 24px rgba(15,23,42,0.06);
}

.header-title {
    font-size: 32px;
    font-weight: 900;
    color: #0f172a;
    margin-bottom: 8px;
}

.header-subtitle, .card-subtitle {
    font-size: 15px;
    color: #475569;
    line-height: 1.55;
}

.card-title {
    font-size: 21px;
    font-weight: 900;
    color: #0f172a;
    margin-bottom: 8px;
}

.badge-wrap {
    display: flex;
    flex-wrap: wrap;
    gap: 8px;
    margin: 12px 0;
}

.badge {
    background: #eef6ff;
    border: 1px solid #bfdbfe;
    color: #1d4ed8;
    border-radius: 999px;
    padding: 7px 11px;
    font-size: 13px;
    font-weight: 800;
}

.badge-green {
    background: #dcfce7;
    border: 1px solid #bbf7d0;
    color: #166534;
}

.badge-yellow {
    background: #fef9c3;
    border: 1px solid #fde68a;
    color: #854d0e;
}

.badge-red {
    background: #fee2e2;
    border: 1px solid #fecaca;
    color: #991b1b;
}

.report-box {
    background: #ffffff;
    border: 1px solid #d8dee9;
    border-left: 7px solid #2563eb;
    border-radius: 16px;
    padding: 24px 28px;
    margin: 18px 0;
    box-shadow: 0 8px 24px rgba(15,23,42,0.055);
}

.report-title {
    font-size: 24px;
    font-weight: 900;
    color: #0f172a;
    margin-bottom: 12px;
}

.report-text {
    font-size: 15px;
    line-height: 1.75;
    color: #334155;
    text-align: justify;
}

.alert-grid {
    display: grid;
    grid-template-columns: repeat(4, 1fr);
    gap: 14px;
    margin: 14px 0;
}

.alert-card {
    border-radius: 14px;
    padding: 16px;
    border: 1px solid #d8dee9;
    background: white;
}

.alert-title {
    font-size: 13px;
    color: #64748b;
    font-weight: 800;
}

.alert-value {
    font-size: 28px;
    color: #0f172a;
    font-weight: 950;
    margin-top: 4px;
}

.alert-red {
    background: #fee2e2;
    border-color: #fecaca;
}

.alert-yellow {
    background: #fef9c3;
    border-color: #fde68a;
}

.alert-green {
    background: #dcfce7;
    border-color: #bbf7d0;
}

.alert-blue {
    background: #dbeafe;
    border-color: #bfdbfe;
}

div[data-testid="stMetric"] {
    background: rgba(255,255,255,0.9);
    border: 1px solid #d8dee9;
    border-radius: 14px;
    padding: 14px 16px;
    box-shadow: 0 4px 12px rgba(15,23,42,0.04);
}

div.stDownloadButton > button {
    border-radius: 12px;
    padding: 12px 18px;
    font-weight: 850;
}

.footer {
    text-align: center;
    color: #64748b;
    font-size: 13px;
    margin-top: 50px;
    padding: 22px 0 12px 0;
    border-top: 1px solid #d8dee9;
}
</style>
""", unsafe_allow_html=True)


def clean(value):
    if value is None or pd.isna(value) or str(value) == "None":
        return ""
    return str(value)


def parse_period(value):
    text = str(value).lower().strip()

    if text in ["", "nan", "none", "н.д."]:
        return None

    q = None
    year = None

    if "1 квартал" in text or "i квартал" in text:
        q = 1
    elif "2 квартал" in text or "ii квартал" in text:
        q = 2
    elif "3 квартал" in text or "iii квартал" in text:
        q = 3
    elif "4 квартал" in text or "iv квартал" in text:
        q = 4

    year_match = re.search(r"20\d{2}", text)

    if year_match:
        year = int(year_match.group())

    if year and q:
        return year * 10 + q

    return None


def quarter_to_number(q):
    return {"I": 1, "II": 2, "III": 3, "IV": 4}.get(str(q), 1)


def get_goal_code(code):
    parts = str(code).split(".")
    return parts[0] + "." if parts else ""


def get_task_code(code):
    parts = str(code).split(".")
    if len(parts) >= 2:
        return f"{parts[0]}.{parts[1]}."
    return ""


def to_number(value):
    text = str(value).replace(",", ".").strip()

    if text.lower() in ["", "nan", "none", "н.д.", "x", "х", "так", "ні", "да", "нет"]:
        return None

    match = re.search(r"-?\d+(\.\d+)?", text)

    if match:
        try:
            return float(match.group())
        except Exception:
            return None

    return None


def normalize_text(value):
    return str(value).strip().lower().replace("і", "i")


def status_score(status):
    status = str(status)

    if status == "Виконано":
        return 100
    if status == "Виконано частково":
        return 50
    if status == "Виконується":
        return 40
    if status == "Потребує уваги":
        return 25
    if status == "Прострочено":
        return 0
    if status == "Не розпочато":
        return 0
    if status == "Не подано":
        return 0

    return 0


def plan_fact_percent(actual, target):
    actual_num = to_number(actual)
    target_num = to_number(target)

    actual_text = normalize_text(actual)
    target_text = normalize_text(target)

    if actual_num is not None and target_num is not None and target_num != 0:
        return round(min((actual_num / target_num) * 100, 150), 1)

    if target_text in ["так", "yes"] or actual_text in ["так", "ні", "yes", "no"]:
        if actual_text in ["так", "yes"]:
            return 100
        if actual_text in ["ні", "no"]:
            return 0

    return None


def traffic_light(score):
    if score >= 70:
        return "У графіку"
    if score >= 35:
        return "Потребує уваги"
    return "Відстає"


def risk_score_calc(row, selected_quarter_num, selected_period_num):
    status = str(row.get("status", "Не подано"))
    actual = row.get("numeric_value", "")
    target = row.get("selected_target", "")
    risks_text = str(row.get("risks", "")).strip()
    end_num = row.get("end_num", None)

    score = 0
    reasons = []

    if status == "Не подано":
        score += 45
        reasons.append("за активним заходом не подано дані")

    if status in ["Прострочено", "Потребує уваги"]:
        score += 35
        reasons.append("проблемний статус виконання")

    if risks_text:
        score += 20
        reasons.append("у заявці зазначені ризики або відхилення")

    pf = plan_fact_percent(actual, target)
    expected_progress = selected_quarter_num * 25

    if pf is not None:
        if pf < max(expected_progress - 25, 0):
            score += 35
            reasons.append("факт суттєво нижчий за очікуваний прогрес")
        elif pf < expected_progress:
            score += 20
            reasons.append("факт нижчий за очікуваний прогрес")

    if pd.notna(end_num) and end_num < selected_period_num and status != "Виконано":
        score += 45
        reasons.append("строк виконання минув, захід не виконано")

    score = min(score, 100)

    if not reasons:
        reasons.append("критичних відхилень не виявлено")

    return score, "; ".join(reasons)


def forecast_to_q4(current_completion, selected_quarter):
    qn = quarter_to_number(selected_quarter)
    if qn == 0:
        return current_completion

    forecast = current_completion / qn * 4
    return round(min(forecast, 100), 1)


@st.cache_data
def load_strat_matrix():
    df = pd.read_excel(FILE_PATH, sheet_name=SHEET_NAME, header=None, engine="openpyxl")
    data = df.iloc[7:].copy()

    result = pd.DataFrame({
        "type_marker": data.iloc[:, 1],
        "code": data.iloc[:, 2],
        "name": data.iloc[:, 3],
        "indicator": data.iloc[:, 5],
        "unit": data.iloc[:, 6],
        "base_2021": data.iloc[:, 7],
        "fact_2024": data.iloc[:, 8],
        "expected_2025": data.iloc[:, 9],
        "target_2026": data.iloc[:, 10],
        "target_2027": data.iloc[:, 11],
        "target_2028": data.iloc[:, 12],
        "department": data.iloc[:, 17],
        "start_period": data.iloc[:, 22],
        "end_period": data.iloc[:, 23],
    })

    result = result.dropna(subset=["code"])
    result["code"] = result["code"].astype(str).str.strip()
    result["type_marker"] = result["type_marker"].astype(str).str.strip()

    def classify(row):
        marker = str(row["type_marker"]).lower()
        code = str(row["code"]).strip()

        if "стратегічна ціль" in marker:
            return "goal"
        if "завдання" in marker:
            return "task"
        if code.count(".") >= 3:
            return "measure"
        return "other"

    result["object_type"] = result.apply(classify, axis=1)
    return result


def load_requests():
    response = supabase.table("monitoring_requests").select("*").execute()
    return pd.DataFrame(response.data or [])


def prepare_period_data(strat_df, requests_df, year, quarter, department):
    measures = strat_df[strat_df["object_type"] == "measure"].copy()
    goals = strat_df[strat_df["object_type"] == "goal"].copy()

    measures["goal_code"] = measures["code"].apply(get_goal_code)
    measures["task_code"] = measures["code"].apply(get_task_code)
    measures["strategic_goal"] = measures["goal_code"].map(goals.set_index("code")["name"].to_dict())

    measures["start_num"] = measures["start_period"].apply(parse_period)
    measures["end_num"] = measures["end_period"].apply(parse_period)

    selected_q_num = quarter_to_number(quarter)
    selected_period_num = int(year) * 10 + selected_q_num

    active = measures[
        (
            measures["start_num"].isna()
            |
            (measures["start_num"] <= selected_period_num)
        )
        &
        (
            measures["end_num"].isna()
            |
            (measures["end_num"] >= selected_period_num)
        )
    ].copy()

    if department != "Усі":
        active = active[active["department"].astype(str) == str(department)]

    if requests_df.empty:
        requests_df = pd.DataFrame(columns=[
            "year", "quarter", "department", "strat_code", "status", "numeric_value",
            "risks", "progress_text", "approval_status", "submitted_at",
            "responsible_person", "phone", "email", "file_names", "file_urls",
            "admin_comment"
        ])

    required = [
        "year", "quarter", "department", "strat_code", "status", "numeric_value",
        "risks", "progress_text", "approval_status", "submitted_at",
        "responsible_person", "phone", "email", "file_names", "file_urls",
        "admin_comment"
    ]

    for col in required:
        if col not in requests_df.columns:
            requests_df[col] = ""

    approved_requests = requests_df[
        requests_df["approval_status"].astype(str) == "Погоджено"
    ].copy()

    if approved_requests.empty:
        period_requests = pd.DataFrame(columns=[
            "strat_code", "status", "numeric_value", "risks", "progress_text",
            "submitted_at", "responsible_person", "phone", "email", "file_names", "file_urls"
        ])
    else:
        period_requests = approved_requests[
            (approved_requests["year"].astype(str) == str(year)) &
            (approved_requests["quarter"].astype(str) == str(quarter))
        ].copy()

        if department != "Усі":
            period_requests = period_requests[period_requests["department"].astype(str) == str(department)]

        if not period_requests.empty:
            period_requests = (
                period_requests
                .sort_values("submitted_at")
                .groupby("strat_code")
                .tail(1)
            )

    active = active.merge(
        period_requests[
            [
                "strat_code", "status", "numeric_value", "risks", "progress_text",
                "submitted_at", "responsible_person", "phone", "email", "file_names", "file_urls"
            ]
        ],
        left_on="code",
        right_on="strat_code",
        how="left"
    )

    active["status"] = active["status"].fillna("Не подано")
    active["numeric_value"] = active["numeric_value"].fillna("")
    active["risks"] = active["risks"].fillna("")
    active["progress_text"] = active["progress_text"].fillna("")
    active["selected_target"] = active[f"target_{year}"] if f"target_{year}" in active.columns else ""

    active["status_score"] = active["status"].apply(status_score)
    active["plan_fact_percent"] = active.apply(
        lambda r: plan_fact_percent(r["numeric_value"], r["selected_target"]),
        axis=1
    )

    active["performance_score"] = active.apply(
        lambda r: r["plan_fact_percent"] if pd.notna(r["plan_fact_percent"]) else r["status_score"],
        axis=1
    )

    risk_results = active.apply(
        lambda r: risk_score_calc(r, selected_q_num, selected_period_num),
        axis=1
    )

    active["risk_score"] = [x[0] for x in risk_results]
    active["risk_reason"] = [x[1] for x in risk_results]
    active["auto_risk"] = active["risk_score"].apply(
        lambda x: "Критичний ризик" if x >= 70 else ("Середній ризик" if x >= 35 else "Низький ризик")
    )
    active["traffic_light"] = active["performance_score"].apply(traffic_light)

    return active


def generate_analytical_text(active, requests_df, year, quarter, department):
    total_active = len(active)
    submitted_count = len(active[active["status"] != "Не подано"])
    coverage = round(submitted_count / total_active * 100, 1) if total_active else 0
    completion = round(active["performance_score"].mean(), 1) if total_active else 0
    risk_count = len(active[active["auto_risk"].isin(["Критичний ризик", "Середній ризик"])])
    critical_count = len(active[active["auto_risk"] == "Критичний ризик"])
    low_risk_count = len(active[active["auto_risk"] == "Низький ризик"])
    without_data = len(active[active["status"] == "Не подано"])
    forecast_q4 = forecast_to_q4(completion, quarter)

    goal_progress = (
        active
        .groupby(["goal_code", "strategic_goal"])
        .agg(
            active_measures=("code", "count"),
            completion=("performance_score", "mean"),
            risk_score=("risk_score", "mean"),
            risk_count=("auto_risk", lambda x: x.isin(["Критичний ризик", "Середній ризик"]).sum()),
        )
        .reset_index()
    )

    dep_progress = (
        active
        .groupby("department")
        .agg(
            active_measures=("code", "count"),
            completion=("performance_score", "mean"),
            risk_score=("risk_score", "mean"),
            risk_count=("auto_risk", lambda x: x.isin(["Критичний ризик", "Середній ризик"]).sum()),
        )
        .reset_index()
    )

    best_goal_text = "н/д"
    worst_goal_text = "н/д"

    if not goal_progress.empty:
        best_goal = goal_progress.sort_values("completion", ascending=False).iloc[0]
        worst_goal = goal_progress.sort_values("completion", ascending=True).iloc[0]
        best_goal_text = f"СЦ {best_goal['goal_code']} — {round(best_goal['completion'], 1)}%"
        worst_goal_text = f"СЦ {worst_goal['goal_code']} — {round(worst_goal['completion'], 1)}%"

    best_dep_text = "н/д"
    worst_dep_text = "н/д"

    if not dep_progress.empty:
        best_dep = dep_progress.sort_values("completion", ascending=False).iloc[0]
        worst_dep = dep_progress.sort_values("completion", ascending=True).iloc[0]
        best_dep_text = f"{best_dep['department']} — {round(best_dep['completion'], 1)}%"
        worst_dep_text = f"{worst_dep['department']} — {round(worst_dep['completion'], 1)}%"

    if completion >= 70 and risk_count / max(total_active, 1) <= 0.2:
        general_assessment = (
            "Загальний стан реалізації стратегічного плану за обраний період можна оцінити як відносно контрольований. "
            "Більшість активних заходів демонструє прийнятний рівень виконання, а частка ризикових позицій не є критичною для портфеля в цілому."
        )
    elif completion >= 40:
        general_assessment = (
            "Загальний стан реалізації стратегічного плану характеризується наявністю помірних відхилень. "
            "Частина заходів перебуває в межах очікуваного темпу, однак за окремими напрямами спостерігається недостатній рівень подання даних або ризик недосягнення планових значень."
        )
    else:
        general_assessment = (
            "Загальний стан реалізації стратегічного плану свідчить про підвищений ризик недовиконання. "
            "Низький рівень фактичного прогресу, неповне покриття моніторингом або наявність значної кількості ризикових заходів потребують управлінської уваги."
        )

    if coverage >= 80:
        coverage_assessment = (
            "Рівень покриття моніторингом є достатнім для формування узагальнених висновків. "
            "Подані дані дозволяють оцінювати стан виконання не лише за окремими заходами, а й на рівні стратегічних цілей та відповідальних структурних підрозділів."
        )
    elif coverage >= 40:
        coverage_assessment = (
            "Рівень покриття моніторингом є частковим. "
            "Це дозволяє сформувати попередню оцінку стану виконання, однак окремі висновки потребують обережної інтерпретації через наявність заходів без поданих даних."
        )
    else:
        coverage_assessment = (
            "Рівень покриття моніторингом є недостатнім. "
            "Значна кількість активних заходів не має поданих квартальних даних, що обмежує якість аналітичної оцінки та підвищує невизначеність щодо реального стану виконання."
        )

    risk_share = round(risk_count / total_active * 100, 1) if total_active else 0

    if risk_share >= 40:
        risk_assessment = (
            "Ризиковий профіль портфеля є підвищеним. "
            "Висока частка заходів із середнім або критичним ризиком може свідчити про системні проблеми у виконанні, недостатню динаміку прогресу або слабке документальне підтвердження результатів."
        )
    elif risk_share >= 20:
        risk_assessment = (
            "Ризиковий профіль є помірним. "
            "Проблемні заходи не домінують у загальній структурі, але потребують регулярного контролю та точкової комунікації з відповідальними департаментами."
        )
    else:
        risk_assessment = (
            "Ризиковий профіль є відносно низьким. "
            "Більшість активних заходів не має ознак суттєвого ризику недосягнення, хоча окремі позиції потребують подальшого спостереження."
        )

    text = f"""
За результатами автоматизованого аналізу виконання стратегічного плану за {quarter} квартал {year} року сформовано узагальнену аналітичну оцінку. Обраний зріз охоплює: {department if department != "Усі" else "усі відповідальні структурні підрозділи"}. У межах обраного періоду активними визначено {total_active} заходів, з яких за {submitted_count} заходами подано моніторингову інформацію. Рівень покриття моніторингом становить {coverage}%.

Середній розрахунковий рівень виконання стратегічного плану за обраний квартал становить {completion}%. Прогнозний рівень виконання до IV кварталу за поточним темпом оцінюється на рівні {forecast_q4}%. {general_assessment}

У структурі активних заходів {low_risk_count} заходів мають низький рівень ризику, {risk_count} заходів віднесено до середнього або критичного ризику, з них {critical_count} заходів мають критичний ризик. Кількість заходів без поданих даних становить {without_data}. {risk_assessment}

З погляду повноти інформаційної бази слід зазначити, що подані квартальні дані є ключовою умовою для коректного оцінювання виконання. {coverage_assessment}

У розрізі стратегічних цілей найкращий рівень виконання зафіксовано за напрямом {best_goal_text}. Найнижчий рівень виконання спостерігається за напрямом {worst_goal_text}. Такий розподіл свідчить про нерівномірність темпів реалізації між різними блоками стратегічного плану та потребує подальшого аналізу причин відхилень.

У розрізі відповідальних підрозділів найвищий рівень виконання демонструє {best_dep_text}, тоді як найнижчий показник зафіксовано у {worst_dep_text}. Ці результати можуть використовуватися для визначення пріоритетів управлінської комунікації, уточнення причин затримок та формування переліку питань для доопрацювання з відповідальними виконавцями.

Загалом система моніторингу дозволяє відстежувати не лише факт подання інформації, а й якість виконання, наявність ризиків, відхилення від очікуваного квартального прогресу та потенційні проблеми з досягненням планових значень. За результатами аналізу доцільно зосередити увагу на заходах із критичним ризиком, заходах без поданих даних, а також на стратегічних цілях і департаментах із найнижчим рівнем виконання.
""".strip()

    return text


def create_excel_report(active, requests_df, risk_table, goal_progress, dep_progress, year, quarter, department):
    output = BytesIO()

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        active.to_excel(writer, sheet_name="Активні заходи", index=False)
        risk_table.to_excel(writer, sheet_name="Ризики", index=False)
        requests_df.to_excel(writer, sheet_name="Реєстр заявок", index=False)
        goal_progress.to_excel(writer, sheet_name="СЦ", index=False)
        dep_progress.to_excel(writer, sheet_name="Департаменти", index=False)

        workbook = writer.book

        header_format = workbook.add_format({
            "bold": True,
            "bg_color": "#D9EAF7",
            "border": 1
        })

        for sheet_name in writer.sheets:
            worksheet = writer.sheets[sheet_name]
            worksheet.freeze_panes(1, 0)

            if sheet_name == "Активні заходи":
                df_sheet = active
            elif sheet_name == "Ризики":
                df_sheet = risk_table
            elif sheet_name == "Реєстр заявок":
                df_sheet = requests_df
            elif sheet_name == "СЦ":
                df_sheet = goal_progress
            else:
                df_sheet = dep_progress

            for col_num, value in enumerate(df_sheet.columns.values):
                worksheet.write(0, col_num, value, header_format)
                worksheet.set_column(col_num, col_num, 18)

        summary = workbook.add_worksheet("Пояснення")
        summary.write(0, 0, "Звіт")
        summary.write(0, 1, f"{year} рік, {quarter} квартал")
        summary.write(1, 0, "Департамент")
        summary.write(1, 1, department)
        summary.write(2, 0, "Дата формування")
        summary.write(2, 1, datetime.now().strftime("%d.%m.%Y %H:%M"))

    output.seek(0)
    return output


def create_docx_report(text, metrics, year, quarter, department):
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("АНАЛІТИЧНА ДОВІДКА")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"щодо стану виконання стратегічного плану за {quarter} квартал {year} року")
    run.font.size = Pt(12)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Зріз: {department}").italic = True

    document.add_paragraph("Ключові показники").runs[0].bold = True

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Показник"
    hdr_cells[1].text = "Значення"

    for key, value in metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)

    document.add_paragraph("Аналітичний висновок").runs[0].bold = True

    for paragraph in text.split("\n\n"):
        p = document.add_paragraph(paragraph.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.size = Pt(11)

    document.add_paragraph("Сформовано автоматизованою системою моніторингу стратегічного плану.").italic = True
    document.add_paragraph("Розроблено департаментом стратегічного планування та макроекономічного прогнозування.").italic = True

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


st.markdown('<div class="ua-line"></div>', unsafe_allow_html=True)

st.markdown("""
<div class="ministry-label">
🇺🇦 Міністерство економіки, довкілля та сільського господарства України
</div>
""", unsafe_allow_html=True)

st.markdown(f"""
<div class="header-box">
    <div class="header-title">Експорт звітів та автоматична аналітична довідка</div>
    <div class="header-subtitle">
        Сторінка формує управлінські матеріали на основі даних стратегічного плану,
        поданих заявок, погоджених значень, автоматичних ризиків і квартального прогресу.
        Система самостійно генерує текст довідки та таблиці для експорту.
    </div>
    <div class="badge-wrap">
        <div class="badge">● Excel-звіт</div>
        <div class="badge">● Реєстр заявок</div>
        <div class="badge">● Таблиця ризиків</div>
        <div class="badge">● Аналітична довідка DOCX</div>
        <div class="badge">● Оновлено: {datetime.now().strftime("%d.%m.%Y %H:%M")}</div>
    </div>
</div>
""", unsafe_allow_html=True)

strat_df = load_strat_matrix()
requests_df = load_requests()

if requests_df.empty:
    requests_df = pd.DataFrame(columns=[
        "year", "quarter", "department", "strat_code", "status", "numeric_value",
        "risks", "progress_text", "approval_status", "submitted_at",
        "responsible_person", "phone", "email", "file_names", "file_urls",
        "admin_comment"
    ])

measures_all = strat_df[strat_df["object_type"] == "measure"].copy()

st.markdown('<div class="card"><div class="card-title">Параметри звіту</div><div class="card-subtitle">Оберіть період і відповідальний підрозділ для формування матеріалів.</div>', unsafe_allow_html=True)

f1, f2, f3 = st.columns(3)

with f1:
    selected_year = st.selectbox("Рік", [2026, 2027, 2028])

with f2:
    selected_quarter = st.selectbox("Квартал", ["I", "II", "III", "IV"])

with f3:
    departments = sorted(measures_all["department"].dropna().astype(str).unique().tolist())
    selected_department = st.selectbox("Департамент", ["Усі"] + departments)

st.markdown('</div>', unsafe_allow_html=True)

active = prepare_period_data(
    strat_df,
    requests_df,
    selected_year,
    selected_quarter,
    selected_department
)

if active.empty:
    st.warning("Для обраного періоду активних заходів не знайдено.")
    st.stop()

total_active = len(active)
submitted_count = len(active[active["status"] != "Не подано"])
coverage = round(submitted_count / total_active * 100, 1) if total_active else 0
completion = round(active["performance_score"].mean(), 1) if total_active else 0
risk_count = len(active[active["auto_risk"].isin(["Критичний ризик", "Середній ризик"])])
critical_count = len(active[active["auto_risk"] == "Критичний ризик"])
without_data = len(active[active["status"] == "Не подано"])
forecast_q4 = forecast_to_q4(completion, selected_quarter)

metrics = {
    "Активних заходів": total_active,
    "Подано моніторинг": submitted_count,
    "Покриття моніторингом": f"{coverage}%",
    "Виконання СП": f"{completion}%",
    "Прогноз до IV кварталу": f"{forecast_q4}%",
    "Заходів у ризику": risk_count,
    "Критичних ризиків": critical_count,
    "Без поданих даних": without_data
}

st.markdown(f"""
<div class="alert-grid">
    <div class="alert-card alert-blue">
        <div class="alert-title">Виконання СП</div>
        <div class="alert-value">{completion}%</div>
    </div>
    <div class="alert-card alert-green">
        <div class="alert-title">Прогноз до IV кварталу</div>
        <div class="alert-value">{forecast_q4}%</div>
    </div>
    <div class="alert-card alert-yellow">
        <div class="alert-title">Покриття моніторингом</div>
        <div class="alert-value">{coverage}%</div>
    </div>
    <div class="alert-card alert-red">
        <div class="alert-title">Заходів у ризику</div>
        <div class="alert-value">{risk_count}</div>
    </div>
</div>
""", unsafe_allow_html=True)

goal_progress = (
    active
    .groupby(["goal_code", "strategic_goal"])
    .agg(
        Активних_заходів=("code", "count"),
        Виконання=("performance_score", "mean"),
        Ризикових=("auto_risk", lambda x: x.isin(["Критичний ризик", "Середній ризик"]).sum()),
        Середній_ризик=("risk_score", "mean")
    )
    .reset_index()
)

goal_progress["Виконання"] = goal_progress["Виконання"].round(1)
goal_progress["Середній_ризик"] = goal_progress["Середній_ризик"].round(1)

dep_progress = (
    active
    .groupby("department")
    .agg(
        Активних_заходів=("code", "count"),
        Виконання=("performance_score", "mean"),
        Подано=("status", lambda x: (x != "Не подано").sum()),
        Ризикових=("auto_risk", lambda x: x.isin(["Критичний ризик", "Середній ризик"]).sum()),
        Середній_ризик=("risk_score", "mean")
    )
    .reset_index()
)

dep_progress["Виконання"] = dep_progress["Виконання"].round(1)
dep_progress["Покриття_%"] = (dep_progress["Подано"] / dep_progress["Активних_заходів"] * 100).round(1)
dep_progress["Середній_ризик"] = dep_progress["Середній_ризик"].round(1)

risk_table = active[
    active["auto_risk"].isin(["Критичний ризик", "Середній ризик"])
].copy()

period_requests = requests_df.copy()

if not period_requests.empty:
    period_requests = period_requests[
        (period_requests["year"].astype(str) == str(selected_year)) &
        (period_requests["quarter"].astype(str) == str(selected_quarter))
    ].copy()

    if selected_department != "Усі":
        period_requests = period_requests[period_requests["department"].astype(str) == str(selected_department)]

analytical_text = generate_analytical_text(
    active,
    requests_df,
    selected_year,
    selected_quarter,
    selected_department
)

st.markdown("""
<div class="report-box">
    <div class="report-title">Автоматично сформована аналітична довідка</div>
    <div class="report-text">
""", unsafe_allow_html=True)

for paragraph in analytical_text.split("\n\n"):
    st.markdown(f"<p class='report-text'>{paragraph}</p>", unsafe_allow_html=True)

st.markdown("</div></div>", unsafe_allow_html=True)

st.markdown('<div class="card"><div class="card-title">Графіки до довідки</div><div class="card-subtitle">Візуальні матеріали для пояснення стану виконання та ризиків.</div>', unsafe_allow_html=True)

c1, c2 = st.columns(2)

with c1:
    fig = px.bar(
        goal_progress.sort_values("Виконання", ascending=False),
        x="strategic_goal",
        y="Виконання",
        text="Виконання",
        hover_data=["Активних_заходів", "Ризикових", "Середній_ризик"],
        title="Виконання за стратегічними цілями",
        labels={"strategic_goal": "Стратегічна ціль", "Виконання": "Виконання, %"}
    )
    fig.update_layout(xaxis_tickangle=-25)
    st.plotly_chart(fig, use_container_width=True)

with c2:
    risk_counts = active.groupby("auto_risk").size().reset_index(name="Кількість")
    fig = px.pie(
        risk_counts,
        names="auto_risk",
        values="Кількість",
        hole=0.45,
        title="Структура ризиків"
    )
    st.plotly_chart(fig, use_container_width=True)

c3, c4 = st.columns(2)

with c3:
    fig = px.bar(
        dep_progress.sort_values("Виконання", ascending=False),
        x="department",
        y="Виконання",
        text="Виконання",
        hover_data=["Активних_заходів", "Покриття_%", "Ризикових"],
        title="Виконання за департаментами",
        labels={"department": "Департамент", "Виконання": "Виконання, %"}
    )
    st.plotly_chart(fig, use_container_width=True)

with c4:
    status_counts = active.groupby("status").size().reset_index(name="Кількість")
    fig = px.pie(
        status_counts,
        names="status",
        values="Кількість",
        hole=0.45,
        title="Статуси активних заходів"
    )
    st.plotly_chart(fig, use_container_width=True)

st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="card"><div class="card-title">Експорт матеріалів</div><div class="card-subtitle">Завантажте звіти у потрібному форматі.</div>', unsafe_allow_html=True)

excel_file = create_excel_report(
    active,
    period_requests,
    risk_table,
    goal_progress,
    dep_progress,
    selected_year,
    selected_quarter,
    selected_department
)

docx_file = create_docx_report(
    analytical_text,
    metrics,
    selected_year,
    selected_quarter,
    selected_department
)

e1, e2, e3, e4 = st.columns(4)

with e1:
    st.download_button(
        "Завантажити Excel-звіт",
        data=excel_file,
        file_name=f"monitoring_report_{selected_year}_{selected_quarter}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )

with e2:
    requests_output = BytesIO()
    period_requests.to_excel(requests_output, index=False, engine="openpyxl")
    requests_output.seek(0)

    st.download_button(
        "Реєстр заявок",
        data=requests_output,
        file_name=f"requests_registry_{selected_year}_{selected_quarter}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )

with e3:
    risk_output = BytesIO()
    risk_table.to_excel(risk_output, index=False, engine="openpyxl")
    risk_output.seek(0)

    st.download_button(
        "Таблиця ризиків",
        data=risk_output,
        file_name=f"risk_table_{selected_year}_{selected_quarter}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )

with e4:
    st.download_button(
        "Аналітична довідка DOCX",
        data=docx_file,
        file_name=f"analytical_note_{selected_year}_{selected_quarter}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="card"><div class="card-title">Таблиці для перевірки</div>', unsafe_allow_html=True)

tab1, tab2, tab3, tab4 = st.tabs(["Активні заходи", "Ризики", "Реєстр заявок", "Показники за департаментами"])

with tab1:
    show_active = active.rename(columns={
        "code": "Код",
        "name": "Захід",
        "indicator": "Індикатор",
        "unit": "Одиниця виміру",
        "department": "Департамент",
        "selected_target": "Планове значення",
        "numeric_value": "Фактичне значення",
        "status": "Статус",
        "performance_score": "Виконання, %",
        "auto_risk": "Ризик",
        "risk_score": "Risk score",
        "risk_reason": "Причина ризику"
    })

    active_cols = [
        "Код", "Захід", "Індикатор", "Одиниця виміру", "Департамент",
        "Планове значення", "Фактичне значення", "Статус",
        "Виконання, %", "Ризик", "Risk score", "Причина ризику"
    ]

    st.dataframe(
        show_active[[c for c in active_cols if c in show_active.columns]],
        use_container_width=True,
        hide_index=True
    )

with tab2:
    if risk_table.empty:
        st.success("Ризикових заходів за обраний період не виявлено.")
    else:
        show_risk = risk_table.rename(columns={
            "code": "Код",
            "name": "Захід",
            "indicator": "Індикатор",
            "department": "Департамент",
            "selected_target": "Планове значення",
            "numeric_value": "Фактичне значення",
            "status": "Статус",
            "auto_risk": "Ризик",
            "risk_score": "Risk score",
            "risk_reason": "Причина ризику"
        })

        risk_cols = [
            "Код", "Захід", "Індикатор", "Департамент", "Планове значення",
            "Фактичне значення", "Статус", "Ризик", "Risk score", "Причина ризику"
        ]

        st.dataframe(
            show_risk[[c for c in risk_cols if c in show_risk.columns]],
            use_container_width=True,
            hide_index=True
        )

with tab3:
    st.dataframe(period_requests, use_container_width=True, hide_index=True)

with tab4:
    st.dataframe(dep_progress, use_container_width=True, hide_index=True)

st.markdown('</div>', unsafe_allow_html=True)

st.markdown("""
<div class="footer">
    Міністерство економіки, довкілля та сільського господарства України<br>
    Розроблено департаментом стратегічного планування та макроекономічного прогнозування<br>
    Версія DEMO 1.0 | Експорт звітів та автоматична аналітична довідка
</div>
""", unsafe_allow_html=True)
